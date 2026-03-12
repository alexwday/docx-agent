from __future__ import annotations

from docx import Document

from word_agent import WordAgent
from word_engine import EngineConfig, WordDocumentService


def test_end_to_end_agent_template_fill(tmp_path):
    target = tmp_path / "target.docx"
    support_a = tmp_path / "support-a.docx"
    support_b = tmp_path / "support-b.docx"

    doc = Document()
    doc.add_heading("Overview", level=1)
    doc.add_paragraph("TODO: write overview.")
    doc.add_heading("Risks", level=1)
    doc.add_paragraph("TODO: write risks.")
    doc.save(str(target))

    Document().save(str(support_a))
    Document().save(str(support_b))

    service = WordDocumentService(config=EngineConfig(allowed_roots=[tmp_path]))
    agent = WordAgent(service=service)

    plan = agent.plan_template_fill(
        target_doc=str(target),
        support_docs=[str(support_a), str(support_b)],
        objective="Create concise, executive-ready content",
    )
    assert plan["status"] == "ok"
    assert len(plan["section_plan"]) == 2

    enriched_plan = []
    for item in plan["section_plan"]:
        generated = agent.generate_section_content(
            section_plan_item=item,
            context_bundle={
                "objective": plan["objective"],
                "support_docs": plan["support_docs"],
            },
        )
        assert generated["status"] == "ok"
        item_copy = dict(item)
        item_copy["paragraphs"] = generated["paragraphs"]
        enriched_plan.append(item_copy)

    applied = agent.apply_section_plan(str(target), enriched_plan)
    assert applied["status"] == "ok"
    assert applied["applied_count"] == 2

    validation = agent.validate_document_result(str(target), ["Overview", "Risks"])
    assert validation["status"] == "ok"
    assert validation["all_expected_present"] is True

    updated = Document(str(target))
    all_text = "\n".join(paragraph.text for paragraph in updated.paragraphs)
    assert "Draft generated using support-a.docx, support-b.docx." in all_text
