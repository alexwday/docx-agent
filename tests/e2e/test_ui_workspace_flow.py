from __future__ import annotations

from pathlib import Path

from docx import Document

from word_agent import WordAgent
from word_engine import EngineConfig, WordDocumentService
from word_ui import WordUIWorkspace


def _build_target(path):
    doc = Document()
    doc.add_heading("Overview", level=1)
    doc.add_paragraph("TODO: overview")
    doc.add_heading("Risks", level=1)
    doc.add_paragraph("TODO: risks")
    doc.save(str(path))


def test_ui_workspace_state_flow(tmp_path):
    target = tmp_path / "target.docx"
    support = tmp_path / "support.docx"
    _build_target(target)
    Document().save(str(support))

    service = WordDocumentService(config=EngineConfig(allowed_roots=[tmp_path]))
    agent = WordAgent(service=service)
    workspace = WordUIWorkspace(agent=agent)

    created = workspace.create_session()
    assert created["status"] == "ok"
    session_id = created["session"]["session_id"]

    sent = workspace.send_message(session_id, "Fill this template with concise content.")
    assert sent["status"] == "ok"
    assert sent["message"]["role"] == "user"

    added_context = workspace.add_context_file(session_id, str(support))
    assert added_context["status"] == "ok"
    assert added_context["context_files"] == [str(support.resolve())]

    added_target = workspace.add_editable_target(session_id, str(target))
    assert added_target["status"] == "ok"
    assert added_target["editable_targets"] == [str(target.resolve())]

    selected = workspace.select_preview_file(session_id, str(target))
    assert selected["status"] == "ok"
    assert selected["preview"]["selected_file"] == str(target.resolve())

    planned = workspace.run_plan(session_id, objective="Generate executive-ready content")
    assert planned["status"] == "ok"
    assert planned["section_plan"]
    plan_id = planned["plan_id"]

    applied = workspace.apply_plan(session_id, plan_id)
    assert applied["status"] == "ok"
    assert applied["result"]["status"] == "ok"
    assert applied["result"]["applied_count"] == 2
    assert applied["session"]["preview"]["revision_id"] is not None
    assert applied["session"]["preview"]["artifact_format"] == "html"
    preview_artifact_path = applied["session"]["preview"]["artifact_path"]
    assert preview_artifact_path is not None
    assert Path(preview_artifact_path).exists()

    validated = workspace.validate_result(session_id, expected_sections=["Overview", "Risks"])
    assert validated["status"] == "ok"
    assert validated["result"]["status"] == "ok"
    assert validated["result"]["all_expected_present"] is True

    state = workspace.get_session_state(session_id)
    assert state["status"] == "ok"
    assert len(state["session"]["messages"]) >= 4
    assert state["session"]["preview"]["selected_file"] == str(target.resolve())

    updated_doc = Document(str(target))
    combined_text = "\n".join(paragraph.text for paragraph in updated_doc.paragraphs)
    assert "Draft generated using support.docx." in combined_text


def test_ui_workspace_requires_editable_target_for_run_plan(tmp_path):
    target = tmp_path / "target.docx"
    _build_target(target)

    service = WordDocumentService(config=EngineConfig(allowed_roots=[tmp_path]))
    workspace = WordUIWorkspace(agent=WordAgent(service=service))

    session_id = workspace.create_session()["session"]["session_id"]
    result = workspace.run_plan(
        session_id=session_id,
        objective="Generate content",
        target_doc=str(target),
    )
    assert result["status"] == "error"
    assert result["error_code"] == "INVALID_ARGUMENT"
