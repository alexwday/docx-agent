from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt


def build_sample_document(path: Path) -> None:
    doc = Document()
    doc.add_heading("Section A", level=1)
    doc.add_paragraph("Instruction A1")
    doc.add_paragraph("Instruction A2")
    doc.add_heading("Section B", level=1)
    doc.add_paragraph("Instruction B1")
    doc.save(str(path))


def build_styled_section_document(path: Path) -> None:
    doc = Document()
    doc.add_heading("Executive Summary", level=1)
    p1 = doc.add_paragraph("Replace this instruction block.")
    run = p1.runs[0]
    run.bold = True
    run.font.size = Pt(14)
    p2 = doc.add_paragraph("Second instruction line.")
    p2.runs[0].italic = True
    doc.add_heading("Next Section", level=1)
    doc.add_paragraph("Keep this content.")
    doc.save(str(path))


def paragraph_texts(path: Path) -> list[str]:
    doc = Document(str(path))
    return [paragraph.text for paragraph in doc.paragraphs]
