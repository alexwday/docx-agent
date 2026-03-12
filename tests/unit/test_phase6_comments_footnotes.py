from __future__ import annotations

from pathlib import Path
import zipfile

from docx import Document


COMMENTS_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:comment w:id="1" w:author="Alice" w:initials="AL" w:date="2026-02-28T10:00:00Z">
    <w:p><w:r><w:t>First comment</w:t></w:r></w:p>
  </w:comment>
  <w:comment w:id="2" w:author="Bob" w:initials="BO" w:date="2026-02-28T10:05:00Z">
    <w:p><w:r><w:t>Second comment</w:t></w:r></w:p>
  </w:comment>
</w:comments>
"""

FOOTNOTES_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:footnote w:id="-1" w:type="separator"><w:p><w:r><w:t>---</w:t></w:r></w:p></w:footnote>
  <w:footnote w:id="0" w:type="continuationSeparator"><w:p><w:r><w:t>---</w:t></w:r></w:p></w:footnote>
  <w:footnote w:id="1"><w:p><w:r><w:t>Footnote one</w:t></w:r></w:p></w:footnote>
  <w:footnote w:id="2"><w:p><w:r><w:t>Footnote two</w:t></w:r></w:p></w:footnote>
</w:footnotes>
"""


def _inject_docx_part(docx_path: Path, part_name: str, xml_content: str) -> None:
    temp_path = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(docx_path, "r") as source_zip:
        with zipfile.ZipFile(temp_path, "w") as target_zip:
            for info in source_zip.infolist():
                if info.filename == part_name:
                    continue
                target_zip.writestr(info, source_zip.read(info.filename))
            target_zip.writestr(part_name, xml_content.encode("utf-8"))
    temp_path.replace(docx_path)


def test_get_document_comments_no_comments_part(make_service, tmp_path):
    service = make_service(tmp_path)
    target = tmp_path / "no-comments.docx"
    Document().save(str(target))

    result = service.get_document_comments(str(target))
    assert result["status"] == "ok"
    assert result["experimental"] is True
    assert result["comments"] == []
    assert result["total_comments"] == 0


def test_get_document_comments_parses_entries(make_service, tmp_path):
    service = make_service(tmp_path)
    target = tmp_path / "with-comments.docx"
    Document().save(str(target))
    _inject_docx_part(target, "word/comments.xml", COMMENTS_XML)

    result = service.get_document_comments(str(target))
    assert result["status"] == "ok"
    assert result["experimental"] is True
    assert result["total_comments"] == 2
    assert result["comments"][0]["author"] == "Alice"
    assert result["comments"][0]["text"] == "First comment"
    assert result["comments"][1]["author"] == "Bob"
    assert result["comments"][1]["text"] == "Second comment"


def test_get_document_comments_invalid_xml_returns_error(make_service, tmp_path):
    service = make_service(tmp_path)
    target = tmp_path / "bad-comments.docx"
    Document().save(str(target))
    _inject_docx_part(target, "word/comments.xml", "<bad xml>")

    result = service.get_document_comments(str(target))
    assert result["status"] == "error"
    assert result["error_code"] == "DOCX_ERROR"


def test_get_document_footnotes_no_footnotes_part(make_service, tmp_path):
    service = make_service(tmp_path)
    target = tmp_path / "no-footnotes.docx"
    Document().save(str(target))

    result = service.get_document_footnotes(str(target))
    assert result["status"] == "ok"
    assert result["experimental"] is True
    assert result["footnotes"] == []
    assert result["total_footnotes"] == 0


def test_get_document_footnotes_parses_user_entries(make_service, tmp_path):
    service = make_service(tmp_path)
    target = tmp_path / "with-footnotes.docx"
    Document().save(str(target))
    _inject_docx_part(target, "word/footnotes.xml", FOOTNOTES_XML)

    result = service.get_document_footnotes(str(target))
    assert result["status"] == "ok"
    assert result["experimental"] is True
    assert result["total_footnotes"] == 2
    assert [item["id"] for item in result["footnotes"]] == ["1", "2"]
    assert [item["text"] for item in result["footnotes"]] == ["Footnote one", "Footnote two"]


def test_get_document_footnotes_invalid_xml_returns_error(make_service, tmp_path):
    service = make_service(tmp_path)
    target = tmp_path / "bad-footnotes.docx"
    Document().save(str(target))
    _inject_docx_part(target, "word/footnotes.xml", "<bad xml>")

    result = service.get_document_footnotes(str(target))
    assert result["status"] == "error"
    assert result["error_code"] == "DOCX_ERROR"
