"""Unit tests for OpenAI LLM integration in WordAgent."""

from __future__ import annotations

from pathlib import Path
from types import SimpleNamespace
from unittest.mock import MagicMock, patch

import pytest
from docx import Document

from word_agent import WordAgent
from word_engine import EngineConfig, WordDocumentService
from word_ui import WordUIWorkspace


def _make_mock_client(content: str = "Mock LLM response"):
    """Create a mock OpenAI client that returns a fixed response."""
    mock_client = MagicMock()
    mock_choice = SimpleNamespace(message=SimpleNamespace(content=content))
    mock_response = SimpleNamespace(choices=[mock_choice])
    mock_client.chat.completions.create.return_value = mock_response
    return mock_client


def _agent_with_mock(tmp_path: Path, content: str = "Mock LLM response") -> tuple[WordAgent, MagicMock]:
    """Create a WordAgent with a mocked OpenAI client."""
    service = WordDocumentService(config=EngineConfig(allowed_roots=[tmp_path]))
    agent = WordAgent(service=service, model="gpt-4.1", api_key="test-key")
    mock_client = _make_mock_client(content)
    agent._client = mock_client
    return agent, mock_client


# ── chat() tests ───────────────────────────────────────────────────────────


class TestChat:
    def test_chat_returns_string(self, tmp_path):
        agent, mock_client = _agent_with_mock(tmp_path, content="Hello! How can I help?")
        result = agent.chat([{"role": "user", "content": "hello"}])
        assert isinstance(result, str)
        assert result == "Hello! How can I help?"

    def test_chat_includes_system_context(self, tmp_path):
        agent, mock_client = _agent_with_mock(tmp_path, content="Got it.")
        agent.chat(
            [{"role": "user", "content": "hello"}],
            system_context="Active target: report.docx",
        )
        call_args = mock_client.chat.completions.create.call_args
        messages = call_args.kwargs["messages"]
        system_msg = messages[0]["content"]
        assert "Active target: report.docx" in system_msg

    def test_chat_fallback_on_api_error(self, tmp_path):
        agent, mock_client = _agent_with_mock(tmp_path)
        mock_client.chat.completions.create.side_effect = Exception("API down")
        result = agent.chat([{"role": "user", "content": "hello"}])
        assert isinstance(result, str)
        assert "trouble" in result.lower() or "try again" in result.lower()


# ── generate_section_content() tests ───────────────────────────────────────


class TestGenerateSectionContent:
    def test_llm_response_parsed_into_paragraphs(self, tmp_path):
        agent, mock_client = _agent_with_mock(
            tmp_path,
            content="First paragraph of content.\n\nSecond paragraph with more detail.",
        )
        result = agent.generate_section_content(
            {"heading_text": "Introduction", "instruction": "Write an intro"},
            {"objective": "Create a report", "support_docs": []},
        )
        assert result["status"] == "ok"
        assert len(result["paragraphs"]) == 2
        assert "First paragraph" in result["paragraphs"][0]
        assert "Second paragraph" in result["paragraphs"][1]

    def test_prompt_includes_heading_and_objective(self, tmp_path):
        agent, mock_client = _agent_with_mock(tmp_path, content="Some content.")
        agent.generate_section_content(
            {"heading_text": "Methodology", "instruction": "Describe the methods"},
            {"objective": "Write a research paper", "support_docs": []},
        )
        call_args = mock_client.chat.completions.create.call_args
        messages = call_args.kwargs["messages"]
        user_msg = messages[1]["content"]
        assert "Methodology" in user_msg
        assert "Write a research paper" in user_msg

    def test_fallback_on_api_error(self, tmp_path):
        agent, mock_client = _agent_with_mock(tmp_path)
        mock_client.chat.completions.create.side_effect = Exception("API down")
        result = agent.generate_section_content(
            {"heading_text": "Intro"},
            {"objective": "Write content", "support_docs": []},
        )
        assert result["status"] == "ok"
        assert len(result["paragraphs"]) == 3
        assert "Intro" in result["paragraphs"][0]

    def test_support_doc_content_included(self, tmp_path):
        doc_path = tmp_path / "ref.docx"
        doc = Document()
        doc.add_paragraph("Reference material about climate change.")
        doc.save(str(doc_path))

        agent, mock_client = _agent_with_mock(tmp_path, content="Generated content.")
        agent.generate_section_content(
            {"heading_text": "Background", "instruction": "Write background"},
            {"objective": "Report", "support_docs": [str(doc_path)]},
        )
        call_args = mock_client.chat.completions.create.call_args
        messages = call_args.kwargs["messages"]
        user_msg = messages[1]["content"]
        assert "Reference material" in user_msg or "ref.docx" in user_msg


# ── plan_template_fill() tests ─────────────────────────────────────────────


class TestPlanTemplateFill:
    def test_llm_instructions_used_when_available(self, tmp_path):
        target = tmp_path / "target.docx"
        doc = Document()
        doc.add_heading("Introduction", level=1)
        doc.add_heading("Methods", level=1)
        doc.save(str(target))

        agent, mock_client = _agent_with_mock(
            tmp_path,
            content="Write a compelling introduction\nDescribe the methodology used",
        )
        result = agent.plan_template_fill(str(target), [], "Write a paper")
        assert result["status"] == "ok"
        plan = result["section_plan"]
        assert len(plan) == 2
        assert plan[0]["instruction"] == "Write a compelling introduction"
        assert plan[1]["instruction"] == "Describe the methodology used"

    def test_fallback_instructions_on_llm_failure(self, tmp_path):
        target = tmp_path / "target.docx"
        doc = Document()
        doc.add_heading("Introduction", level=1)
        doc.save(str(target))

        agent, mock_client = _agent_with_mock(tmp_path)
        mock_client.chat.completions.create.side_effect = Exception("API error")
        result = agent.plan_template_fill(str(target), [], "Write a paper")
        assert result["status"] == "ok"
        plan = result["section_plan"]
        assert "Write a paper for section 'Introduction'" == plan[0]["instruction"]

    def test_fallback_on_line_count_mismatch(self, tmp_path):
        target = tmp_path / "target.docx"
        doc = Document()
        doc.add_heading("Section A", level=1)
        doc.add_heading("Section B", level=1)
        doc.save(str(target))

        # LLM returns wrong number of lines — should fall back
        agent, mock_client = _agent_with_mock(tmp_path, content="Only one instruction")
        result = agent.plan_template_fill(str(target), [], "Objective")
        assert result["status"] == "ok"
        plan = result["section_plan"]
        assert "Objective for section" in plan[0]["instruction"]
        assert "Objective for section" in plan[1]["instruction"]


# ── workspace chat_with_agent() tests ──────────────────────────────────────


class TestWorkspaceChatWithAgent:
    def _make_workspace(self, tmp_path: Path) -> WordUIWorkspace:
        service = WordDocumentService(config=EngineConfig(allowed_roots=[tmp_path]))
        agent = WordAgent(service=service, model="gpt-4.1", api_key="test-key")
        agent._client = _make_mock_client("Hello from the agent!")
        return WordUIWorkspace(agent=agent, allowed_roots=[tmp_path])

    def test_chat_with_agent_returns_ok(self, tmp_path):
        ws = self._make_workspace(tmp_path)
        session_id = ws.create_session()["session"]["session_id"]
        ws.send_message(session_id, "hello")
        result = ws.chat_with_agent(session_id, "hello")
        assert result["status"] == "ok"
        assert result["message"]["role"] == "assistant"
        assert result["message"]["text"] == "Hello from the agent!"

    def test_chat_with_agent_empty_text(self, tmp_path):
        ws = self._make_workspace(tmp_path)
        session_id = ws.create_session()["session"]["session_id"]
        result = ws.chat_with_agent(session_id, "   ")
        assert result["status"] == "error"

    def test_chat_with_agent_invalid_session(self, tmp_path):
        ws = self._make_workspace(tmp_path)
        result = ws.chat_with_agent("no-such-session", "hello")
        assert result["status"] == "error"
