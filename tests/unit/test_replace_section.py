from __future__ import annotations

from docx import Document

from tests.unit.helpers import build_styled_section_document, paragraph_texts
from word_engine.errors import ErrorCode
from word_engine.service import ServiceError


def test_replace_section_dry_run(make_service, tmp_path):
    service = make_service(tmp_path)
    target = tmp_path / "dry-run.docx"
    build_styled_section_document(target)

    result = service.replace_section_content(
        str(target),
        selector={"mode": "heading_exact", "value": "Executive Summary", "occurrence": 1},
        new_paragraphs=["Generated paragraph"],
        dry_run=True,
    )
    assert result["status"] == "ok"
    assert result["replaced_range"]["start_index"] == 1
    assert "Replace this instruction block." in result["preview"]["existing_paragraphs"][0]


def test_replace_section_preserve_style(make_service, tmp_path):
    service = make_service(tmp_path)
    target = tmp_path / "preserve.docx"
    build_styled_section_document(target)

    result = service.replace_section_content(
        str(target),
        selector={"mode": "heading_exact", "value": "Executive Summary", "occurrence": 1},
        new_paragraphs=["Generated summary line one", "Generated summary line two"],
        preserve_style=True,
        dry_run=False,
    )
    assert result["status"] == "ok"

    doc = Document(str(target))
    assert doc.paragraphs[1].text == "Generated summary line one"
    assert doc.paragraphs[1].runs[0].bold is True
    assert doc.paragraphs[2].text == "Generated summary line two"


def test_replace_section_anchors(make_service, tmp_path):
    service = make_service(tmp_path)
    target = tmp_path / "anchors.docx"
    doc = Document()
    doc.add_paragraph("BEGIN")
    doc.add_paragraph("Old line 1")
    doc.add_paragraph("Old line 2")
    doc.add_paragraph("END")
    doc.save(str(target))

    result = service.replace_section_content(
        str(target),
        selector={"mode": "anchors", "start_text": "BEGIN", "end_text": "END"},
        new_paragraphs=["New line A"],
        preserve_style=True,
    )
    assert result["status"] == "ok"
    lines = paragraph_texts(target)
    assert "Old line 1" not in lines
    assert "New line A" in lines


def test_replace_section_selector_not_found(make_service, tmp_path):
    service = make_service(tmp_path)
    target = tmp_path / "missing.docx"
    build_styled_section_document(target)

    result = service.replace_section_content(
        str(target),
        selector={"mode": "heading_exact", "value": "Does Not Exist", "occurrence": 1},
        new_paragraphs=["irrelevant"],
    )
    assert result["status"] == "error"
    assert result["error_code"] == "SELECTOR_NOT_FOUND"


def test_replace_section_atomic_failure_no_partial_write(make_service, tmp_path, monkeypatch):
    service = make_service(tmp_path)
    target = tmp_path / "atomic.docx"
    build_styled_section_document(target)
    before = paragraph_texts(target)

    def broken_save(*args, **kwargs):
        raise ServiceError(ErrorCode.DOCX_ERROR, "forced save failure")

    monkeypatch.setattr(service, "_save_document_atomic", broken_save)

    result = service.replace_section_content(
        str(target),
        selector={"mode": "heading_exact", "value": "Executive Summary", "occurrence": 1},
        new_paragraphs=["new content"],
    )
    assert result["status"] == "error"
    assert result["error_code"] == "DOCX_ERROR"

    after = paragraph_texts(target)
    assert before == after
