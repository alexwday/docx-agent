from __future__ import annotations

from pathlib import Path

from docx import Document

from word_agent import WordAgent
from word_engine import EngineConfig, WordDocumentService
from word_ui import WordUIWorkspace


def _new_workspace(tmp_path: Path, session_store: Path | None = None) -> WordUIWorkspace:
    service = WordDocumentService(config=EngineConfig(allowed_roots=[tmp_path]))
    return WordUIWorkspace(
        agent=WordAgent(service=service),
        allowed_roots=[tmp_path],
        session_store_path=str(session_store) if session_store is not None else None,
    )


def test_workspace_rejects_target_outside_allowed_roots(tmp_path):
    allowed_dir = tmp_path / "allowed"
    allowed_dir.mkdir()
    denied_dir = tmp_path / "denied"
    denied_dir.mkdir()
    target = denied_dir / "target.docx"
    Document().save(str(target))

    service = WordDocumentService(config=EngineConfig(allowed_roots=[allowed_dir]))
    workspace = WordUIWorkspace(
        agent=WordAgent(service=service),
        allowed_roots=[allowed_dir],
    )

    session_id = workspace.create_session()["session"]["session_id"]
    result = workspace.add_editable_target(session_id, str(target))
    assert result["status"] == "error"
    assert result["error_code"] == "PATH_NOT_ALLOWED"


def test_workspace_persists_and_recovers_sessions(tmp_path):
    target = tmp_path / "target.docx"
    support = tmp_path / "support.docx"
    Document().save(str(target))
    Document().save(str(support))
    store = tmp_path / "state" / "sessions.json"

    first = _new_workspace(tmp_path, session_store=store)
    created = first.create_session(session_id="session-a")
    session_id = created["session"]["session_id"]
    assert session_id == "session-a"
    assert first.add_context_file(session_id, str(support))["status"] == "ok"
    assert first.add_editable_target(session_id, str(target))["status"] == "ok"
    assert first.send_message(session_id, "Test persisted conversation.")["status"] == "ok"

    second = _new_workspace(tmp_path, session_store=store)
    recovered = second.get_session_state("session-a")
    assert recovered["status"] == "ok"
    session = recovered["session"]
    assert session["context_files"] == [str(support.resolve())]
    assert session["editable_targets"] == [str(target.resolve())]
    assert len(session["messages"]) == 1
    assert session["messages"][0]["text"] == "Test persisted conversation."
