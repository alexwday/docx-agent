from __future__ import annotations

from pathlib import Path

from docx import Document

from word_agent import WordAgent
from word_engine import EngineConfig, WordDocumentService
from word_ui import DocxPreviewRenderer, WordUIWorkspace


def _build_doc(path: Path) -> None:
    doc = Document()
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph("This paragraph is visible in preview.")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    doc.save(str(path))


def test_docx_preview_renderer_creates_html_artifact(tmp_path):
    target = tmp_path / "preview.docx"
    _build_doc(target)
    renderer = DocxPreviewRenderer()

    result = renderer.render_docx_to_html(str(target), revision_id="rev-1")
    assert result["artifact_format"] == "html"
    artifact_path = Path(result["artifact_path"])
    assert artifact_path.exists()

    html_text = artifact_path.read_text(encoding="utf-8")
    assert "<h1>Executive Summary</h1>" in html_text
    assert "This paragraph is visible in preview." in html_text
    assert "<h2>Tables</h2>" in html_text


def test_workspace_refresh_preview_returns_artifact_metadata(tmp_path):
    target = tmp_path / "target.docx"
    _build_doc(target)

    service = WordDocumentService(config=EngineConfig(allowed_roots=[tmp_path]))
    workspace = WordUIWorkspace(agent=WordAgent(service=service))
    session_id = workspace.create_session()["session"]["session_id"]
    workspace.add_editable_target(session_id, str(target))

    refresh = workspace.refresh_preview(session_id, str(target), revision_id="r1")
    assert refresh["status"] == "ok"
    assert refresh["preview"]["artifact_format"] == "html"
    assert Path(refresh["preview"]["artifact_path"]).exists()
    assert refresh["preview_artifact"]["paragraph_count"] >= 2


def test_workspace_refresh_preview_reports_render_failure(tmp_path):
    target = tmp_path / "target.docx"
    _build_doc(target)
    service = WordDocumentService(config=EngineConfig(allowed_roots=[tmp_path]))
    workspace = WordUIWorkspace(agent=WordAgent(service=service))
    session_id = workspace.create_session()["session"]["session_id"]
    assert workspace.add_editable_target(session_id, str(target))["status"] == "ok"
    target.unlink()

    refresh = workspace.refresh_preview(session_id, str(target), revision_id="r2")
    assert refresh["status"] == "error"
    assert refresh["error_code"] in {"FILE_NOT_FOUND", "PREVIEW_RENDER_FAILED"}
