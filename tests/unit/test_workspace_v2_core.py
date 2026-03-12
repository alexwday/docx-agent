from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document

from word_ui.workspace_v2 import WordUIWorkspaceV2


class _FakeSessionsRepo:
    def __init__(self) -> None:
        self._sessions: dict[str, dict[str, Any]] = {}

    def get_session(self, session_id: str):
        if session_id == "missing":
            return None
        if session_id not in self._sessions:
            self._sessions[session_id] = {"session_id": session_id, "user_id": "123456789", "metadata": {}}
        return self._sessions[session_id]

    def update_session(self, session_id: str, *, title=None, status=None, metadata=None):
        session = self.get_session(session_id)
        if session is None:
            return None
        if title is not None:
            session["title"] = title
        if status is not None:
            session["status"] = status
        if metadata is not None:
            session["metadata"] = metadata
        return dict(session)


class _FakeMessagesRepo:
    def __init__(self) -> None:
        self._messages: list[dict[str, Any]] = []
        self._counter = 0

    def create_message(
        self,
        session_id: str,
        *,
        role: str,
        content_text: str | None,
        content_json: dict[str, Any] | None = None,
        parent_message_id: str | None = None,
        processing_state: str = "completed",
        processing_started_at=None,
        processing_ended_at=None,
        error=None,
    ):
        self._counter += 1
        msg = {
            "message_id": f"m{self._counter}",
            "session_id": session_id,
            "sequence_no": self._counter,
            "role": role,
            "content_text": content_text,
            "content_json": content_json or {},
            "parent_message_id": parent_message_id,
            "processing_state": processing_state,
            "processing_started_at": processing_started_at,
            "processing_ended_at": processing_ended_at,
            "error": error,
            "created_at": "2026-03-02T00:00:00Z",
        }
        self._messages.append(msg)
        return dict(msg)

    def list_messages(self, session_id: str, *, limit: int = 120, after_sequence_no=None):
        rows = [m for m in self._messages if m["session_id"] == session_id]
        return rows[-limit:]

    def update_message_content_and_state(
        self,
        session_id: str,
        message_id: str,
        *,
        content_text: str | None = None,
        content_json: dict[str, Any] | None = None,
        processing_state: str | None = None,
        processing_started_at=None,
        processing_ended_at=None,
        error=None,
    ):
        for msg in self._messages:
            if msg["session_id"] == session_id and msg["message_id"] == message_id:
                if content_text is not None:
                    msg["content_text"] = content_text
                if content_json is not None:
                    msg["content_json"] = content_json
                if processing_state is not None:
                    msg["processing_state"] = processing_state
                if processing_started_at is not None:
                    msg["processing_started_at"] = processing_started_at
                if processing_ended_at is not None:
                    msg["processing_ended_at"] = processing_ended_at
                msg["error"] = error
                return dict(msg)
        return None

    def get_message(self, session_id: str, message_id: str):
        for msg in self._messages:
            if msg["session_id"] == session_id and msg["message_id"] == message_id:
                return dict(msg)
        return None


class _FakeEventsRepo:
    def __init__(self) -> None:
        self.events: list[dict[str, Any]] = []

    def create_event(self, session_id: str, message_id: str, *, event_type: str, payload: dict[str, Any], event_index=None):
        row = {
            "event_id": f"e{len(self.events)+1}",
            "session_id": session_id,
            "message_id": message_id,
            "event_type": event_type,
            "payload": payload,
        }
        self.events.append(row)
        return row

    def list_events(self, session_id: str, message_id: str):
        return [e for e in self.events if e["session_id"] == session_id and e["message_id"] == message_id]

    def list_recent_events(
        self,
        session_id: str,
        *,
        limit: int = 200,
        exclude_message_id: str | None = None,
        event_types: list[str] | None = None,
    ):
        filtered = [e for e in self.events if e["session_id"] == session_id]
        if exclude_message_id is not None:
            filtered = [e for e in filtered if e["message_id"] != exclude_message_id]
        if event_types:
            wanted = set(event_types)
            filtered = [e for e in filtered if e["event_type"] in wanted]
        return list(reversed(filtered))[:limit]


class _FakeArtifactsRepo:
    def __init__(self) -> None:
        self.artifacts: dict[str, dict[str, Any]] = {}
        self._counter = 0

    def create_artifact(self, session_id: str, **kwargs):
        self._counter += 1
        artifact_id = f"a{self._counter}"
        row = {"artifact_id": artifact_id, "session_id": session_id, **kwargs}
        self.artifacts[artifact_id] = row
        return dict(row)

    def update_artifact(self, session_id: str, artifact_id: str, *, lifecycle_state=None, metadata=None):
        row = self.artifacts[artifact_id]
        if lifecycle_state is not None:
            row["lifecycle_state"] = lifecycle_state
        if metadata is not None:
            row["metadata"] = metadata
        return dict(row)

    def get_artifact(self, session_id: str, artifact_id: str):
        return dict(self.artifacts[artifact_id])

    def list_artifacts(self, session_id: str, *, artifact_type=None, limit=200):
        rows = [r for r in self.artifacts.values() if r["session_id"] == session_id]
        if artifact_type is not None:
            rows = [r for r in rows if r["artifact_type"] == artifact_type]
        return rows[:limit]

    def list_artifact_panes(self, session_id: str):
        return {"uploaded_documents": [], "research_outputs": [], "report_documents": []}


class _FakeKnowledgeRepo:
    def __init__(self) -> None:
        self.units: list[dict[str, Any]] = []

    def create_knowledge_unit(self, session_id: str, artifact_id: str, *, unit_type: str, content: str, sequence_no: int = 0, metadata=None):
        row = {
            "knowledge_id": f"k{len(self.units)+1}",
            "session_id": session_id,
            "artifact_id": artifact_id,
            "unit_type": unit_type,
            "content": content,
            "sequence_no": sequence_no,
            "metadata": metadata or {},
        }
        self.units.append(row)
        return row

    def list_knowledge_units(self, session_id: str, *, artifact_id=None, unit_type=None, limit=500):
        rows = [u for u in self.units if u["session_id"] == session_id]
        if artifact_id is not None:
            rows = [u for u in rows if u["artifact_id"] == artifact_id]
        if unit_type is not None:
            rows = [u for u in rows if u["unit_type"] == unit_type]
        return rows[:limit]


class _FakeDataSourcesRepo:
    def list_sources(self, *, enabled_only=True, source_type=None, source_ids=None):
        rows = [
            {
                "source_id": "risk_db.suppliers",
                "name": "Suppliers",
                "source_type": "postgres_table",
                "location": {
                    "schema": "risk_db",
                    "table": "suppliers",
                    "research_output_format": "xlsx",
                },
                "schema_json": {"fields": [{"name": "supplier_name"}, {"name": "risk_score"}]},
            },
            {
                "source_id": "sales_db.orders",
                "name": "Orders",
                "source_type": "postgres_table",
                "location": {
                    "schema": "sales_db",
                    "table": "orders",
                    "research_output_format": "pdf",
                },
                "schema_json": {"fields": [{"name": "order_id"}, {"name": "status"}]},
            },
        ]
        if source_type is not None:
            rows = [r for r in rows if r["source_type"] == source_type]
        if source_ids:
            wanted = set(source_ids)
            rows = [r for r in rows if r["source_id"] in wanted]
        return rows


class _FakeAgent:
    model = "gpt-4.1"

    def chat(self, messages, system_context=""):
        return f"assistant-response ({len(messages)} messages)"


class _PlanningAgent:
    model = "gpt-4.1"

    def __init__(self) -> None:
        self._iterations = 0

    def chat(self, messages, system_context=""):
        content = ""
        if messages and isinstance(messages[0], dict):
            content = str(messages[0].get("content") or "")
        if "INTERNAL_RESEARCH_ITERATION_JSON" in content:
            if self._iterations == 0:
                self._iterations += 1
                return (
                    "{"
                    "\"action\":\"call_more\","
                    "\"source_queries\":[{\"source_id\":\"sales_db.orders\","
                    "\"research_statement\":\"Investigate delayed orders by region and root causes.\"}],"
                    "\"reasoning\":\"Orders source is most relevant for delay analysis.\""
                    "}"
                )
            return "{\"action\":\"finish\",\"source_queries\":[],\"reasoning\":\"No further sources needed.\"}"
        return "assistant-response (planned)"


class _FollowupPlanningAgent:
    model = "gpt-4.1"

    def __init__(self) -> None:
        self._iteration = 0

    def chat(self, messages, system_context=""):
        content = ""
        if messages and isinstance(messages[0], dict):
            content = str(messages[0].get("content") or "")
        if "INTERNAL_RESEARCH_ITERATION_JSON" in content:
            if self._iteration == 0:
                self._iteration += 1
                return (
                    "{"
                    "\"action\":\"call_more\","
                    "\"source_queries\":[{\"source_id\":\"risk_db.suppliers\","
                    "\"research_statement\":\"Assess supplier risk concentration and critical entities.\"}],"
                    "\"reasoning\":\"Start with supplier risk baseline.\""
                    "}"
                )
            if self._iteration == 1:
                self._iteration += 1
                return (
                    "{"
                    "\"action\":\"call_more\","
                    "\"source_queries\":[{\"source_id\":\"sales_db.orders\","
                    "\"research_statement\":\"Assess delayed order impacts linked to supplier risk.\"}],"
                    "\"reasoning\":\"Add order delays to cross-check downstream impact.\""
                    "}"
                )
            return "{\"action\":\"finish\",\"source_queries\":[],\"reasoning\":\"Sufficient coverage.\"}"
        return "assistant-response (followup)"


class _ReportGapPlannerAgent:
    model = "gpt-4.1"

    def __init__(self) -> None:
        self._gap_calls = 0

    def chat(self, messages, system_context=""):
        content = ""
        if messages and isinstance(messages[0], dict):
            content = str(messages[0].get("content") or "")
        if "REPORT_GAP_FILL_PLAN_JSON" in content:
            if self._gap_calls == 0:
                self._gap_calls += 1
                return (
                    "{"
                    "\"action\":\"call_more\","
                    "\"calls\":[{"
                    "\"section_key\":\"Executive Summary::Purpose\","
                    "\"source_id\":\"risk_db.suppliers\","
                    "\"research_statement\":\"Identify supplier concentration risks for executive summary priorities.\""
                    "}],"
                    "\"reasoning\":\"Need focused risk evidence for executive summary purpose.\""
                    "}"
                )
            return "{\"action\":\"finish\",\"calls\":[],\"reasoning\":\"Coverage now sufficient.\"}"
        if "INTERNAL_RESEARCH_ITERATION_JSON" in content:
            return "{\"action\":\"finish\",\"source_queries\":[],\"reasoning\":\"No internal follow-up needed.\"}"
        return "assistant-response (report-gap-planned)"


def _new_workspace(tmp_path: Path) -> WordUIWorkspaceV2:
    ws = object.__new__(WordUIWorkspaceV2)
    ws.sessions = _FakeSessionsRepo()
    ws.messages = _FakeMessagesRepo()
    ws.events = _FakeEventsRepo()
    ws.artifacts = _FakeArtifactsRepo()
    ws.knowledge = _FakeKnowledgeRepo()
    ws.data_sources = _FakeDataSourcesRepo()
    ws.agent = _FakeAgent()
    ws.preview_renderer = None
    ws.allowed_roots = [tmp_path.resolve()]
    ws.contract_version = "v2"
    return ws


def test_respond_uses_filtered_sources_and_logs_events(tmp_path: Path):
    ws = _new_workspace(tmp_path)
    result = ws.respond(
        "s1",
        message="Need supplier risk analysis",
        data_source_filters=["risk_db.suppliers", "unknown.source"],
    )
    assert result["status"] == "ok"
    assert result["filter_result"]["mode"] == "filtered"
    assert result["filter_result"]["effective_source_ids"] == ["risk_db.suppliers"]
    assert result["filter_result"]["ignored_source_ids"] == ["unknown.source"]
    assert result["assistant_message"]["processing_state"] == "completed"
    assert result["artifacts_created"]
    event_types = [e["event_type"] for e in ws.events.events]
    assert "ui_data_source_filter_applied" in event_types
    assert "uploaded_documents_context_injected" in event_types
    assert "prior_agent_interactions_injected" in event_types
    assert "model_response" in event_types
    assert "artifact_created" in event_types
    assert "tool_call_request" in event_types
    assert "tool_call_response" in event_types

    artifact_types = [artifact["artifact_type"] for artifact in ws.artifacts.artifacts.values()]
    assert "research_markdown" in artifact_types
    assert "research_output_doc" in artifact_types

    tool_requests = [e for e in ws.events.events if e["event_type"] == "tool_call_request"]
    tool_responses = [e for e in ws.events.events if e["event_type"] == "tool_call_response"]
    tool_names = {e["payload"]["tool_name"] for e in tool_requests}
    assert "database_research" in tool_names
    assert "research_internal_data_sources" in tool_names
    assert "research_uploaded_documents" in tool_names
    assert len(tool_requests) == len(tool_responses)

    internal_payloads = [
        payload
        for payload in (event["payload"] for event in tool_responses)
        if payload.get("tool_name") == "research_internal_data_sources"
    ]
    assert internal_payloads
    internal_result = internal_payloads[-1]["result"]
    assert internal_result["status"] == "completed"
    assert internal_result["finding_count"] >= 1
    assert internal_result["artifacts_created"]
    assert internal_result["source_results"]
    first_finding = internal_result["findings"][0]
    assert "execution" in first_finding
    assert first_finding["execution"]["handler"] in (
        "mock_supplier_risk_retriever",
        "mock_sales_orders_retriever",
    )
    first_source_result = internal_result["source_results"][0]
    assert first_source_result["research_markdown"] is not None
    assert first_source_result["research_document"] is not None
    assert first_source_result["research_document"]["format"] == "xlsx"

    db_requests = [
        event["payload"]
        for event in tool_requests
        if event["payload"].get("tool_name") == "database_research"
    ]
    assert db_requests
    request_args = db_requests[0]["arguments"]
    assert request_args["source_id"] == "risk_db.suppliers"
    assert "supplier risk" in request_args["research_statement"].lower()

    db_responses = [
        event["payload"]
        for event in tool_responses
        if event["payload"].get("tool_name") == "database_research"
    ]
    assert db_responses
    db_result = db_responses[0]["result"]
    assert db_result["research_markdown"] is not None
    assert db_result["research_document"] is not None
    assert db_result["research_document"]["format"] == "xlsx"
    assert Path(db_result["research_document"]["storage_uri"]).read_bytes().startswith(b"PK")


def test_model_planner_selects_sources_and_research_statements(tmp_path: Path):
    ws = _new_workspace(tmp_path)
    ws.agent = _PlanningAgent()

    result = ws.respond("s1", message="Investigate delayed order performance issues")
    assert result["status"] == "ok"
    assert result["selected_source_ids"] == ["sales_db.orders"]

    selected_events = [
        event
        for event in ws.events.events
        if event["event_type"] == "agent_data_source_selected"
    ]
    assert selected_events
    selection_payload = selected_events[-1]["payload"]
    assert selection_payload["selection_mode"] == "model_iterative"
    assert selection_payload["selected_source_ids"] == ["sales_db.orders"]

    db_requests = [
        event["payload"]
        for event in ws.events.events
        if event["event_type"] == "tool_call_request"
        and event["payload"].get("tool_name") == "database_research"
    ]
    assert db_requests
    args = db_requests[0]["arguments"]
    assert args["source_id"] == "sales_db.orders"
    assert args["research_statement"] == "Investigate delayed orders by region and root causes."


def test_model_planner_followup_executes_multistep_calls(tmp_path: Path):
    ws = _new_workspace(tmp_path)
    ws.agent = _FollowupPlanningAgent()

    result = ws.respond("s1", message="Cross-analyze supplier risk and order delays")
    assert result["status"] == "ok"
    assert set(result["selected_source_ids"]) == {"risk_db.suppliers", "sales_db.orders"}

    selected_events = [
        event
        for event in ws.events.events
        if event["event_type"] == "agent_data_source_selected"
    ]
    assert selected_events
    payload = selected_events[-1]["payload"]
    assert payload["selection_mode"] == "model_iterative"
    assert payload["planned_call_count"] >= 2
    assert payload["executed_call_count"] >= 2

    db_requests = [
        event["payload"]
        for event in ws.events.events
        if event["event_type"] == "tool_call_request"
        and event["payload"].get("tool_name") == "database_research"
    ]
    assert len(db_requests) >= 2
    statements_by_source = {
        request["arguments"]["source_id"]: request["arguments"]["research_statement"]
        for request in db_requests
    }
    assert "risk_db.suppliers" in statements_by_source
    assert "sales_db.orders" in statements_by_source
    assert "supplier risk" in statements_by_source["risk_db.suppliers"].lower()
    assert "order" in statements_by_source["sales_db.orders"].lower()


def test_second_turn_injects_prior_agent_interactions(tmp_path: Path):
    ws = _new_workspace(tmp_path)
    first = ws.respond("s1", message="Need supplier risk analysis")
    assert first["status"] == "ok"

    second = ws.respond("s1", message="Continue with mitigation actions")
    assert second["status"] == "ok"

    prior_events = [
        event
        for event in ws.events.events
        if event["event_type"] == "prior_agent_interactions_injected"
    ]
    assert len(prior_events) >= 2
    assert prior_events[-1]["payload"]["interaction_count"] >= 1


def test_database_research_uses_source_preset_pdf_output(tmp_path: Path):
    ws = _new_workspace(tmp_path)
    result = ws.respond("s1", message="Analyze delayed orders in sales_db orders")
    assert result["status"] == "ok"

    db_responses = [
        event["payload"]
        for event in ws.events.events
        if event["event_type"] == "tool_call_response"
        and event["payload"].get("tool_name") == "database_research"
    ]
    assert db_responses
    sales_results = [
        item["result"] for item in db_responses if item["result"].get("source_id") == "sales_db.orders"
    ]
    assert sales_results
    sales_result = sales_results[-1]
    assert sales_result["research_document"] is not None
    assert sales_result["research_document"]["format"] == "pdf"
    assert Path(sales_result["research_document"]["storage_uri"]).read_bytes().startswith(b"%PDF-")


def test_uploaded_document_tool_call_returns_completed_when_context_exists(tmp_path: Path):
    ws = _new_workspace(tmp_path)
    ws.knowledge.create_knowledge_unit(
        "s1",
        "a-upload",
        unit_type="summary",
        sequence_no=0,
        content="Supplier risk trend indicates delayed shipments and rising defect rates.",
    )

    result = ws.respond("s1", message="Summarize supplier risk from uploaded docs")
    assert result["status"] == "ok"

    tool_responses = [
        event["payload"]
        for event in ws.events.events
        if event["event_type"] == "tool_call_response"
    ]
    uploaded_payloads = [
        payload
        for payload in tool_responses
        if payload.get("tool_name") == "research_uploaded_documents"
    ]
    assert uploaded_payloads
    uploaded_result = uploaded_payloads[-1]["result"]
    assert uploaded_result["status"] == "completed"
    assert uploaded_result["unit_count"] >= 1
    assert uploaded_result["artifact_count"] >= 1


def test_upload_artifact_ingests_text_into_knowledge_units(tmp_path: Path):
    ws = _new_workspace(tmp_path)
    file_path = tmp_path / "upload.txt"
    file_path.write_text("alpha beta gamma " * 300, encoding="utf-8")

    result = ws.upload_artifact("s1", file_path=str(file_path), artifact_type="upload")
    assert result["status"] == "ok"
    assert result["ingestion_state"] == "completed"
    assert result["ingestion_summary"]["summary_units"] == 1
    assert result["ingestion_summary"]["chunk_units"] >= 1
    unit_types = [unit["unit_type"] for unit in ws.knowledge.units]
    assert "summary" in unit_types
    assert "chunk" in unit_types


def test_report_workflow_runs_multi_turn_and_creates_final_artifact(tmp_path: Path):
    ws = _new_workspace(tmp_path)

    first = ws.respond("s1", message="Create a report based on supplier risk trends")
    assert first["status"] == "ok"
    assert first["artifacts_created"]
    assert "Suggested primary sections" in first["assistant_message"]["content_text"]

    second = ws.respond("s1", message="use suggested")
    assert second["status"] == "ok"
    assert "Primary sections confirmed" in second["assistant_message"]["content_text"]

    third = ws.respond("s1", message="use suggested")
    assert third["status"] == "ok"
    assert "instructions" in third["assistant_message"]["content_text"].lower()

    fourth = ws.respond("s1", message="use defaults")
    assert fourth["status"] == "ok"
    assert "Report generation completed." in fourth["assistant_message"]["content_text"]
    assert fourth["artifacts_created"]

    event_types = [event["event_type"] for event in ws.events.events]
    assert "report_structure_proposed" in event_types
    assert "report_structure_confirmed" in event_types
    assert "report_section_instructions_captured" in event_types
    assert "report_generation_started" in event_types
    assert "report_generation_completed" in event_types
    assert "tool_call_request" in event_types
    assert "tool_call_response" in event_types

    generate_requests = [
        event["payload"]
        for event in ws.events.events
        if event["event_type"] == "tool_call_request"
        and event["payload"].get("tool_name") == "generate_report_document"
    ]
    export_requests = [
        event["payload"]
        for event in ws.events.events
        if event["event_type"] == "tool_call_request"
        and event["payload"].get("tool_name") == "export_report_document"
    ]
    generate_responses = [
        event["payload"]
        for event in ws.events.events
        if event["event_type"] == "tool_call_response"
        and event["payload"].get("tool_name") == "generate_report_document"
    ]
    export_responses = [
        event["payload"]
        for event in ws.events.events
        if event["event_type"] == "tool_call_response"
        and event["payload"].get("tool_name") == "export_report_document"
    ]
    assert generate_requests
    assert export_requests
    assert generate_responses
    assert export_responses
    assert generate_responses[-1]["result"]["status"] == "completed"
    assert generate_responses[-1]["result"]["section_count"] >= 1
    gap_fill_calls = int(generate_responses[-1]["result"]["gap_fill_call_count"])
    low_quality_remaining = int(generate_responses[-1]["result"].get("low_quality_section_count", 0))
    assert gap_fill_calls >= 0
    assert low_quality_remaining >= 0
    assert gap_fill_calls >= 1 or low_quality_remaining == 0
    assert export_responses[-1]["result"]["status"] == "completed"
    assert len(export_responses[-1]["result"]["artifacts_created"]) >= 1
    export_artifacts = [
        item for item in ws.artifacts.artifacts.values() if item["artifact_type"] == "export_file"
    ]
    assert export_artifacts
    export_formats = {item["format"] for item in export_artifacts}
    assert {"docx", "pdf", "xlsx"}.issubset(export_formats)
    pdf_exports = [item for item in export_artifacts if item["format"] == "pdf"]
    xlsx_exports = [item for item in export_artifacts if item["format"] == "xlsx"]
    assert pdf_exports
    assert xlsx_exports
    assert Path(pdf_exports[0]["storage_uri"]).read_bytes().startswith(b"%PDF-")
    assert Path(xlsx_exports[0]["storage_uri"]).read_bytes().startswith(b"PK")


def test_report_workflow_emits_plan_card_in_content_json(tmp_path: Path):
    ws = _new_workspace(tmp_path)

    first = ws.respond("s1", message="Create a report based on supplier risk trends")
    assert first["status"] == "ok"
    content_json = first["assistant_message"]["content_json"]
    assert "report_plan_card" in content_json
    plan_card = content_json["report_plan_card"]
    assert plan_card["status"] == "scaffolding"
    assert plan_card["plan_id"]
    assert plan_card["sections"]
    first_section = plan_card["sections"][0]
    assert first_section["section_id"]
    assert first_section["depth"] == 0

    event_types = [event["event_type"] for event in ws.events.events]
    assert "report_plan_card_created" in event_types
    assert "report_plan_state_updated" in event_types


def test_report_workflow_applies_manual_report_plan_state_updates(tmp_path: Path):
    ws = _new_workspace(tmp_path)

    first = ws.respond("s1", message="Create a report for supplier resilience")
    assert first["status"] == "ok"
    plan_card = first["assistant_message"]["content_json"]["report_plan_card"]
    assert plan_card["sections"]
    edited_plan = dict(plan_card)
    edited_sections = list(plan_card["sections"])
    edited_sections.append(
        {
            "section_id": "custom-section-1",
            "title": "Financial Risk",
            "depth": 0,
            "instructions": "Focus on liquidity and margin exposure.",
            "instruction_source": "user",
            "status": "has_instructions",
            "subsections": [],
        }
    )
    edited_plan["sections"] = edited_sections
    edited_plan["updated_by"] = "user"

    second = ws.respond("s1", message="use suggested", report_plan_state=edited_plan)
    assert second["status"] == "ok"
    next_card = second["assistant_message"]["content_json"]["report_plan_card"]
    titles = [item["title"] for item in next_card["sections"]]
    assert "Financial Risk" in titles

    state_update_events = [
        event
        for event in ws.events.events
        if event["event_type"] == "report_plan_state_updated"
        and event["payload"].get("reason") == "user_plan_state_applied"
    ]
    assert state_update_events


def test_report_workflow_start_now_action_triggers_generation(tmp_path: Path):
    ws = _new_workspace(tmp_path)

    first = ws.respond("s1", message="Create a report on supplier risk and order delays")
    assert first["status"] == "ok"
    plan_card = dict(first["assistant_message"]["content_json"]["report_plan_card"])
    plan_card["status"] = "ready"
    plan_card["updated_by"] = "user"

    started = ws.respond("s1", report_plan_state=plan_card, report_plan_action="start_now")
    assert started["status"] == "ok"
    assert "Report generation completed." in started["assistant_message"]["content_text"]
    final_card = started["assistant_message"]["content_json"]["report_plan_card"]
    assert final_card["status"] == "completed"

    event_types = [event["event_type"] for event in ws.events.events]
    assert "report_plan_start_now_triggered" in event_types
    assert "report_generation_completed" in event_types

    instruction_events = [
        event for event in ws.events.events if event["event_type"] == "report_section_instructions_captured"
    ]
    assert instruction_events
    assert instruction_events[-1]["payload"]["mode"] in ("start_now_defaults", "start_now_existing")


def test_report_content_map_includes_cross_section_dependencies(tmp_path: Path):
    ws = _new_workspace(tmp_path)
    ws.respond("s1", message="Create a report for supplier resilience")
    ws.respond("s1", message="use suggested")
    ws.respond("s1", message="use suggested")
    final = ws.respond("s1", message="use defaults")
    assert final["status"] == "ok"

    working_artifacts = [
        item for item in ws.artifacts.artifacts.values() if item["artifact_type"] == "report_working_doc"
    ]
    assert working_artifacts
    working_path = Path(working_artifacts[-1]["storage_uri"])
    doc = Document(str(working_path))
    combined_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "Cross-section dependencies:" in combined_text


def test_internal_research_call_normalization_replaces_generic_statement():
    calls = WordUIWorkspaceV2._normalize_internal_research_calls(
        calls=[{"source_id": "risk_db.suppliers", "research_statement": "research this"}],
        allowed_source_ids=["risk_db.suppliers"],
        default_statement="Analyze supplier concentration and risk signals for top vendors.",
    )
    assert calls
    assert calls[0]["research_statement"] == "Analyze supplier concentration and risk signals for top vendors."


def test_report_gap_fill_call_normalization_replaces_generic_statement(tmp_path: Path):
    ws = _new_workspace(tmp_path)
    workflow = {
        "topic": "Supplier Risk",
        "section_hierarchy": {"": ["Executive Summary"], "Executive Summary": []},
    }
    normalized = ws._normalize_report_gap_fill_calls(
        calls=[
            {
                "section_key": "Executive Summary",
                "source_id": "risk_db.suppliers",
                "research_statement": "analyze this",
            }
        ],
        allowed_section_keys=["Executive Summary"],
        allowed_source_ids=["risk_db.suppliers"],
        executed_signatures=set(),
        workflow=workflow,
    )
    assert normalized
    assert "section 'Executive Summary'" in normalized[0]["research_statement"]


def test_report_workflow_accepts_structured_section_instruction_mapping(tmp_path: Path):
    ws = _new_workspace(tmp_path)

    ws.respond("s1", message="Create a report for supplier resilience")
    ws.respond("s1", message="use suggested")
    ws.respond("s1", message="use suggested")
    final = ws.respond(
        "s1",
        message=(
            "Executive Summary: Purpose -> Prioritize top 3 supplier risks and immediate mitigation actions.\n"
            "all -> Keep the tone concise and executive-ready."
        ),
    )

    assert final["status"] == "ok"
    assert "Report generation completed." in final["assistant_message"]["content_text"]

    captured_events = [
        event
        for event in ws.events.events
        if event["event_type"] == "report_section_instructions_captured"
    ]
    assert captured_events
    payload = captured_events[-1]["payload"]
    assert payload["mode"] == "structured"
    assert any("Executive Summary::Purpose" == key for key in payload["provided_keys"])

    working_artifacts = [
        item for item in ws.artifacts.artifacts.values() if item["artifact_type"] == "report_working_doc"
    ]
    assert working_artifacts
    working_path = Path(working_artifacts[-1]["storage_uri"])
    doc = Document(str(working_path))
    combined_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "Prioritize top 3 supplier risks" in combined_text
    assert "Research reference:" in combined_text


def test_report_workflow_supports_nested_subsection_depth(tmp_path: Path):
    ws = _new_workspace(tmp_path)

    first = ws.respond("s1", message="Create a report for supplier resilience")
    assert first["status"] == "ok"
    second = ws.respond("s1", message="use suggested")
    assert second["status"] == "ok"

    deeper = ws.respond("s1", message="add deeper")
    assert deeper["status"] == "ok"
    assert "subsection" in deeper["assistant_message"]["content_text"].lower()

    workflow = ws.sessions.get_session("s1")["metadata"]["report_workflow"]
    assert workflow["state"] == "awaiting_subsections"
    assert int(workflow.get("structure_max_depth") or 1) >= 2
    hierarchy = workflow.get("section_hierarchy") or {}
    assert "" in hierarchy
    has_depth3 = any(key.count("::") >= 2 for key in hierarchy.keys() if key)
    assert has_depth3 or int(workflow.get("structure_max_depth") or 1) == 2

    finalize = ws.respond("s1", message="final structure")
    assert finalize["status"] == "ok"
    assert "instructions" in finalize["assistant_message"]["content_text"].lower()

    done = ws.respond("s1", message="use defaults")
    assert done["status"] == "ok"
    assert "Report generation completed." in done["assistant_message"]["content_text"]


def test_report_gap_fill_target_selection_prioritizes_low_quality_sections(tmp_path: Path):
    ws = _new_workspace(tmp_path)
    section_keys = [
        "Executive Summary::",
        "Risk Posture::",
        "Mitigation Actions::",
        "Appendix::",
    ]
    assignments = {
        "Executive Summary::": {"quality_score": 3, "reference_text": "strong coverage"},
        "Risk Posture::": {"quality_score": 1, "reference_text": "partial coverage"},
        "Mitigation Actions::": {"quality_score": 0, "reference_text": ""},
        "Appendix::": {"quality_score": 2, "reference_text": "adequate coverage"},
    }

    targets = ws._select_report_gap_fill_section_targets(
        section_keys=section_keys,
        assignments=assignments,
        min_quality_score=2,
        max_targets=3,
    )

    assert targets == ["Mitigation Actions::", "Risk Posture::"]


def test_report_gap_fill_can_use_llm_planner_calls(tmp_path: Path):
    ws = _new_workspace(tmp_path)
    ws.agent = _ReportGapPlannerAgent()

    ws.respond("s1", message="Create a report for supplier risk")
    ws.respond("s1", message="use suggested")
    ws.respond("s1", message="use suggested")
    final = ws.respond("s1", message="use defaults")

    assert final["status"] == "ok"
    generate_responses = [
        event["payload"]
        for event in ws.events.events
        if event["event_type"] == "tool_call_response"
        and event["payload"].get("tool_name") == "generate_report_document"
    ]
    assert generate_responses
    result = generate_responses[-1]["result"]
    assert result["status"] == "completed"
    assert result.get("gap_fill_mode") in ("llm_iterative", "llm_with_heuristic_fallback")
    assert int(result["gap_fill_call_count"]) >= 1

    db_requests = [
        event["payload"]
        for event in ws.events.events
        if event["event_type"] == "tool_call_request"
        and event["payload"].get("tool_name") == "database_research"
    ]
    assert db_requests
    has_planned_query = any(
        "supplier concentration risks" in str(item["arguments"].get("research_statement", "")).lower()
        for item in db_requests
    )
    assert has_planned_query


def test_preview_artifact_renders_pdf_and_xlsx_html(tmp_path: Path):
    ws = _new_workspace(tmp_path)

    pdf_path = tmp_path / "preview.pdf"
    ws._write_minimal_pdf(
        export_path=pdf_path,
        title="Preview PDF",
        lines=["line one", "line two"],
    )
    pdf_artifact = ws.artifacts.create_artifact(
        session_id="s1",
        artifact_type="export_file",
        lifecycle_state="final",
        format="pdf",
        filename=pdf_path.name,
        storage_uri=str(pdf_path),
        mime_type="application/pdf",
        size_bytes=pdf_path.stat().st_size,
        metadata={},
    )
    pdf_preview = ws.preview_artifact("s1", pdf_artifact["artifact_id"])
    assert pdf_preview["status"] == "ok"
    assert pdf_preview["preview_format"] == "html"
    assert "PDF Preview" in (pdf_preview["preview_content"] or "")
    assert "data:application/pdf;base64," in (pdf_preview["preview_content"] or "")

    xlsx_path = tmp_path / "preview.xlsx"
    ws._write_minimal_xlsx(
        export_path=xlsx_path,
        sheet_name="Report Content",
        rows=[["header_a", "header_b"], ["alpha", "beta"]],
        metadata_rows=[["key", "value"], ["topic", "Preview XLSX"]],
    )
    xlsx_artifact = ws.artifacts.create_artifact(
        session_id="s1",
        artifact_type="export_file",
        lifecycle_state="final",
        format="xlsx",
        filename=xlsx_path.name,
        storage_uri=str(xlsx_path),
        mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        size_bytes=xlsx_path.stat().st_size,
        metadata={},
    )
    xlsx_preview = ws.preview_artifact("s1", xlsx_artifact["artifact_id"])
    assert xlsx_preview["status"] == "ok"
    assert xlsx_preview["preview_format"] == "html"
    html = xlsx_preview["preview_content"] or ""
    assert "Spreadsheet Preview" in html
    assert "<table" in html
    assert "Report Content" in html


def test_pdf_preview_large_file_mode_disables_inline_embed(tmp_path: Path):
    ws = _new_workspace(tmp_path)
    large_pdf = tmp_path / "large-preview.pdf"
    large_pdf.write_bytes(b"%PDF-1.4\n" + b"x" * (4 * 1024 * 1024 + 32))

    html = ws._render_pdf_preview_html(large_pdf)
    assert "Large File Mode" in html
    assert "Inline rendering is disabled" in html
    assert "data:application/pdf;base64," not in html


def test_xlsx_preview_uses_large_file_mode_limits(tmp_path: Path):
    ws = _new_workspace(tmp_path)
    large_xlsx = tmp_path / "large-preview.xlsx"
    large_xlsx.write_bytes(b"PK" + b"x" * (2 * 1024 * 1024 + 20))

    captured: dict[str, int] = {}

    def fake_extract(*, path: Path, max_sheets: int, max_rows: int, max_cols: int):
        captured["max_sheets"] = max_sheets
        captured["max_rows"] = max_rows
        captured["max_cols"] = max_cols
        return [{"name": "Sheet1", "rows": [["A", "B"], ["1", "2"]]}]

    ws._extract_xlsx_preview_data = fake_extract  # type: ignore[method-assign]
    html = ws._render_xlsx_preview_html(large_xlsx)
    assert "Preview mode: balanced mode" in html
    assert captured == {"max_sheets": 2, "max_rows": 22, "max_cols": 10}
