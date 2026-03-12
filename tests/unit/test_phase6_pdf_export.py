from __future__ import annotations

from pathlib import Path

from docx import Document

from word_engine.errors import ErrorCode
from word_engine.service import ServiceError


def test_convert_to_pdf_missing_source_returns_file_not_found(make_service, tmp_path):
    service = make_service(tmp_path)
    missing = tmp_path / "missing.docx"
    output = tmp_path / "missing.pdf"

    result = service.convert_to_pdf(str(missing), str(output))
    assert result["status"] == "error"
    assert result["error_code"] == "FILE_NOT_FOUND"


def test_convert_to_pdf_invalid_output_extension(make_service, tmp_path):
    service = make_service(tmp_path)
    source = tmp_path / "source.docx"
    Document().save(str(source))

    result = service.convert_to_pdf(str(source), str(tmp_path / "out.txt"))
    assert result["status"] == "error"
    assert result["error_code"] == "INVALID_PATH"


def test_convert_to_pdf_success_with_mock_backend(make_service, tmp_path, monkeypatch):
    service = make_service(tmp_path)
    source = tmp_path / "source.docx"
    output = tmp_path / "result.pdf"
    Document().save(str(source))

    def fake_convert(src: Path, dst: Path) -> str:
        dst.write_bytes(b"%PDF-1.4\n%mock\n")
        return "mock-backend"

    monkeypatch.setattr(service, "_perform_pdf_conversion", fake_convert)

    result = service.convert_to_pdf(str(source), str(output))
    assert result["status"] == "ok"
    assert result["output_path"] == str(output.resolve())
    assert result["method"] == "mock-backend"
    assert result["experimental"] is True
    assert output.exists()


def test_convert_to_pdf_default_output_path(make_service, tmp_path, monkeypatch):
    service = make_service(tmp_path)
    source = tmp_path / "source-default.docx"
    Document().save(str(source))

    def fake_convert(src: Path, dst: Path) -> str:
        dst.write_bytes(b"%PDF-1.4\n%mock\n")
        return "mock-backend"

    monkeypatch.setattr(service, "_perform_pdf_conversion", fake_convert)

    result = service.convert_to_pdf(str(source))
    expected = source.with_suffix(".pdf").resolve()
    assert result["status"] == "ok"
    assert result["output_path"] == str(expected)
    assert expected.exists()


def test_convert_to_pdf_backend_failure(make_service, tmp_path, monkeypatch):
    service = make_service(tmp_path)
    source = tmp_path / "source-fail.docx"
    output = tmp_path / "result-fail.pdf"
    Document().save(str(source))

    def failing_convert(src: Path, dst: Path) -> str:
        raise ServiceError(ErrorCode.DOCX_ERROR, "no converter backend available")

    monkeypatch.setattr(service, "_perform_pdf_conversion", failing_convert)

    result = service.convert_to_pdf(str(source), str(output))
    assert result["status"] == "error"
    assert result["error_code"] == "DOCX_ERROR"
