from __future__ import annotations

from docx import Document

from word_ui.workspace_v2 import WordUIWorkspaceV2


def test_summarize_ingestion_text_truncates():
    text = "alpha " * 500
    summary = WordUIWorkspaceV2._summarize_ingestion_text(text, max_chars=120)
    assert len(summary) <= 120
    assert summary.endswith("...")


def test_chunk_text_for_ingestion_produces_multiple_chunks():
    text = " ".join(f"token{i}" for i in range(1000))
    chunks = WordUIWorkspaceV2._chunk_text_for_ingestion(
        text,
        chunk_chars=200,
        overlap_chars=40,
        max_chunks=20,
    )
    assert len(chunks) > 1
    assert len(chunks) <= 20
    assert all(chunk.strip() for chunk in chunks)


def test_keyword_terms_deduplicates_and_filters_short_tokens():
    terms = WordUIWorkspaceV2._keyword_terms("AI ai report report q2 data")
    assert "report" in terms
    assert "data" in terms
    assert "ai" not in terms
    assert terms.count("report") == 1


def test_extract_text_for_ingestion_reads_txt(tmp_path):
    path = tmp_path / "sample.txt"
    path.write_text("hello\nworld", encoding="utf-8")
    text = WordUIWorkspaceV2._extract_text_for_ingestion(path)
    assert "hello" in text
    assert "world" in text


def test_extract_text_for_ingestion_reads_docx(tmp_path):
    path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("Heading")
    doc.add_paragraph("Body line")
    doc.save(str(path))
    text = WordUIWorkspaceV2._extract_text_for_ingestion(path)
    assert "Heading" in text
    assert "Body line" in text


def test_resolve_relation_name_uses_source_id_fallback():
    schema, relation = WordUIWorkspaceV2._resolve_relation_name({}, "risk_db.suppliers")
    assert schema == "risk_db"
    assert relation == "suppliers"


def test_execute_search_index_probe_returns_metadata_result():
    ws = object.__new__(WordUIWorkspaceV2)
    result = ws._execute_search_index_probe(
        source={
            "source_id": "search.customer_notes",
            "name": "Customer Notes",
            "source_type": "search_index",
            "location": {"index": "customer-notes"},
            "schema_json": {"fields": [{"name": "note_text"}]},
        },
        query_terms=["customer", "risk"],
    )
    assert result["status"] == "completed"
    assert result["handler"] == "search_index_probe"
    assert result["mode"] == "metadata_only"
    assert result["index_name"] == "customer-notes"
    assert "customer" in result["matched_terms"]
