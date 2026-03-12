from __future__ import annotations

import importlib.util
import logging
from concurrent.futures import ProcessPoolExecutor

from docx import Document
import pytest

from tests.unit.helpers import build_sample_document, build_styled_section_document, paragraph_texts
from word_engine.config import EngineConfig
from word_engine.service import WordDocumentService


def _multiprocess_insert_worker(args: tuple[str, str, int]) -> str:
    root_path, target_path, worker_id = args
    service = WordDocumentService(config=EngineConfig(allowed_roots=[root_path]))
    result = service.insert_paragraphs(
        target_path,
        after_paragraph_index=1,
        paragraphs=[{"text": f"mp-marker-{worker_id}"}],
    )
    return result["status"]


def test_search_and_replace_idempotent_without_limit(make_service, tmp_path):
    service = make_service(tmp_path)
    target = tmp_path / "idempotent-unbounded.docx"
    build_sample_document(target)

    first = service.search_and_replace(
        str(target),
        find_text="Instruction",
        replace_text="Instruction",
    )
    second = service.search_and_replace(
        str(target),
        find_text="Instruction",
        replace_text="Instruction",
    )

    assert first["status"] == "ok"
    assert second["status"] == "ok"
    assert first["replacements"] == second["replacements"] == 3

    lines = paragraph_texts(target)
    assert lines.count("Instruction A1") == 1
    assert lines.count("Instruction A2") == 1
    assert lines.count("Instruction B1") == 1


def test_search_and_replace_idempotent_with_limit(make_service, tmp_path):
    service = make_service(tmp_path)
    target = tmp_path / "idempotent-limited.docx"
    build_sample_document(target)

    first = service.search_and_replace(
        str(target),
        find_text="Instruction",
        replace_text="Instruction",
        max_replacements=2,
    )
    second = service.search_and_replace(
        str(target),
        find_text="Instruction",
        replace_text="Instruction",
        max_replacements=2,
    )

    assert first["status"] == "ok"
    assert second["status"] == "ok"
    assert first["replacements"] == second["replacements"] == 2


def test_replace_section_content_repeated_is_stable(make_service, tmp_path):
    service = make_service(tmp_path)
    target = tmp_path / "repeat-section.docx"
    build_styled_section_document(target)

    payload = ["Generated line 1", "Generated line 2"]
    first = service.replace_section_content(
        str(target),
        selector={"mode": "heading_exact", "value": "Executive Summary", "occurrence": 1},
        new_paragraphs=payload,
        preserve_style=True,
    )
    snapshot_after_first = paragraph_texts(target)

    second = service.replace_section_content(
        str(target),
        selector={"mode": "heading_exact", "value": "Executive Summary", "occurrence": 1},
        new_paragraphs=payload,
        preserve_style=True,
    )
    snapshot_after_second = paragraph_texts(target)

    assert first["status"] == "ok"
    assert second["status"] == "ok"
    assert snapshot_after_first == snapshot_after_second


def test_get_document_outline_is_deterministic(make_service, tmp_path):
    service = make_service(tmp_path)
    target = tmp_path / "det-outline.docx"
    build_sample_document(target)

    first = service.get_document_outline(str(target))
    second = service.get_document_outline(str(target))

    assert first["status"] == "ok"
    assert second["status"] == "ok"
    assert first["headings"] == second["headings"]
    assert [item["paragraph_index"] for item in first["headings"]] == sorted(
        item["paragraph_index"] for item in first["headings"]
    )


def test_find_text_is_deterministic_and_ordered(make_service, tmp_path):
    service = make_service(tmp_path)
    target = tmp_path / "det-find.docx"
    build_sample_document(target)

    first = service.find_text(str(target), "Instruction", match_case=True)
    second = service.find_text(str(target), "Instruction", match_case=True)

    assert first["status"] == "ok"
    assert second["status"] == "ok"
    assert first["matches"] == second["matches"]

    pairs = [(match["paragraph_index"], match["start"]) for match in first["matches"]]
    assert pairs == sorted(pairs)
    for match in first["matches"]:
        assert match["end"] >= match["start"]


def test_path_allowlist_rejection(make_service, tmp_path):
    service = make_service(tmp_path)
    outside = tmp_path.parent / "outside.docx"
    result = service.create_document(str(outside))
    assert result["status"] == "error"
    assert result["error_code"] == "PATH_NOT_ALLOWED"


def test_file_size_guard(make_service, tmp_path):
    tiny_service = WordDocumentService(config=EngineConfig(allowed_roots=[tmp_path], max_file_size_bytes=1))
    target = tmp_path / "oversize.docx"
    Document().save(str(target))

    result = tiny_service.get_document_info(str(target))
    assert result["status"] == "error"
    assert result["error_code"] == "FILE_TOO_LARGE"


def test_file_size_guard_large_fixture_backed(make_service, tmp_path):
    target = tmp_path / "oversize-large-fixture.docx"
    doc = Document()
    doc.add_heading("Large Fixture", level=1)
    for i in range(1200):
        doc.add_paragraph(f"Paragraph {i} " + ("data-block " * 20))
    doc.save(str(target))

    file_size = target.stat().st_size
    strict_limit = max(1024, file_size // 2)
    strict_service = WordDocumentService(
        config=EngineConfig(allowed_roots=[tmp_path], max_file_size_bytes=strict_limit)
    )

    result = strict_service.get_document_info(str(target))
    assert result["status"] == "error"
    assert result["error_code"] == "FILE_TOO_LARGE"


def test_selector_not_found_missing_start_anchor(make_service, tmp_path):
    service = make_service(tmp_path)
    target = tmp_path / "missing-start-anchor.docx"
    doc = Document()
    doc.add_paragraph("A")
    doc.add_paragraph("B")
    doc.save(str(target))

    result = service.replace_section_content(
        str(target),
        selector={"mode": "anchors", "start_text": "BEGIN", "end_text": "END"},
        new_paragraphs=["x"],
    )
    assert result["status"] == "error"
    assert result["error_code"] == "SELECTOR_NOT_FOUND"


def test_selector_not_found_missing_end_anchor(make_service, tmp_path):
    service = make_service(tmp_path)
    target = tmp_path / "missing-end-anchor.docx"
    doc = Document()
    doc.add_paragraph("BEGIN")
    doc.add_paragraph("Body")
    doc.save(str(target))

    result = service.replace_section_content(
        str(target),
        selector={"mode": "anchors", "start_text": "BEGIN", "end_text": "END"},
        new_paragraphs=["x"],
    )
    assert result["status"] == "error"
    assert result["error_code"] == "SELECTOR_NOT_FOUND"


def test_delete_paragraph_range_bound_error(make_service, tmp_path):
    service = make_service(tmp_path)
    target = tmp_path / "range-bound.docx"
    build_sample_document(target)

    result = service.delete_paragraph_range(str(target), 0, 100)
    assert result["status"] == "error"
    assert result["error_code"] == "PARAGRAPH_INDEX_OUT_OF_RANGE"


def test_concurrent_writers_do_not_corrupt_doc(make_service, tmp_path):
    from concurrent.futures import ThreadPoolExecutor

    service = make_service(tmp_path)
    target = tmp_path / "no-corruption.docx"
    build_sample_document(target)

    def worker(i: int):
        return service.insert_paragraphs(
            str(target),
            after_paragraph_index=1,
            paragraphs=[{"text": f"marker-{i}"}],
        )

    with ThreadPoolExecutor(max_workers=8) as executor:
        results = list(executor.map(worker, range(12)))

    assert all(result["status"] == "ok" for result in results)
    lines = paragraph_texts(target)
    for i in range(12):
        assert f"marker-{i}" in lines

    # A readable document here confirms no structural write corruption.
    reopened = Document(str(target))
    assert len(reopened.paragraphs) >= 12


def test_multiprocessing_writers_do_not_corrupt_doc(make_service, tmp_path):
    target = tmp_path / "no-corruption-multiprocess.docx"
    build_sample_document(target)

    worker_args = [(str(tmp_path), str(target), i) for i in range(8)]
    try:
        with ProcessPoolExecutor(max_workers=4) as executor:
            statuses = list(executor.map(_multiprocess_insert_worker, worker_args))
    except PermissionError as exc:
        pytest.skip(f"multiprocessing process pool is restricted in this environment: {exc}")

    assert statuses
    assert all(status == "ok" for status in statuses)

    lines = paragraph_texts(target)
    for i in range(8):
        assert f"mp-marker-{i}" in lines

    reopened = Document(str(target))
    assert len(reopened.paragraphs) >= 8


def test_structured_logging_fields_success_and_failure(make_service, tmp_path, caplog):
    service = make_service(tmp_path)
    target = tmp_path / "log-fields.docx"
    build_sample_document(target)

    logger_name = "word_engine.service"
    caplog.set_level(logging.INFO, logger=logger_name)

    ok_response = service.get_document_info(str(target))
    assert ok_response["status"] == "ok"

    err_response = service.get_paragraph_text(str(target), 999)
    assert err_response["status"] == "error"

    success_records = [r for r in caplog.records if r.msg == "word_engine_success"]
    failure_records = [r for r in caplog.records if r.msg == "word_engine_failure"]

    assert success_records
    assert failure_records

    success = success_records[-1]
    assert getattr(success, "event", None)
    assert getattr(success, "file_path", None)
    assert getattr(success, "status", None) == "ok"
    assert isinstance(getattr(success, "duration_ms", None), int)

    failure = failure_records[-1]
    assert getattr(failure, "event", None)
    assert getattr(failure, "file_path", None)
    assert getattr(failure, "status", None) == "error"
    assert isinstance(getattr(failure, "duration_ms", None), int)
    assert getattr(failure, "error_code", None) == "PARAGRAPH_INDEX_OUT_OF_RANGE"


def test_mcp_runtime_smoke_or_install_guidance():
    from word_mcp_server.server import _create_mcp

    fastmcp_available = importlib.util.find_spec("fastmcp") is not None
    if fastmcp_available:
        mcp = _create_mcp()
        assert mcp is not None
        assert hasattr(mcp, "run")
    else:
        try:
            _create_mcp()
        except RuntimeError as exc:
            message = str(exc)
            assert "fastmcp is not installed" in message
            assert "pip install -e '.[mcp]'" in message
        else:  # pragma: no cover
            raise AssertionError("Expected RuntimeError guidance when fastmcp is unavailable")
