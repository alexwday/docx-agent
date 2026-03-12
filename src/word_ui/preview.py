"""DOCX preview rendering helpers for the UI workspace."""

from __future__ import annotations

import html
from pathlib import Path
import re
from typing import Any

from docx import Document

HEADING_STYLE_PATTERN = re.compile(r"^heading\s+(\d+)$", re.IGNORECASE)


class PreviewRenderError(Exception):
    """Raised when preview rendering cannot be completed."""


class DocxPreviewRenderer:
    """Renders a read-only HTML preview artifact from a DOCX file."""

    def __init__(self, preview_dir_name: str = ".docx-agent-preview") -> None:
        self.preview_dir_name = preview_dir_name

    def render_docx_to_html(
        self,
        file_path: str,
        revision_id: str,
        output_dir: str | None = None,
    ) -> dict[str, Any]:
        source = Path(file_path).expanduser().resolve()
        if source.suffix.lower() != ".docx":
            raise PreviewRenderError("only .docx files are supported for preview rendering")
        if not source.exists():
            raise PreviewRenderError(f"source document not found: {source}")
        if not revision_id:
            raise PreviewRenderError("revision_id must be a non-empty string")

        document = Document(str(source))
        if output_dir is None:
            artifact_dir = source.parent / self.preview_dir_name
        else:
            artifact_dir = Path(output_dir).expanduser().resolve()
        artifact_dir.mkdir(parents=True, exist_ok=True)

        safe_revision = re.sub(r"[^A-Za-z0-9._-]", "-", revision_id)
        artifact_path = artifact_dir / f"{source.stem}.{safe_revision}.preview.html"
        html_content = self._render_html(document=document, source_name=source.name)
        artifact_path.write_text(html_content, encoding="utf-8")

        return {
            "artifact_path": str(artifact_path),
            "artifact_format": "html",
            "source_path": str(source),
            "paragraph_count": len(document.paragraphs),
            "table_count": len(document.tables),
        }

    def _render_html(self, document: Document, source_name: str) -> str:
        body_parts: list[str] = []
        for paragraph in document.paragraphs:
            style_name = paragraph.style.name if paragraph.style else ""
            level = self._extract_heading_level(style_name)
            escaped_text = html.escape(paragraph.text)
            if not escaped_text.strip():
                continue
            if level is not None:
                body_parts.append(f"<h{level}>{escaped_text}</h{level}>")
            else:
                body_parts.append(f"<p>{escaped_text}</p>")

        if document.tables:
            body_parts.append("<section><h2>Tables</h2>")
            for index, table in enumerate(document.tables, start=1):
                body_parts.append(f"<h3>Table {index}</h3>")
                body_parts.append("<table><tbody>")
                for row in table.rows:
                    body_parts.append("<tr>")
                    for cell in row.cells:
                        body_parts.append(f"<td>{html.escape(cell.text)}</td>")
                    body_parts.append("</tr>")
                body_parts.append("</tbody></table>")
            body_parts.append("</section>")

        if not body_parts:
            body_parts.append("<p><em>No textual preview content is available.</em></p>")

        return "\n".join(
            [
                "<!doctype html>",
                "<html>",
                "<head>",
                '<meta charset="utf-8" />',
                f"<title>Preview: {html.escape(source_name)}</title>",
                "<style>",
                "body { font-family: Georgia, 'Times New Roman', serif; margin: 24px auto; max-width: 900px; line-height: 1.5; color: #1e1e1e; }",
                "h1, h2, h3, h4, h5, h6 { font-family: 'Avenir Next', 'Segoe UI', sans-serif; margin-top: 1.1em; }",
                "p { margin: 0.4em 0; }",
                "table { border-collapse: collapse; margin: 0.8em 0; width: 100%; }",
                "td { border: 1px solid #d0d0d0; padding: 6px 8px; vertical-align: top; }",
                "em { color: #555; }",
                "</style>",
                "</head>",
                "<body>",
                f"<header><h1>{html.escape(source_name)}</h1></header>",
                *body_parts,
                "</body>",
                "</html>",
            ]
        )

    def _extract_heading_level(self, style_name: str) -> int | None:
        match = HEADING_STYLE_PATTERN.match(style_name.strip())
        if match is None:
            return None
        level = int(match.group(1))
        return min(max(level, 1), 6)
