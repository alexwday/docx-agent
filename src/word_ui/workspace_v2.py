"""Postgres-backed V2 workspace service for session-centric orchestration."""

from __future__ import annotations

from datetime import UTC, datetime
import json
import logging
import mimetypes
from pathlib import Path
import shutil
from typing import Any, Callable
from uuid import uuid4

from word_agent import WordAgent

from .preview import DocxPreviewRenderer, PreviewRenderError
from .retrievers import (
    FunctionSourceRetriever,
    GenericMetadataRetriever,
    MockComplianceRetriever,
    MockEmployeeRetriever,
    MockFinancialDataRetriever,
    MockSalesOrdersRetriever,
    MockSupplierRiskRetriever,
    PostgresRelationProbeRetriever,
    RetrieverRegistry,
    SearchIndexMetadataRetriever,
)
from data_sources.config import DataSourcesConfig
from data_sources.db import DataSourcesDB
from data_sources.retrieve.supp_financials import SuppFinancialsRetriever


logger = logging.getLogger(__name__)


class WordUIWorkspaceV2:
    """Postgres-backed orchestration workspace."""

    def __init__(
        self,
        *,
        dsn: str | None = None,
        store: Any | None = None,
        agent: WordAgent | None = None,
        preview_renderer: DocxPreviewRenderer | None = None,
        allowed_roots: list[str | Path] | None = None,
        contract_version: str = "v2",
        openai_api_key: str | None = None,
        openai_model: str = "gpt-4.1",
        mock_retriever_registry: dict[str, Callable[..., dict[str, Any]]] | None = None,
        retriever_registry: RetrieverRegistry | None = None,
    ) -> None:
        try:
            from word_store import (
                ArtifactKnowledgeRepository,
                DataSourcesRepository,
                MessageEventsRepository,
                MessagesRepository,
                PostgresStore,
                SessionArtifactsRepository,
                SessionsRepository,
            )
        except ModuleNotFoundError as exc:
            raise RuntimeError(
                "word_store dependencies are unavailable. Install psycopg[binary] to use WordUIWorkspaceV2."
            ) from exc

        self.store = store or PostgresStore(dsn=dsn)
        self.sessions = SessionsRepository(self.store)
        self.messages = MessagesRepository(self.store)
        self.events = MessageEventsRepository(self.store)
        self.artifacts = SessionArtifactsRepository(self.store)
        self.knowledge = ArtifactKnowledgeRepository(self.store)
        self.data_sources = DataSourcesRepository(self.store)
        self.agent = agent or WordAgent(model=openai_model, api_key=openai_api_key)
        self.preview_renderer = preview_renderer or DocxPreviewRenderer()
        self.allowed_roots = self._resolve_allowed_roots(allowed_roots)
        self.contract_version = contract_version
        self.retriever_registry = retriever_registry or self._build_retriever_registry(
            custom_source_retrievers=mock_retriever_registry
        )

    # Auth/session ----------------------------------------------------------

    def login(self, employee_id: str) -> dict[str, Any]:
        if not self._is_valid_user_id(employee_id):
            return self._error("INVALID_ARGUMENT", "employee_id must be a 9-digit string")
        return self._ok(user={"user_id": employee_id})

    def list_user_sessions(
        self,
        user_id: str,
        *,
        status: str | None = None,
        limit: int = 50,
    ) -> dict[str, Any]:
        if not self._is_valid_user_id(user_id):
            return self._error("INVALID_ARGUMENT", "user_id must be a 9-digit string")
        rows = self.sessions.list_sessions(user_id, status=status, limit=limit)
        return self._ok(sessions=rows, next_cursor=None)

    def create_session(
        self,
        user_id: str,
        *,
        title: str | None = None,
        metadata: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        if not self._is_valid_user_id(user_id):
            return self._error("INVALID_ARGUMENT", "user_id must be a 9-digit string")
        session_meta = dict(metadata or {})
        if title:
            session_meta["title_source"] = "manual"
        row = self.sessions.create_session(user_id, title=title, metadata=session_meta)
        return self._ok(session=row)

    def delete_session(self, session_id: str) -> dict[str, Any]:
        session = self.sessions.get_session(session_id)
        if session is None:
            return self._error("NOT_FOUND", f"unknown session_id: {session_id}")
        self.sessions.update_session(session_id, status="deleted")
        return self._ok()

    def rename_session(self, session_id: str, title: str) -> dict[str, Any]:
        session = self.sessions.get_session(session_id)
        if session is None:
            return self._error("NOT_FOUND", f"unknown session_id: {session_id}")
        title = title.strip()
        if not title:
            return self._error("INVALID_ARGUMENT", "title must be a non-empty string")
        existing_meta = dict(session.get("metadata") or {})
        existing_meta["title_source"] = "manual"
        self.sessions.update_session(session_id, title=title, metadata=existing_meta)
        return self._ok()

    def hydrate_session(self, session_id: str) -> dict[str, Any]:
        session = self.sessions.get_session(session_id)
        if session is None:
            return self._error("NOT_FOUND", f"unknown session_id: {session_id}")
        messages = self.messages.list_messages(session_id, limit=2000)
        panes = self.artifacts.list_artifact_panes(session_id)
        active_preview = panes["report_documents"][0]["artifact_id"] if panes["report_documents"] else None
        return self._ok(
            session=session,
            messages=messages,
            artifact_panes=panes,
            active_preview_artifact_id=active_preview,
        )

    # Core turn processing --------------------------------------------------

    def respond(
        self,
        session_id: str,
        *,
        message: str | None = None,
        data_source_filters: list[str] | None = None,
        response_mode: str = "auto",
        report_plan_state: dict[str, Any] | None = None,
        report_plan_action: str | None = None,
    ) -> dict[str, Any]:
        session = self.sessions.get_session(session_id)
        if session is None:
            return self._error("NOT_FOUND", f"unknown session_id: {session_id}")
        if (
            not message
            and not data_source_filters
            and not isinstance(report_plan_state, dict)
            and not report_plan_action
        ):
            return self._error("INVALID_ARGUMENT", "at least one actionable input is required")

        user_row: dict[str, Any] | None = None
        if message and message.strip():
            user_row = self.messages.create_message(
                session_id=session_id,
                role="user",
                content_text=message.strip(),
                content_json={"text": message.strip()},
                processing_state="completed",
            )
            self._maybe_auto_rename_session(session_id)

        started_at = datetime.now(UTC)
        assistant_row = self.messages.create_message(
            session_id=session_id,
            role="assistant",
            content_text="",
            content_json={},
            parent_message_id=user_row["message_id"] if user_row else None,
            processing_state="pending",
            processing_started_at=started_at,
        )

        filter_result = self._resolve_data_source_filters(data_source_filters)
        if filter_result["error"]:
            self.messages.update_message_content_and_state(
                session_id=session_id,
                message_id=assistant_row["message_id"],
                content_text="",
                content_json={},
                processing_state="failed",
                processing_ended_at=datetime.now(UTC),
                error={
                    "error_code": filter_result["error_code"],
                    "message": filter_result["error_message"],
                },
            )
            return self._error(filter_result["error_code"], filter_result["error_message"])

        self.events.create_event(
            session_id=session_id,
            message_id=assistant_row["message_id"],
            event_type="tool_definitions_injected",
            payload={
                "tools": [
                    "research_internal_data_sources",
                    "database_research",
                    "research_uploaded_documents",
                    "generate_report_document",
                    "export_report_document",
                ]
            },
        )

        effective_sources = filter_result["effective_sources"]
        self.events.create_event(
            session_id=session_id,
            message_id=assistant_row["message_id"],
            event_type="ui_data_source_filter_applied",
            payload={
                "mode": filter_result["mode"],
                "requested_source_ids": filter_result["requested_source_ids"],
                "effective_source_ids": filter_result["effective_source_ids"],
                "ignored_source_ids": filter_result["ignored_source_ids"],
            },
        )

        self.events.create_event(
            session_id=session_id,
            message_id=assistant_row["message_id"],
            event_type="available_data_sources_injected",
            payload={
                "count": len(effective_sources),
                "sources": [
                    {
                        "source_id": item["source_id"],
                        "name": item["name"],
                        "source_type": item["source_type"],
                    }
                    for item in effective_sources
                ],
            },
        )

        latest_messages = self._build_llm_messages(session_id)
        self.events.create_event(
            session_id=session_id,
            message_id=assistant_row["message_id"],
            event_type="conversation_context_injected",
            payload={"message_count": len(latest_messages)},
        )

        uploaded_context_units = self._select_uploaded_context(
            session_id=session_id,
            query_text=message or "",
            limit=6,
        )
        self.events.create_event(
            session_id=session_id,
            message_id=assistant_row["message_id"],
            event_type="uploaded_documents_context_injected",
            payload={
                "knowledge_unit_count": len(uploaded_context_units),
                "artifact_ids": sorted({item["artifact_id"] for item in uploaded_context_units}),
            },
        )
        prior_agent_interactions = self._select_prior_agent_interactions(
            session_id=session_id,
            assistant_message_id=assistant_row["message_id"],
            limit=8,
        )
        self.events.create_event(
            session_id=session_id,
            message_id=assistant_row["message_id"],
            event_type="prior_agent_interactions_injected",
            payload={
                "interaction_count": len(prior_agent_interactions),
                "event_types": sorted(
                    {item.get("event_type", "") for item in prior_agent_interactions if item.get("event_type")}
                ),
            },
        )

        selected_source_ids: list[str] = []
        planned_calls: list[dict[str, str]] = []

        internal_research = self._run_internal_data_source_research_tool(
            session_id=session_id,
            assistant_message_id=assistant_row["message_id"],
            query_text=message or "",
            effective_sources=effective_sources,
            selected_source_ids=selected_source_ids,
            planned_calls=planned_calls,
            selection_mode="model_iterative",
            selection_reasoning_summary="",
        )
        selected_source_ids = list(internal_research.get("selected_source_ids") or selected_source_ids)
        self.events.create_event(
            session_id=session_id,
            message_id=assistant_row["message_id"],
            event_type="agent_data_source_selected",
            payload={
                "selected_source_ids": selected_source_ids,
                "selection_reasoning_summary": internal_research.get(
                    "selection_reasoning_summary",
                    "Internal research planning completed.",
                ),
                "selection_inputs": {
                    "response_mode": response_mode,
                    "has_user_message": bool(user_row),
                },
                "selection_mode": internal_research.get("selection_mode", "model_iterative"),
                "planned_call_count": int(internal_research.get("planned_call_count") or 0),
                "executed_call_count": int(internal_research.get("executed_call_count") or 0),
            },
        )
        uploaded_research = self._run_uploaded_documents_research_tool(
            session_id=session_id,
            assistant_message_id=assistant_row["message_id"],
            query_text=message or "",
            uploaded_context_units=uploaded_context_units,
        )

        report_research_summary = self._summarize_research_for_prompt(
            internal_research, uploaded_research,
        )
        report_result = self._handle_report_workflow(
            session=session,
            session_id=session_id,
            message_text=message or "",
            assistant_message_id=assistant_row["message_id"],
            triggering_message_id=user_row["message_id"] if user_row else None,
            available_sources=effective_sources,
            preferred_source_ids=selected_source_ids,
            report_plan_state=report_plan_state,
            report_plan_action=report_plan_action,
            conversation_context=latest_messages,
            research_summary=report_research_summary or None,
        )
        pre_report_artifacts = self._merge_artifact_ids(
            internal_research.get("artifacts_created", []),
            uploaded_research.get("artifacts_created", []),
        )
        if report_result is not None:
            if report_result.get("error"):
                failed_row = self.messages.update_message_content_and_state(
                    session_id=session_id,
                    message_id=assistant_row["message_id"],
                    content_text="",
                    content_json={},
                    processing_state="failed",
                    processing_ended_at=datetime.now(UTC),
                    error={
                        "error_code": report_result["error_code"],
                        "message": report_result["error_message"],
                    },
                )
                self.events.create_event(
                    session_id=session_id,
                    message_id=assistant_row["message_id"],
                    event_type="error",
                    payload={"message": report_result["error_message"]},
                )
                return self._error(
                    report_result["error_code"],
                    report_result["error_message"],
                    user_message=user_row,
                    assistant_message=failed_row,
                    operation={"operation_id": assistant_row["message_id"], "state": "failed"},
                )

            completed_row = self.messages.update_message_content_and_state(
                session_id=session_id,
                message_id=assistant_row["message_id"],
                content_text=report_result["response_text"],
                content_json=report_result.get(
                    "content_json",
                    {"text": report_result["response_text"]},
                ),
                processing_state="completed",
                processing_ended_at=datetime.now(UTC),
                error=None,
            )
            return self._ok(
                user_message=user_row,
                assistant_message=completed_row,
                operation={
                    "operation_id": assistant_row["message_id"],
                    "state": completed_row["processing_state"],
                },
                filter_result={
                    "mode": filter_result["mode"],
                    "requested_source_ids": filter_result["requested_source_ids"],
                    "effective_source_ids": filter_result["effective_source_ids"],
                    "ignored_source_ids": filter_result["ignored_source_ids"],
                },
                selected_source_ids=selected_source_ids,
                artifacts_created=self._merge_artifact_ids(
                    pre_report_artifacts,
                    report_result.get("artifacts_created", []),
                ),
                artifacts_updated=report_result.get("artifacts_updated", []),
            )

        try:
            llm_messages = latest_messages
            system_context = self._build_system_context(
                filter_mode=filter_result["mode"],
                sources=effective_sources,
                selected_source_ids=selected_source_ids,
                uploaded_context_units=uploaded_context_units,
                prior_agent_interactions=prior_agent_interactions,
                internal_research=internal_research,
                uploaded_research=uploaded_research,
                report_plan_context_note=(
                    "The report plan was recently updated by the user. Review the updated plan before proceeding."
                    if isinstance(report_plan_state, dict)
                    else None
                ),
            )
            self.events.create_event(
                session_id=session_id,
                message_id=assistant_row["message_id"],
                event_type="orchestrator_system_prompt",
                payload={
                    "filter_mode": filter_result["mode"],
                    "available_source_ids": filter_result["effective_source_ids"],
                    "selected_source_ids": selected_source_ids,
                    "uploaded_knowledge_unit_count": len(uploaded_context_units),
                    "prior_agent_interaction_count": len(prior_agent_interactions),
                },
            )
            self.events.create_event(
                session_id=session_id,
                message_id=assistant_row["message_id"],
                event_type="model_request",
                payload={
                    "model": self.agent.model,
                    "message_count": len(llm_messages),
                },
            )
            response_text = self.agent.chat(llm_messages, system_context)
            completed_row = self.messages.update_message_content_and_state(
                session_id=session_id,
                message_id=assistant_row["message_id"],
                content_text=response_text,
                content_json={"text": response_text},
                processing_state="completed",
                processing_ended_at=datetime.now(UTC),
                error=None,
            )
            self.events.create_event(
                session_id=session_id,
                message_id=assistant_row["message_id"],
                event_type="model_response",
                payload={"text": response_text},
            )
        except Exception as exc:  # noqa: BLE001
            failed_row = self.messages.update_message_content_and_state(
                session_id=session_id,
                message_id=assistant_row["message_id"],
                content_text="",
                content_json={},
                processing_state="failed",
                processing_ended_at=datetime.now(UTC),
                error={"message": str(exc)},
            )
            self.events.create_event(
                session_id=session_id,
                message_id=assistant_row["message_id"],
                event_type="error",
                payload={"message": str(exc)},
            )
            return self._error(
                "PROCESSING_FAILED",
                "assistant processing failed",
                user_message=user_row,
                assistant_message=failed_row,
                operation={"operation_id": assistant_row["message_id"], "state": "failed"},
            )

        return self._ok(
            user_message=user_row,
            assistant_message=completed_row,
            operation={
                "operation_id": assistant_row["message_id"],
                "state": completed_row["processing_state"],
            },
            filter_result={
                "mode": filter_result["mode"],
                "requested_source_ids": filter_result["requested_source_ids"],
                "effective_source_ids": filter_result["effective_source_ids"],
                "ignored_source_ids": filter_result["ignored_source_ids"],
            },
            selected_source_ids=selected_source_ids,
            artifacts_created=pre_report_artifacts,
            artifacts_updated=[],
        )

    def get_operation(self, session_id: str, operation_id: str) -> dict[str, Any]:
        row = self.messages.get_message(session_id=session_id, message_id=operation_id)
        if row is None or row["role"] != "assistant":
            return self._error("NOT_FOUND", f"operation not found: {operation_id}")
        return self._ok(
            operation={
                "operation_id": row["message_id"],
                "state": row["processing_state"],
                "started_at": row["processing_started_at"],
                "ended_at": row["processing_ended_at"],
                "error": row["error"],
            }
        )

    # Artifacts -------------------------------------------------------------

    def upload_artifact(
        self,
        session_id: str,
        *,
        file_path: str,
        artifact_type: str = "upload",
    ) -> dict[str, Any]:
        session = self.sessions.get_session(session_id)
        if session is None:
            return self._error("NOT_FOUND", f"unknown session_id: {session_id}")
        normalized = self._normalize_path(file_path)
        if isinstance(normalized, dict):
            return normalized
        path = Path(normalized)
        if not path.exists() or not path.is_file():
            return self._error("INVALID_ARGUMENT", f"file does not exist: {path}")
        if not self._is_path_allowed(path):
            return self._error("PERMISSION_DENIED", f"path is outside allowed roots: {path}")

        mime_type, _ = mimetypes.guess_type(str(path))
        artifact = self.artifacts.create_artifact(
            session_id=session_id,
            artifact_type=artifact_type,
            lifecycle_state="final",
            format=path.suffix.lower().lstrip(".") or "bin",
            filename=path.name,
            storage_uri=str(path),
            mime_type=mime_type,
            size_bytes=path.stat().st_size,
            metadata={"ingestion_state": "queued"},
        )

        if artifact_type != "upload":
            return self._ok(artifact=artifact, ingestion_state="skipped")

        ingestion = self._ingest_artifact_content(session_id=session_id, artifact=artifact)
        updated = self.artifacts.get_artifact(session_id, artifact["artifact_id"]) or artifact
        return self._ok(
            artifact=updated,
            ingestion_state=ingestion["ingestion_state"],
            ingestion_summary=ingestion,
        )

    def list_session_artifacts(
        self,
        session_id: str,
        *,
        artifact_type: str | None = None,
        limit: int = 200,
    ) -> dict[str, Any]:
        session = self.sessions.get_session(session_id)
        if session is None:
            return self._error("NOT_FOUND", f"unknown session_id: {session_id}")
        rows = self.artifacts.list_artifacts(session_id, artifact_type=artifact_type, limit=limit)
        return self._ok(artifacts=rows, next_cursor=None)

    def preview_artifact(self, session_id: str, artifact_id: str) -> dict[str, Any]:
        row = self.artifacts.get_artifact(session_id, artifact_id)
        if row is None:
            return self._error("NOT_FOUND", f"artifact not found: {artifact_id}")
        storage_uri = row["storage_uri"]
        path = Path(storage_uri)
        if path.exists() and path.is_file():
            suffix = path.suffix.lower()
            if suffix in (".txt", ".md", ".json", ".html"):
                return self._ok(
                    artifact_id=artifact_id,
                    preview_format="html",
                    preview_content=self._text_to_html(path.read_text(encoding="utf-8", errors="replace")),
                    preview_url=None,
                )
            if suffix == ".docx":
                revision_id = str(uuid4())
                try:
                    artifact = self.preview_renderer.render_docx_to_html(str(path), revision_id)
                except PreviewRenderError as exc:
                    return self._error("PREVIEW_NOT_AVAILABLE", str(exc))
                html = Path(artifact["artifact_path"]).read_text(encoding="utf-8")
                return self._ok(
                    artifact_id=artifact_id,
                    preview_format="html",
                    preview_content=html,
                    preview_url=None,
                )
            if suffix == ".pdf":
                return self._ok(
                    artifact_id=artifact_id,
                    preview_format="html",
                    preview_content=self._render_pdf_preview_html(path),
                    preview_url=None,
                )
            if suffix == ".xlsx":
                return self._ok(
                    artifact_id=artifact_id,
                    preview_format="html",
                    preview_content=self._render_xlsx_preview_html(path),
                    preview_url=None,
                )
        return self._ok(
            artifact_id=artifact_id,
            preview_format=None,
            preview_content=None,
            preview_url=storage_uri,
        )

    # Events/data sources ---------------------------------------------------

    def message_events(self, session_id: str, message_id: str) -> dict[str, Any]:
        session = self.sessions.get_session(session_id)
        if session is None:
            return self._error("NOT_FOUND", f"unknown session_id: {session_id}")
        rows = self.events.list_events(session_id, message_id)
        return self._ok(events=rows)

    def list_data_source_catalog(
        self,
        *,
        enabled_only: bool = True,
        source_type: str | None = None,
    ) -> dict[str, Any]:
        rows = self.data_sources.list_sources(enabled_only=enabled_only, source_type=source_type)
        return self._ok(sources=rows)

    # Internal helpers ------------------------------------------------------

    def _handle_report_workflow(
        self,
        *,
        session: dict[str, Any],
        session_id: str,
        message_text: str,
        assistant_message_id: str,
        triggering_message_id: str | None,
        available_sources: list[dict[str, Any]],
        preferred_source_ids: list[str],
        report_plan_state: dict[str, Any] | None,
        report_plan_action: str | None,
        conversation_context: list[dict[str, str]] | None = None,
        research_summary: str | None = None,
    ) -> dict[str, Any] | None:
        text = (message_text or "").strip()
        action = str(report_plan_action or "").strip().lower() or None
        session_meta = dict(session.get("metadata") or {})
        workflow = session_meta.get("report_workflow")

        if workflow and not workflow.get("active", True):
            workflow = None

        if workflow is not None and isinstance(report_plan_state, dict):
            workflow = self._apply_report_plan_state_to_workflow(workflow, report_plan_state)
            self._mark_report_plan_updated(
                workflow,
                updated_by=str(report_plan_state.get("updated_by") or "user"),
                status=report_plan_state.get("status"),
            )
            self._persist_report_workflow(session_id, session_meta, workflow)
            self._emit_report_plan_state_updated_event(
                session_id=session_id,
                message_id=assistant_message_id,
                workflow=workflow,
                reason="user_plan_state_applied",
            )

        if workflow is None and action == "start_now":
            return {
                "error": True,
                "error_code": "INVALID_ARGUMENT",
                "error_message": "no active report workflow to start",
            }

        if workflow is None:
            interpreted = self._interpret_report_turn(
                workflow_state=None,
                workflow=None,
                user_text=text,
                conversation_context=conversation_context,
                research_summary=research_summary,
            )
            if interpreted is None:
                interpreted = self._interpret_report_turn_fallback(
                    workflow_state=None, user_text=text, workflow=None,
                )
            if interpreted is None or interpreted.get("intent") != "start_report":
                return None

            topic = interpreted.get("topic") or self._extract_report_topic(text)
            primary_sections = interpreted.get("suggested_sections") or self._suggest_primary_sections(topic)
            working_artifact = self._create_report_working_artifact(
                session_id=session_id,
                topic=topic,
                created_from_message_id=triggering_message_id,
            )
            if working_artifact is None:
                return {
                    "error": True,
                    "error_code": "PROCESSING_FAILED",
                    "error_message": "failed to create report working document",
                }

            workflow = {
                "active": True,
                "state": "awaiting_primary_sections",
                "topic": topic,
                "working_artifact_id": working_artifact["artifact_id"],
                "plan_id": str(uuid4()),
                "plan_title": topic[:180],
                "plan_summary": text[:600],
                "plan_status": "scaffolding",
                "plan_updated_by": "agent",
                "plan_updated_at": datetime.now(UTC).isoformat(),
                "primary_sections": primary_sections,
                "subsections": {},
                "section_hierarchy": {"": list(primary_sections)},
                "structure_max_depth": 1,
                "section_instructions": {},
                "section_instruction_sources": {},
                "section_ids": {},
                "created_at": datetime.now(UTC).isoformat(),
            }
            self._sync_legacy_structure_fields(workflow)
            self._ensure_report_section_ids(workflow)
            self._persist_report_workflow(session_id, session_meta, workflow)
            self._write_report_working_doc(
                session_id=session_id,
                workflow=workflow,
                stage="structure",
            )
            self.events.create_event(
                session_id=session_id,
                message_id=assistant_message_id,
                event_type="report_structure_proposed",
                payload={
                    "topic": topic,
                    "suggested_primary_sections": primary_sections,
                    "working_artifact_id": working_artifact["artifact_id"],
                },
            )
            self.events.create_event(
                session_id=session_id,
                message_id=assistant_message_id,
                event_type="report_plan_card_created",
                payload={
                    "plan_id": workflow["plan_id"],
                    "plan_status": workflow.get("plan_status", "scaffolding"),
                    "section_count": len(self._list_report_section_entries(workflow)),
                },
            )
            self._emit_report_plan_state_updated_event(
                session_id=session_id,
                message_id=assistant_message_id,
                workflow=workflow,
                reason="plan_initialized",
            )
            response_text = interpreted.get("response_text") or (
                f"I started a working report document for '{topic}'. "
                "Suggested primary sections are: "
                + ", ".join(primary_sections)
                + ".\n\nWould you like to use these sections, or would you prefer different ones?"
            )
            return self._report_workflow_response(
                response_text=response_text,
                workflow=workflow,
                artifacts_created=[working_artifact["artifact_id"]],
                artifacts_updated=[working_artifact["artifact_id"]],
            )

        state = workflow.get("state")
        topic = workflow.get("topic", "Report")
        working_artifact_id = workflow.get("working_artifact_id")
        if not working_artifact_id:
            return {
                "error": True,
                "error_code": "PROCESSING_FAILED",
                "error_message": "report workflow missing working_artifact_id",
            }

        if action == "start_now":
            self.events.create_event(
                session_id=session_id,
                message_id=assistant_message_id,
                event_type="report_plan_start_now_triggered",
                payload={
                    "plan_id": workflow.get("plan_id"),
                    "plan_status": workflow.get("plan_status"),
                    "workflow_state": state,
                },
            )
            instructions, instruction_meta = self._resolve_start_now_section_instructions(workflow)
            return self._generate_report_from_workflow(
                session_id=session_id,
                session_meta=session_meta,
                workflow=workflow,
                assistant_message_id=assistant_message_id,
                available_sources=available_sources,
                preferred_source_ids=preferred_source_ids,
                topic=topic,
                working_artifact_id=working_artifact_id,
                instructions=instructions,
                instruction_meta=instruction_meta,
            )

        if state == "awaiting_primary_sections":
            interpreted = self._interpret_report_turn(
                workflow_state="awaiting_primary_sections",
                workflow=workflow,
                user_text=text,
                conversation_context=conversation_context,
                research_summary=research_summary,
            )
            if interpreted is None:
                interpreted = self._interpret_report_turn_fallback(
                    workflow_state="awaiting_primary_sections", user_text=text, workflow=workflow,
                )
            intent = (interpreted or {}).get("intent", "accept_suggested")
            if intent == "provide_sections" and interpreted.get("sections"):
                workflow["primary_sections"] = interpreted["sections"]
            primary_sections = workflow.get("primary_sections") or self._suggest_primary_sections(topic)
            hierarchy = {"": list(primary_sections)}
            for primary in primary_sections:
                parent_key = self._section_key_from_path([primary])
                hierarchy[parent_key] = self._suggest_subsections(primary, topic)
                for subsection in hierarchy[parent_key]:
                    hierarchy.setdefault(self._section_key_from_path([primary, subsection]), [])
            workflow["section_hierarchy"] = hierarchy
            self._sync_legacy_structure_fields(workflow)
            self._ensure_report_section_ids(workflow)
            workflow["state"] = "awaiting_subsections"
            self._mark_report_plan_updated(workflow, updated_by="agent", status="scaffolding")
            self._persist_report_workflow(session_id, session_meta, workflow)
            self._write_report_working_doc(session_id=session_id, workflow=workflow, stage="structure")
            self.events.create_event(
                session_id=session_id,
                message_id=assistant_message_id,
                event_type="report_structure_confirmed",
                payload={
                    "level": "primary",
                    "primary_sections": primary_sections,
                    "subsections": dict(workflow.get("subsections") or {}),
                },
            )
            self._emit_report_plan_state_updated_event(
                session_id=session_id,
                message_id=assistant_message_id,
                workflow=workflow,
                reason="primary_sections_confirmed",
            )
            llm_response = (interpreted or {}).get("response_text", "")
            if not llm_response:
                lines = ["Primary sections confirmed. Here are the suggested subsections:"]
                for primary in primary_sections:
                    subs = (workflow.get("subsections") or {}).get(primary, [])
                    lines.append(f"- {primary}: {', '.join(subs)}")
                lines.append(
                    "\nYou can update subsections, ask me to break them down further, "
                    "or let me know if this structure looks good."
                )
                llm_response = "\n".join(lines)
            return self._report_workflow_response(
                response_text=llm_response,
                workflow=workflow,
                artifacts_created=[],
                artifacts_updated=[working_artifact_id],
            )

        if state == "awaiting_subsections":
            hierarchy = self._get_report_structure_hierarchy(workflow)
            interpreted = self._interpret_report_turn(
                workflow_state="awaiting_subsections",
                workflow=workflow,
                user_text=text,
                conversation_context=conversation_context,
                research_summary=research_summary,
            )
            if interpreted is None:
                interpreted = self._interpret_report_turn_fallback(
                    workflow_state="awaiting_subsections", user_text=text, workflow=workflow,
                )
            intent = (interpreted or {}).get("intent", "accept_structure")
            llm_response = (interpreted or {}).get("response_text", "")

            overrides: dict[str, list[str]] = {}
            if intent == "provide_overrides":
                overrides = (interpreted or {}).get("overrides") or {}

            if overrides:
                for parent_key, children in overrides.items():
                    hierarchy[parent_key] = children
                    parent_path = self._section_path_from_key(parent_key)
                    for child in children:
                        child_key = self._section_key_from_path([*parent_path, child])
                        hierarchy.setdefault(child_key, [])
                workflow["section_hierarchy"] = hierarchy
                self._sync_legacy_structure_fields(workflow)
                self._ensure_report_section_ids(workflow)

            if intent == "expand_deeper":
                expansion = self._expand_report_hierarchy_depth(
                    hierarchy=hierarchy,
                    topic=topic,
                    max_depth=4,
                    max_parents=8,
                )
                workflow["section_hierarchy"] = hierarchy
                self._sync_legacy_structure_fields(workflow)
                self._ensure_report_section_ids(workflow)
                self._mark_report_plan_updated(workflow, updated_by="agent", status="scaffolding")
                self._persist_report_workflow(session_id, session_meta, workflow)
                self._write_report_working_doc(session_id=session_id, workflow=workflow, stage="structure")
                self.events.create_event(
                    session_id=session_id,
                    message_id=assistant_message_id,
                    event_type="report_structure_confirmed",
                    payload={
                        "level": "nested_subsections",
                        "added_parent_count": len(expansion),
                        "current_max_depth": int(workflow.get("structure_max_depth") or 1),
                    },
                )
                self._emit_report_plan_state_updated_event(
                    session_id=session_id,
                    message_id=assistant_message_id,
                    workflow=workflow,
                    reason="nested_subsections_suggested",
                )
                if expansion:
                    if not llm_response:
                        lines = ["I've added nested subsection suggestions for:"]
                        for parent_key, children in list(expansion.items())[:8]:
                            parent_label = " > ".join(self._section_path_from_key(parent_key))
                            lines.append(f"- {parent_label}: {', '.join(children)}")
                        lines.append(
                            "\nFeel free to adjust these, ask for more depth, "
                            "or let me know when the structure looks good."
                        )
                        llm_response = "\n".join(lines)
                    return self._report_workflow_response(
                        response_text=llm_response,
                        workflow=workflow,
                        artifacts_created=[],
                        artifacts_updated=[working_artifact_id],
                    )
                return self._report_workflow_response(
                    response_text=llm_response or (
                        "No additional nested subsection suggestions were added. "
                        "You can make manual updates or let me know when the structure looks good."
                    ),
                    workflow=workflow,
                    artifacts_created=[],
                    artifacts_updated=[working_artifact_id],
                )

            if intent == "provide_overrides" and overrides:
                self._mark_report_plan_updated(workflow, updated_by="agent", status="scaffolding")
                self._persist_report_workflow(session_id, session_meta, workflow)
                self._write_report_working_doc(session_id=session_id, workflow=workflow, stage="structure")
                self._emit_report_plan_state_updated_event(
                    session_id=session_id,
                    message_id=assistant_message_id,
                    workflow=workflow,
                    reason="subsections_updated",
                )
                return self._report_workflow_response(
                    response_text=llm_response or (
                        "Subsection updates captured. You can expand further, "
                        "make more changes, or let me know when the structure looks good."
                    ),
                    workflow=workflow,
                    artifacts_created=[],
                    artifacts_updated=[working_artifact_id],
                )

            # intent == "accept_structure" or "other" with no overrides → finalize
            workflow["state"] = "awaiting_instructions"
            self._mark_report_plan_updated(workflow, updated_by="agent", status="scaffolding")
            self._persist_report_workflow(session_id, session_meta, workflow)
            self._write_report_working_doc(session_id=session_id, workflow=workflow, stage="structure")
            self.events.create_event(
                session_id=session_id,
                message_id=assistant_message_id,
                event_type="report_structure_confirmed",
                payload={
                    "level": "subsections",
                    "primary_sections": workflow.get("primary_sections") or [],
                    "subsections": workflow.get("subsections") or {},
                    "current_max_depth": int(workflow.get("structure_max_depth") or 1),
                },
            )
            self._emit_report_plan_state_updated_event(
                session_id=session_id,
                message_id=assistant_message_id,
                workflow=workflow,
                reason="structure_finalized",
            )
            return self._report_workflow_response(
                response_text=llm_response or (
                    "The structure looks good! Would you like to provide any specific "
                    "instructions for individual sections, or shall I use sensible defaults?"
                ),
                workflow=workflow,
                artifacts_created=[],
                artifacts_updated=[working_artifact_id],
            )

        if state == "awaiting_instructions":
            interpreted = self._interpret_report_turn(
                workflow_state="awaiting_instructions",
                workflow=workflow,
                user_text=text,
                conversation_context=conversation_context,
                research_summary=research_summary,
            )
            if interpreted is None:
                interpreted = self._interpret_report_turn_fallback(
                    workflow_state="awaiting_instructions", user_text=text, workflow=workflow,
                )
            intent = (interpreted or {}).get("intent", "use_defaults")

            if intent == "provide_instructions" and interpreted:
                llm_instructions = interpreted.get("instructions") or {}
                global_inst = interpreted.get("global_instruction")
                if llm_instructions or global_inst:
                    section_entries = self._list_report_section_entries(workflow)
                    all_keys = [entry["key"] for entry in section_entries]
                    instructions: dict[str, str] = {}
                    for key in all_keys:
                        matched = None
                        for candidate_key, inst_text in llm_instructions.items():
                            resolved = self._match_section_target(candidate_key, section_keys=all_keys)
                            if resolved == key:
                                matched = inst_text
                                break
                        if matched:
                            instructions[key] = matched
                        elif global_inst:
                            instructions[key] = global_inst
                        else:
                            instructions[key] = f"Provide clear, evidence-based analysis for {topic}."
                    provided_keys = sorted(k for k in all_keys if k in {
                        self._match_section_target(c, section_keys=all_keys)
                        for c in llm_instructions
                    })
                    defaulted_keys = sorted(k for k in all_keys if k not in set(provided_keys))
                    instruction_meta = {
                        "mode": "structured",
                        "provided_keys": provided_keys,
                        "defaulted_keys": defaulted_keys,
                    }
                else:
                    instructions, instruction_meta = self._resolve_section_instructions(workflow, text)
            else:
                instructions, instruction_meta = self._resolve_section_instructions(workflow, text)

            return self._generate_report_from_workflow(
                session_id=session_id,
                session_meta=session_meta,
                workflow=workflow,
                assistant_message_id=assistant_message_id,
                available_sources=available_sources,
                preferred_source_ids=preferred_source_ids,
                topic=topic,
                working_artifact_id=working_artifact_id,
                instructions=instructions,
                instruction_meta=instruction_meta,
            )

        if state in ("generating", "completed"):
            return self._report_workflow_response(
                response_text="Report workflow is already completed. Ask to create a new report to start another.",
                workflow=workflow,
                artifacts_created=[],
                artifacts_updated=[working_artifact_id],
            )

        return None

    def _persist_report_workflow(self, session_id: str, session_meta: dict[str, Any], workflow: dict[str, Any]) -> None:
        updated_meta = dict(session_meta)
        updated_meta["report_workflow"] = workflow
        self.sessions.update_session(session_id, metadata=updated_meta)

    def _report_workflow_response(
        self,
        *,
        response_text: str,
        workflow: dict[str, Any],
        artifacts_created: list[str],
        artifacts_updated: list[str],
    ) -> dict[str, Any]:
        return {
            "response_text": response_text,
            "content_json": {
                "text": response_text,
                "report_plan_card": self._build_report_plan_card(workflow),
            },
            "artifacts_created": artifacts_created,
            "artifacts_updated": artifacts_updated,
        }

    # LLM-driven report turn interpretation ---------------------------------

    _REPORT_TURN_INTENTS: dict[str, list[str]] = {
        "none": ["start_report", "not_report"],
        "awaiting_primary_sections": ["accept_suggested", "provide_sections", "other"],
        "awaiting_subsections": ["accept_structure", "provide_overrides", "expand_deeper", "other"],
        "awaiting_instructions": ["use_defaults", "provide_instructions", "other"],
    }

    def _build_report_turn_prompt(
        self,
        *,
        workflow_state: str | None,
        workflow: dict[str, Any] | None,
        user_text: str,
        conversation_context: list[dict[str, str]] | None = None,
        research_summary: str | None = None,
    ) -> str:
        state_key = workflow_state or "none"
        valid_intents = self._REPORT_TURN_INTENTS.get(state_key, ["other"])
        lines: list[str] = [
            "You are a report-workflow assistant that interprets user messages.",
            "Analyze the user message and return a JSON object with your classification.",
            "",
            f"CURRENT_WORKFLOW_STATE: {state_key}",
            f"VALID_INTENTS: {json.dumps(valid_intents)}",
            "",
        ]

        if state_key == "none":
            lines.extend([
                "INTENT_DESCRIPTIONS:",
                "- start_report: user wants to create/build/generate/draft a report or document",
                "- not_report: user message is NOT about creating a report (questions, research requests, etc.)",
                "",
                "When intent is start_report, extract:",
                "- topic: the report subject (from context clues, NOT the literal request text)",
                "- suggested_sections: 4-8 section titles relevant to the topic and available data",
                "",
            ])
        elif state_key == "awaiting_primary_sections":
            sections_display = json.dumps(workflow.get("primary_sections", []) if workflow else [])
            lines.extend([
                f"CURRENT_SUGGESTED_SECTIONS: {sections_display}",
                "",
                "INTENT_DESCRIPTIONS:",
                "- accept_suggested: user accepts the suggested sections (yes, ok, use those, looks good, etc.)",
                "- provide_sections: user provides their own section list",
                "- other: unclear or unrelated message",
                "",
                "When intent is provide_sections, extract:",
                "- sections: list of section title strings the user specified",
                "",
            ])
        elif state_key == "awaiting_subsections":
            hierarchy = self._get_report_structure_hierarchy(workflow) if workflow else {}
            compact_hierarchy = {k: v for k, v in hierarchy.items() if v}
            lines.extend([
                f"CURRENT_HIERARCHY: {json.dumps(compact_hierarchy)}",
                "",
                "INTENT_DESCRIPTIONS:",
                "- accept_structure: user is satisfied and wants to finalize the structure (done, looks good, proceed, final structure, etc.)",
                "- provide_overrides: user provides specific subsection changes (e.g. 'Section: sub1, sub2')",
                "- expand_deeper: user wants more nested subsections (deeper, expand, break down further, etc.)",
                "- other: unclear or unrelated message",
                "",
                "When intent is provide_overrides, extract:",
                "- overrides: dict mapping parent section paths to lists of child titles",
                "  Use '::' as path separator (e.g. 'Executive Summary::Purpose': ['sub1', 'sub2'])",
                "",
            ])
        elif state_key == "awaiting_instructions":
            lines.extend([
                "INTENT_DESCRIPTIONS:",
                "- use_defaults: user wants default/automatic instructions (use defaults, just go ahead, reasonable defaults, etc.)",
                "- provide_instructions: user provides specific section instructions",
                "- other: unclear or unrelated message",
                "",
                "When intent is provide_instructions, extract:",
                "- instructions: dict mapping section paths (using '::' separator) to instruction strings",
                "- global_instruction: optional overall instruction for all sections",
                "",
            ])

        if research_summary:
            lines.extend(["RESEARCH_CONTEXT:", research_summary, ""])

        if conversation_context:
            lines.append("RECENT_CONVERSATION:")
            for turn in conversation_context[-6:]:
                role = turn.get("role", "user")
                content = (turn.get("content") or "")[:300]
                lines.append(f"[{role}]: {content}")
            lines.append("")

        lines.extend([
            "RULES:",
            "- Affirmative phrases (yes, ok, sure, that works, use those, looks good, let's go) count as acceptance",
            "- If the user intent is unclear, use 'other' and ask for clarification in response_text",
            "- response_text must be a natural, friendly conversational response (NOT a template)",
            "- Do NOT include markdown code fences in your response",
            "",
            "Return JSON only:",
            json.dumps({
                "intent": "<one of the valid intents>",
                "response_text": "<natural response to the user>",
                "...": "<state-specific extracted data fields>",
            }),
        ])
        return "\n".join(lines)

    def _interpret_report_turn(
        self,
        *,
        workflow_state: str | None,
        workflow: dict[str, Any] | None,
        user_text: str,
        conversation_context: list[dict[str, str]] | None = None,
        research_summary: str | None = None,
    ) -> dict[str, Any] | None:
        state_key = workflow_state or "none"

        if state_key == "none":
            lower = user_text.lower()
            report_hints = ("report", "document", "write up", "writeup", "draft", "memo")
            if not any(hint in lower for hint in report_hints):
                return {"intent": "not_report", "response_text": ""}

        prompt = self._build_report_turn_prompt(
            workflow_state=workflow_state,
            workflow=workflow,
            user_text=user_text,
            conversation_context=conversation_context,
            research_summary=research_summary,
        )
        try:
            response = self.agent.chat(
                [{"role": "user", "content": user_text}],
                system_context=prompt,
            )
        except Exception:
            logger.warning("LLM call failed for report turn interpretation", exc_info=True)
            return None

        payload = self._extract_json_object_from_text(response)
        return self._normalize_report_turn_result(
            payload=payload,
            workflow_state=workflow_state,
            workflow=workflow,
        )

    def _normalize_report_turn_result(
        self,
        *,
        payload: dict[str, Any] | None,
        workflow_state: str | None,
        workflow: dict[str, Any] | None,
    ) -> dict[str, Any] | None:
        if not isinstance(payload, dict):
            return None
        state_key = workflow_state or "none"
        valid_intents = self._REPORT_TURN_INTENTS.get(state_key, ["other"])
        intent = str(payload.get("intent") or "").strip().lower()
        if intent not in valid_intents:
            return None
        response_text = str(payload.get("response_text") or "").strip()
        if not response_text and intent != "not_report":
            return None

        result: dict[str, Any] = {"intent": intent, "response_text": response_text}

        if state_key == "none" and intent == "start_report":
            result["topic"] = str(payload.get("topic") or "").strip()[:180] or "Untitled Report"
            raw_sections = payload.get("suggested_sections") or []
            if isinstance(raw_sections, list):
                result["suggested_sections"] = self._normalize_section_titles(
                    [str(s) for s in raw_sections if s]
                )[:8]
            else:
                result["suggested_sections"] = []

        elif state_key == "awaiting_primary_sections" and intent == "provide_sections":
            raw = payload.get("sections") or []
            if isinstance(raw, list):
                result["sections"] = self._normalize_section_titles(
                    [str(s) for s in raw if s]
                )
            else:
                result["sections"] = []
            if not result["sections"]:
                return None

        elif state_key == "awaiting_subsections" and intent == "provide_overrides":
            raw_overrides = payload.get("overrides") or {}
            if isinstance(raw_overrides, dict) and workflow:
                hierarchy = self._get_report_structure_hierarchy(workflow)
                normalized: dict[str, list[str]] = {}
                all_keys = [k for k in hierarchy.keys() if k]
                for parent_candidate, children in raw_overrides.items():
                    matched = self._match_section_target(
                        str(parent_candidate), section_keys=all_keys
                    )
                    if matched and isinstance(children, list):
                        titles = self._normalize_section_titles(
                            [str(c) for c in children if c]
                        )
                        if titles:
                            normalized[matched] = titles
                result["overrides"] = normalized
                if not normalized:
                    return None
            else:
                return None

        elif state_key == "awaiting_instructions" and intent == "provide_instructions":
            raw_instructions = payload.get("instructions") or {}
            if isinstance(raw_instructions, dict):
                result["instructions"] = {
                    str(k): str(v) for k, v in raw_instructions.items() if k and v
                }
            else:
                result["instructions"] = {}
            result["global_instruction"] = str(payload.get("global_instruction") or "").strip() or None

        return result

    def _interpret_report_turn_fallback(
        self,
        *,
        workflow_state: str | None,
        user_text: str,
        workflow: dict[str, Any] | None,
    ) -> dict[str, Any] | None:
        text = (user_text or "").strip()
        lower = text.lower()
        state_key = workflow_state or "none"

        if state_key == "none":
            if self._is_report_request(text):
                topic = self._extract_report_topic(text)
                sections = self._suggest_primary_sections(topic)
                return {
                    "intent": "start_report",
                    "topic": topic,
                    "suggested_sections": sections,
                    "response_text": "",
                }
            return {"intent": "not_report", "response_text": ""}

        if state_key == "awaiting_primary_sections":
            parsed = self._parse_list_from_text(text)
            if parsed:
                return {
                    "intent": "provide_sections",
                    "sections": parsed,
                    "response_text": "",
                }
            return {"intent": "accept_suggested", "response_text": ""}

        if state_key == "awaiting_subsections":
            if self._is_deeper_structure_request(lower):
                return {"intent": "expand_deeper", "response_text": ""}
            if self._is_finalize_structure_request(lower) or lower in (
                "use suggested", "yes", "ok", "okay",
            ):
                return {"intent": "accept_structure", "response_text": ""}
            hierarchy = self._get_report_structure_hierarchy(workflow) if workflow else {}
            overrides = self._parse_structure_overrides(text=text, hierarchy=hierarchy)
            if overrides:
                return {"intent": "provide_overrides", "overrides": overrides, "response_text": ""}
            shared = self._parse_list_from_text(text)
            if shared and workflow:
                hierarchy = self._get_report_structure_hierarchy(workflow)
                shared_normalized = self._normalize_section_titles(shared)
                computed_overrides: dict[str, list[str]] = {}
                for primary in hierarchy.get("", []):
                    pk = self._section_key_from_path([primary])
                    if pk:
                        computed_overrides[pk] = shared_normalized
                if computed_overrides:
                    return {"intent": "provide_overrides", "overrides": computed_overrides, "response_text": ""}
            return {"intent": "accept_structure", "response_text": ""}

        if state_key == "awaiting_instructions":
            if "use defaults" in lower or lower in ("default", "defaults", "no", "no instructions"):
                return {"intent": "use_defaults", "response_text": ""}
            return {"intent": "provide_instructions", "response_text": ""}

        return None

    @staticmethod
    def _summarize_research_for_prompt(
        internal_research: dict[str, Any],
        uploaded_research: dict[str, Any],
    ) -> str:
        parts: list[str] = []
        internal_summary = str(internal_research.get("summary_text") or "").strip()
        if internal_summary and internal_research.get("status") != "skipped":
            parts.append(f"Internal data research: {internal_summary[:500]}")
        source_results = internal_research.get("source_results") or []
        if isinstance(source_results, list):
            for sr in source_results[:4]:
                if isinstance(sr, dict):
                    name = str(sr.get("source_name") or sr.get("source_id") or "source")
                    finding = str(sr.get("finding_summary") or sr.get("summary") or "")[:200]
                    if finding:
                        parts.append(f"- {name}: {finding}")
        uploaded_summary = str(uploaded_research.get("summary_text") or "").strip()
        if uploaded_summary and uploaded_research.get("status") != "skipped":
            parts.append(f"Uploaded document research: {uploaded_summary[:500]}")
        return "\n".join(parts) if parts else ""

    def _emit_report_plan_state_updated_event(
        self,
        *,
        session_id: str,
        message_id: str,
        workflow: dict[str, Any],
        reason: str,
    ) -> None:
        self.events.create_event(
            session_id=session_id,
            message_id=message_id,
            event_type="report_plan_state_updated",
            payload={
                "plan_id": str(workflow.get("plan_id") or ""),
                "plan_status": str(workflow.get("plan_status") or "scaffolding"),
                "workflow_state": str(workflow.get("state") or ""),
                "updated_by": str(workflow.get("plan_updated_by") or "agent"),
                "updated_at": str(workflow.get("plan_updated_at") or datetime.now(UTC).isoformat()),
                "section_count": len(self._list_report_section_entries(workflow)),
                "reason": reason,
            },
        )

    def _mark_report_plan_updated(
        self,
        workflow: dict[str, Any],
        *,
        updated_by: str,
        status: Any | None = None,
    ) -> None:
        status_value = str(status or "").strip().lower()
        workflow.setdefault("plan_id", str(uuid4()))
        workflow.setdefault("plan_title", str(workflow.get("topic") or "Report")[:180])
        workflow.setdefault("plan_summary", str(workflow.get("topic") or "Report")[:600])
        if status_value in {"scaffolding", "ready", "generating", "completed"}:
            workflow["plan_status"] = status_value
        else:
            workflow["plan_status"] = workflow.get("plan_status") or self._report_plan_status_from_workflow_state(
                str(workflow.get("state") or "")
            )
        workflow["plan_updated_by"] = "user" if str(updated_by).lower() == "user" else "agent"
        workflow["plan_updated_at"] = datetime.now(UTC).isoformat()

    @staticmethod
    def _report_plan_status_from_workflow_state(state: str) -> str:
        mapping = {
            "awaiting_primary_sections": "scaffolding",
            "awaiting_subsections": "scaffolding",
            "awaiting_instructions": "scaffolding",
            "generating": "generating",
            "completed": "completed",
        }
        return mapping.get(state, "scaffolding")

    @staticmethod
    def _workflow_state_from_plan_status(status: str) -> str | None:
        mapping = {
            "ready": "awaiting_instructions",
            "generating": "generating",
            "completed": "completed",
        }
        return mapping.get(status)

    def _ensure_report_section_ids(self, workflow: dict[str, Any]) -> dict[str, str]:
        hierarchy = self._get_report_structure_hierarchy(workflow)
        raw_ids = workflow.get("section_ids")
        section_ids: dict[str, str] = dict(raw_ids) if isinstance(raw_ids, dict) else {}
        valid_keys = {key for key in hierarchy.keys() if key}
        section_ids = {key: value for key, value in section_ids.items() if key in valid_keys and str(value).strip()}
        for key in sorted(valid_keys):
            if key not in section_ids:
                section_ids[key] = str(uuid4())
        workflow["section_ids"] = section_ids
        return section_ids

    def _build_report_plan_card(self, workflow: dict[str, Any]) -> dict[str, Any]:
        if not workflow.get("plan_id"):
            workflow["plan_id"] = str(uuid4())
        hierarchy = self._get_report_structure_hierarchy(workflow)
        section_ids = self._ensure_report_section_ids(workflow)
        instruction_map = workflow.get("section_instructions")
        if not isinstance(instruction_map, dict):
            instruction_map = {}
        instruction_sources = workflow.get("section_instruction_sources")
        if not isinstance(instruction_sources, dict):
            instruction_sources = {}
        plan_status = str(
            workflow.get("plan_status")
            or self._report_plan_status_from_workflow_state(str(workflow.get("state") or ""))
        )
        if plan_status not in {"scaffolding", "ready", "generating", "completed"}:
            plan_status = "scaffolding"
        is_completed = plan_status == "completed" or str(workflow.get("state") or "") == "completed"

        def build_children(parent_key: str, parent_path: list[str]) -> list[dict[str, Any]]:
            nodes: list[dict[str, Any]] = []
            for title in hierarchy.get(parent_key, []):
                path = [*parent_path, title]
                section_key = self._section_key_from_path(path)
                instruction_value = str(instruction_map.get(section_key) or "").strip()
                instruction_text = instruction_value or None
                source_value = str(instruction_sources.get(section_key) or "").lower()
                if source_value not in {"user", "agent"}:
                    source_value = ""
                if is_completed:
                    section_status = "generated"
                elif instruction_text:
                    section_status = "has_instructions"
                else:
                    section_status = "pending"
                nodes.append(
                    {
                        "section_id": section_ids.get(section_key, str(uuid4())),
                        "title": title,
                        "depth": max(len(path) - 1, 0),
                        "instructions": instruction_text,
                        "instruction_source": source_value or None,
                        "status": section_status,
                        "subsections": build_children(section_key, path),
                    }
                )
            return nodes

        title = str(workflow.get("plan_title") or workflow.get("topic") or "Report")
        summary = str(workflow.get("plan_summary") or "")
        return {
            "plan_id": str(workflow.get("plan_id") or ""),
            "title": title,
            "summary": summary,
            "status": plan_status,
            "sections": build_children("", []),
            "updated_by": str(workflow.get("plan_updated_by") or "agent"),
            "updated_at": str(workflow.get("plan_updated_at") or datetime.now(UTC).isoformat()),
        }

    def _apply_report_plan_state_to_workflow(
        self,
        workflow: dict[str, Any],
        report_plan_state: dict[str, Any],
    ) -> dict[str, Any]:
        updated = dict(workflow)

        plan_id = str(report_plan_state.get("plan_id") or "").strip()
        if plan_id:
            updated["plan_id"] = plan_id

        title = str(report_plan_state.get("title") or "").strip()
        if title:
            updated["plan_title"] = title[:180]
            if not str(updated.get("topic") or "").strip():
                updated["topic"] = title[:180]

        summary = str(report_plan_state.get("summary") or "").strip()
        if summary:
            updated["plan_summary"] = summary[:600]

        status = str(report_plan_state.get("status") or "").strip().lower()
        if status in {"scaffolding", "ready", "generating", "completed"}:
            updated["plan_status"] = status
            mapped_state = self._workflow_state_from_plan_status(status)
            if mapped_state:
                updated["state"] = mapped_state
            if status == "completed":
                updated["active"] = False
            elif status in {"scaffolding", "ready", "generating"}:
                updated["active"] = True

        incoming_sections = report_plan_state.get("sections")
        if isinstance(incoming_sections, list):
            hierarchy: dict[str, list[str]] = {"": []}
            section_ids: dict[str, str] = {}
            instructions: dict[str, str] = {}
            instruction_sources: dict[str, str] = {}

            def walk(nodes: list[dict[str, Any]], parent_key: str, parent_path: list[str]) -> None:
                for node in nodes:
                    if not isinstance(node, dict):
                        continue
                    raw_title = str(node.get("title") or "").strip()
                    if len(raw_title) < 2:
                        continue
                    titles = hierarchy.setdefault(parent_key, [])
                    if raw_title in titles:
                        continue
                    titles.append(raw_title)
                    path = [*parent_path, raw_title]
                    section_key = self._section_key_from_path(path)
                    hierarchy.setdefault(section_key, [])

                    raw_section_id = str(node.get("section_id") or "").strip()
                    section_ids[section_key] = raw_section_id or str(uuid4())

                    raw_instruction = node.get("instructions")
                    if isinstance(raw_instruction, str) and raw_instruction.strip():
                        instructions[section_key] = raw_instruction.strip()
                        source_value = str(node.get("instruction_source") or "").strip().lower()
                        instruction_sources[section_key] = source_value if source_value in {"user", "agent"} else "user"

                    child_nodes = node.get("subsections")
                    if isinstance(child_nodes, list) and child_nodes:
                        walk(child_nodes, section_key, path)

            walk(incoming_sections, "", [])
            for key, titles in list(hierarchy.items()):
                hierarchy[key] = self._normalize_section_titles(titles)
            updated["section_hierarchy"] = hierarchy
            self._sync_legacy_structure_fields(updated)
            updated["section_ids"] = section_ids
            updated["section_instructions"] = instructions
            updated["section_instruction_sources"] = instruction_sources
            self._ensure_report_section_ids(updated)

        return updated

    def _default_instruction_for_section_entry(self, *, topic: str, entry: dict[str, Any]) -> str:
        depth = int(entry.get("depth") or 1)
        if depth <= 1:
            return f"Provide clear, evidence-based analysis for {topic}."
        section_label = " > ".join(entry.get("path") or [str(entry.get("title") or "Section")])
        return f"Summarize {section_label} for {topic} with clear findings, evidence, and actions."

    def _resolve_start_now_section_instructions(self, workflow: dict[str, Any]) -> tuple[dict[str, str], dict[str, Any]]:
        topic = str(workflow.get("topic") or "Report")
        section_entries = self._list_report_section_entries(workflow)
        existing_instructions = workflow.get("section_instructions")
        if not isinstance(existing_instructions, dict):
            existing_instructions = {}
        existing_sources = workflow.get("section_instruction_sources")
        if not isinstance(existing_sources, dict):
            existing_sources = {}

        instructions: dict[str, str] = {}
        instruction_sources: dict[str, str] = {}
        provided_keys: list[str] = []
        defaulted_keys: list[str] = []
        for entry in section_entries:
            key = entry["key"]
            instruction_text = str(existing_instructions.get(key) or "").strip()
            if instruction_text:
                instructions[key] = instruction_text
                source = str(existing_sources.get(key) or "user").strip().lower()
                instruction_sources[key] = source if source in {"user", "agent"} else "user"
                provided_keys.append(key)
                continue
            instructions[key] = self._default_instruction_for_section_entry(topic=topic, entry=entry)
            instruction_sources[key] = "agent"
            defaulted_keys.append(key)

        mode = "start_now_existing" if not defaulted_keys else "start_now_defaults"
        return instructions, {
            "mode": mode,
            "provided_keys": sorted(provided_keys),
            "defaulted_keys": sorted(defaulted_keys),
            "instruction_sources": instruction_sources,
        }

    def _generate_report_from_workflow(
        self,
        *,
        session_id: str,
        session_meta: dict[str, Any],
        workflow: dict[str, Any],
        assistant_message_id: str,
        available_sources: list[dict[str, Any]],
        preferred_source_ids: list[str],
        topic: str,
        working_artifact_id: str,
        instructions: dict[str, str],
        instruction_meta: dict[str, Any],
    ) -> dict[str, Any]:
        workflow["section_instructions"] = dict(instructions)
        instruction_sources = workflow.get("section_instruction_sources")
        if not isinstance(instruction_sources, dict):
            instruction_sources = {}
        source_updates = instruction_meta.get("instruction_sources")
        if isinstance(source_updates, dict):
            for key, value in source_updates.items():
                source_value = str(value or "").strip().lower()
                if source_value in {"user", "agent"}:
                    instruction_sources[str(key)] = source_value
        workflow["section_instruction_sources"] = instruction_sources

        self.events.create_event(
            session_id=session_id,
            message_id=assistant_message_id,
            event_type="report_section_instructions_captured",
            payload={
                "instruction_count": len(instructions),
                "mode": instruction_meta.get("mode", "unknown"),
                "provided_keys": list(instruction_meta.get("provided_keys") or []),
                "defaulted_keys": list(instruction_meta.get("defaulted_keys") or []),
            },
        )

        workflow["state"] = "generating"
        self._mark_report_plan_updated(workflow, updated_by="agent", status="generating")
        self._persist_report_workflow(session_id, session_meta, workflow)
        self._emit_report_plan_state_updated_event(
            session_id=session_id,
            message_id=assistant_message_id,
            workflow=workflow,
            reason="generation_started",
        )
        self._write_report_working_doc(
            session_id=session_id,
            workflow=workflow,
            stage="instructions",
            instructions=instructions,
        )
        self.events.create_event(
            session_id=session_id,
            message_id=assistant_message_id,
            event_type="report_generation_started",
            payload={"topic": topic, "working_artifact_id": working_artifact_id},
        )

        generation_result = self._run_generate_report_document_tool(
            session_id=session_id,
            assistant_message_id=assistant_message_id,
            workflow=workflow,
            available_sources=available_sources,
            preferred_source_ids=preferred_source_ids,
        )
        if generation_result.get("status") != "completed":
            workflow["state"] = "awaiting_instructions"
            workflow["last_generation_error"] = str(generation_result.get("error") or "generate_report_document tool failed")
            self._mark_report_plan_updated(workflow, updated_by="agent", status="scaffolding")
            self._persist_report_workflow(session_id, session_meta, workflow)
            self._emit_report_plan_state_updated_event(
                session_id=session_id,
                message_id=assistant_message_id,
                workflow=workflow,
                reason="generation_failed",
            )
            return {
                "error": True,
                "error_code": "TOOL_EXECUTION_FAILED",
                "error_message": workflow["last_generation_error"],
            }

        generated_content = generation_result.get("generated_content") or {}
        workflow.pop("last_generation_error", None)
        self._write_report_working_doc(
            session_id=session_id,
            workflow=workflow,
            stage="content",
            content_map=generated_content,
        )
        self.artifacts.update_artifact(
            session_id=session_id,
            artifact_id=working_artifact_id,
            lifecycle_state="final",
            metadata={
                "report_topic": topic,
                "report_state": "generated",
                "generated_at": datetime.now(UTC).isoformat(),
            },
        )
        self.events.create_event(
            session_id=session_id,
            message_id=assistant_message_id,
            event_type="artifact_updated",
            payload={
                "artifact_id": working_artifact_id,
                "reason": "report_content_generated",
            },
        )
        final_artifact = self._create_report_final_artifact(session_id=session_id, workflow=workflow)
        export_result = self._run_export_report_document_tool(
            session_id=session_id,
            assistant_message_id=assistant_message_id,
            report_artifact=final_artifact,
            requested_formats=["docx", "pdf", "xlsx"],
        )
        self.events.create_event(
            session_id=session_id,
            message_id=assistant_message_id,
            event_type="report_generation_completed",
            payload={
                "working_artifact_id": working_artifact_id,
                "final_artifact_id": final_artifact["artifact_id"] if final_artifact else None,
                "research_artifact_count": generation_result.get("research_artifact_count", 0),
                "assigned_research_count": generation_result.get("assigned_research_count", 0),
                "gap_fill_call_count": generation_result.get("gap_fill_call_count", 0),
                "export_artifact_ids": export_result.get("artifacts_created", []),
            },
        )

        workflow["state"] = "completed"
        workflow["active"] = False
        self._mark_report_plan_updated(workflow, updated_by="agent", status="completed")
        self._persist_report_workflow(session_id, session_meta, workflow)
        self._emit_report_plan_state_updated_event(
            session_id=session_id,
            message_id=assistant_message_id,
            workflow=workflow,
            reason="generation_completed",
        )
        response_lines = [
            "Report generation completed.",
            f"Working report document updated: {working_artifact_id}.",
        ]
        artifacts_created: list[str] = list(generation_result.get("gap_fill_artifacts_created") or [])
        if final_artifact:
            artifacts_created.append(final_artifact["artifact_id"])
            response_lines.append(f"Final report artifact created: {final_artifact['artifact_id']}.")
        exported_artifacts = export_result.get("artifacts_created") or []
        if exported_artifacts:
            artifacts_created.extend(exported_artifacts)
            response_lines.append(f"Export artifacts created: {len(exported_artifacts)}.")
        return self._report_workflow_response(
            response_text="\n".join(response_lines),
            workflow=workflow,
            artifacts_created=artifacts_created,
            artifacts_updated=[working_artifact_id],
        )

    def _create_report_working_artifact(
        self,
        *,
        session_id: str,
        topic: str,
        created_from_message_id: str | None,
    ) -> dict[str, Any] | None:
        try:
            from docx import Document

            output_dir = self._report_output_dir(session_id)
            output_dir.mkdir(parents=True, exist_ok=True)
            slug = self._safe_slug(topic)
            file_path = output_dir / f"{slug}-{uuid4().hex[:8]}-working.docx"
            doc = Document()
            doc.add_heading(f"Report: {topic}", level=0)
            doc.add_paragraph("Report structure and instructions will be refined through this session.")
            doc.save(str(file_path))
            artifact = self.artifacts.create_artifact(
                session_id=session_id,
                artifact_type="report_working_doc",
                lifecycle_state="in_progress",
                format="docx",
                filename=file_path.name,
                storage_uri=str(file_path),
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                size_bytes=file_path.stat().st_size,
                created_from_message_id=created_from_message_id,
                metadata={"report_topic": topic, "report_state": "scaffolding"},
            )
            return artifact
        except Exception:
            return None

    def _create_report_final_artifact(self, *, session_id: str, workflow: dict[str, Any]) -> dict[str, Any] | None:
        working_artifact_id = workflow.get("working_artifact_id")
        if not working_artifact_id:
            return None
        working = self.artifacts.get_artifact(session_id, working_artifact_id)
        if working is None:
            return None
        source_path = Path(working["storage_uri"])
        if not source_path.exists():
            return None
        final_path = source_path.with_name(source_path.stem.replace("-working", "") + "-final.docx")
        shutil.copy2(source_path, final_path)
        return self.artifacts.create_artifact(
            session_id=session_id,
            artifact_type="report_final_doc",
            lifecycle_state="final",
            format="docx",
            filename=final_path.name,
            storage_uri=str(final_path),
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            size_bytes=final_path.stat().st_size,
            source_artifact_id=working_artifact_id,
            metadata={"report_topic": workflow.get("topic", "Report"), "report_state": "final"},
        )

    def _run_export_report_document_tool(
        self,
        *,
        session_id: str,
        assistant_message_id: str,
        report_artifact: dict[str, Any] | None,
        requested_formats: list[str] | None = None,
    ) -> dict[str, Any]:
        tool_name = "export_report_document"
        call_id = str(uuid4())
        formats = self._normalize_export_formats(requested_formats)
        report_artifact_id = str(report_artifact.get("artifact_id")) if isinstance(report_artifact, dict) else None
        self.events.create_event(
            session_id=session_id,
            message_id=assistant_message_id,
            event_type="tool_call_request",
            payload={
                "call_id": call_id,
                "tool_name": tool_name,
                "arguments": {
                    "report_artifact_id": report_artifact_id,
                    "requested_formats": formats,
                },
            },
        )

        if not report_artifact:
            result = {
                "status": "skipped",
                "reason": "missing_report_artifact",
                "requested_formats": formats,
                "artifacts_created": [],
            }
            self.events.create_event(
                session_id=session_id,
                message_id=assistant_message_id,
                event_type="tool_call_response",
                payload={"call_id": call_id, "tool_name": tool_name, "result": result},
            )
            return result

        source_path = Path(str(report_artifact.get("storage_uri") or ""))
        if not source_path.exists() or not source_path.is_file():
            result = {
                "status": "failed",
                "error": f"report artifact file not found: {source_path}",
                "requested_formats": formats,
                "artifacts_created": [],
            }
            self.events.create_event(
                session_id=session_id,
                message_id=assistant_message_id,
                event_type="tool_call_response",
                payload={"call_id": call_id, "tool_name": tool_name, "result": result},
            )
            return result

        export_group_id = str(uuid4())
        topic = str((report_artifact.get("metadata") or {}).get("report_topic") or "Report")
        created_ids: list[str] = []
        exports: list[dict[str, str]] = []
        timestamp = datetime.now(UTC).strftime("%Y%m%d-%H%M%S")

        for export_format in formats:
            export_filename = (
                f"{self._safe_slug(source_path.stem)}-export-{timestamp}-{uuid4().hex[:6]}.{export_format}"
            )
            export_path = source_path.with_name(export_filename)
            mime_type = self._mime_for_export_format(export_format)
            export_detail = self._materialize_report_export(
                source_path=source_path,
                export_path=export_path,
                export_format=export_format,
                topic=topic,
                source_filename=str(report_artifact.get("filename") or source_path.name),
            )

            export_artifact = self.artifacts.create_artifact(
                session_id=session_id,
                artifact_group_id=export_group_id,
                artifact_type="export_file",
                lifecycle_state="final",
                format=export_format,
                filename=export_filename,
                storage_uri=str(export_path),
                mime_type=mime_type,
                size_bytes=export_path.stat().st_size,
                created_from_message_id=assistant_message_id,
                source_artifact_id=report_artifact["artifact_id"],
                metadata={
                    "report_topic": topic,
                    "export_format": export_format,
                    "export_mode": export_detail["mode"],
                    "export_backend": export_detail["backend"],
                    **({"export_note": export_detail["note"]} if export_detail.get("note") else {}),
                },
            )
            self.events.create_event(
                session_id=session_id,
                message_id=assistant_message_id,
                event_type="artifact_created",
                payload={
                    "artifact_id": export_artifact["artifact_id"],
                    "artifact_group_id": export_group_id,
                    "artifact_type": "export_file",
                    "source_artifact_id": report_artifact["artifact_id"],
                    "format": export_format,
                },
            )
            created_ids.append(export_artifact["artifact_id"])
            exports.append(
                {
                    "artifact_id": export_artifact["artifact_id"],
                    "format": export_format,
                    "filename": export_filename,
                    "mode": export_detail["mode"],
                    "backend": export_detail["backend"],
                }
            )

        result = {
            "status": "completed",
            "report_artifact_id": report_artifact["artifact_id"],
            "requested_formats": formats,
            "exports": exports,
            "artifacts_created": created_ids,
        }
        self.events.create_event(
            session_id=session_id,
            message_id=assistant_message_id,
            event_type="tool_call_response",
            payload={"call_id": call_id, "tool_name": tool_name, "result": result},
        )
        return result

    @staticmethod
    def _normalize_export_formats(requested_formats: list[str] | None) -> list[str]:
        allowed = ("docx", "pdf", "xlsx")
        incoming = requested_formats or []
        out: list[str] = []
        seen: set[str] = set()
        for raw in incoming:
            fmt = str(raw or "").strip().lower()
            if fmt not in allowed or fmt in seen:
                continue
            seen.add(fmt)
            out.append(fmt)
        if out:
            return out
        return ["docx"]

    def _resolve_research_output_format(
        self,
        *,
        source: dict[str, Any],
        retriever_output: dict[str, Any],
    ) -> str:
        candidates: list[Any] = []
        if isinstance(retriever_output, dict):
            candidates.extend(
                [
                    retriever_output.get("research_output_format"),
                    retriever_output.get("output_format"),
                    retriever_output.get("preferred_output_format"),
                ]
            )
        location = source.get("location") if isinstance(source, dict) else None
        if isinstance(location, dict):
            candidates.extend(
                [
                    location.get("research_output_format"),
                    location.get("output_format"),
                    location.get("default_output_format"),
                ]
            )
        schema_json = source.get("schema_json") if isinstance(source, dict) else None
        if isinstance(schema_json, dict):
            candidates.extend(
                [
                    schema_json.get("research_output_format"),
                    schema_json.get("output_format"),
                    schema_json.get("default_output_format"),
                ]
            )
        for candidate in candidates:
            if candidate is None:
                continue
            if isinstance(candidate, str) and not candidate.strip():
                continue
            normalized = self._normalize_output_doc_format(candidate)
            if normalized in ("docx", "pdf", "xlsx"):
                return normalized
        return "docx"

    @staticmethod
    def _normalize_output_doc_format(value: Any) -> str:
        raw = str(value or "").strip().lower()
        if raw in ("xlsx", "excel", "xls"):
            return "xlsx"
        if raw in ("pdf",):
            return "pdf"
        if raw in ("docx", "word", "doc"):
            return "docx"
        return "docx"

    @staticmethod
    def _mime_for_export_format(export_format: str) -> str:
        mapping = {
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "pdf": "application/pdf",
            "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        }
        return mapping.get(export_format, "application/octet-stream")

    def _materialize_report_export(
        self,
        *,
        source_path: Path,
        export_path: Path,
        export_format: str,
        topic: str,
        source_filename: str,
    ) -> dict[str, str]:
        if export_format == "docx":
            shutil.copy2(source_path, export_path)
            return {"mode": "native_copy", "backend": "filesystem_copy"}
        if export_format == "xlsx":
            backend = self._export_docx_to_xlsx(
                source_path=source_path,
                export_path=export_path,
                topic=topic,
            )
            if backend == "openpyxl":
                return {"mode": "native_convert", "backend": backend}
            return {
                "mode": "fallback_convert",
                "backend": backend,
                "note": "openpyxl unavailable; generated minimal OOXML workbook.",
            }
        if export_format == "pdf":
            return self._export_docx_to_pdf(
                source_path=source_path,
                export_path=export_path,
                topic=topic,
                source_filename=source_filename,
            )

        export_path.write_text(
            self._mock_export_text(
                topic=topic,
                source_filename=source_filename,
                export_format=export_format,
            ),
            encoding="utf-8",
        )
        return {"mode": "fallback_text", "backend": "utf8_text"}

    def _export_docx_to_xlsx(self, *, source_path: Path, export_path: Path, topic: str) -> str:
        from docx import Document

        doc = Document(str(source_path))
        rows: list[list[str]] = [["paragraph_index", "style_name", "section_path", "text"]]
        current_h1 = ""
        current_h2 = ""
        for index, paragraph in enumerate(doc.paragraphs):
            text = (paragraph.text or "").strip()
            if not text:
                continue
            style_name = paragraph.style.name if paragraph.style else ""
            style_lower = style_name.lower()
            if style_lower == "heading 1":
                current_h1 = text
                current_h2 = ""
            elif style_lower == "heading 2":
                current_h2 = text
            section_parts = [part for part in (current_h1, current_h2) if part]
            section_path = " > ".join(section_parts)
            rows.append([str(index), style_name, section_path, text])

        metadata_rows = [
            ["key", "value"],
            ["topic", topic],
            ["source_docx", str(source_path)],
            ["generated_at", datetime.now(UTC).isoformat()],
        ]

        try:
            from openpyxl import Workbook
        except ModuleNotFoundError:
            self._write_minimal_xlsx(
                export_path=export_path,
                sheet_name="Report Content",
                rows=rows,
                metadata_rows=metadata_rows,
            )
            return "minimal_ooxml"

        wb = Workbook()
        sheet = wb.active
        sheet.title = "Report Content"
        for row in rows:
            sheet.append(row)
        sheet.freeze_panes = "A2"
        sheet.column_dimensions["A"].width = 18
        sheet.column_dimensions["B"].width = 24
        sheet.column_dimensions["C"].width = 44
        sheet.column_dimensions["D"].width = 120

        meta = wb.create_sheet("Metadata")
        for row in metadata_rows:
            meta.append(row)
        wb.save(str(export_path))
        return "openpyxl"

    @staticmethod
    def _write_minimal_xlsx(
        *,
        export_path: Path,
        sheet_name: str,
        rows: list[list[str]],
        metadata_rows: list[list[str]] | None = None,
    ) -> None:
        from io import BytesIO
        from xml.sax.saxutils import escape
        import zipfile

        namespace_main = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
        namespace_rel = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        namespace_pkg_rel = "http://schemas.openxmlformats.org/package/2006/relationships"

        def col_name(index: int) -> str:
            label = ""
            current = index + 1
            while current:
                current, rem = divmod(current - 1, 26)
                label = chr(65 + rem) + label
            return label

        def build_sheet_xml(sheet_rows: list[list[str]]) -> str:
            sheet_data_parts: list[str] = []
            for row_idx, row in enumerate(sheet_rows, start=1):
                cell_parts: list[str] = []
                for col_idx, value in enumerate(row):
                    text = str(value or "")
                    ref = f"{col_name(col_idx)}{row_idx}"
                    if col_idx == 0 and row_idx > 1 and text.isdigit():
                        cell_parts.append(f'<c r="{ref}"><v>{text}</v></c>')
                    else:
                        safe = escape(text)
                        cell_parts.append(
                            f'<c r="{ref}" t="inlineStr"><is><t>{safe}</t></is></c>'
                        )
                sheet_data_parts.append(f"<row r=\"{row_idx}\">{''.join(cell_parts)}</row>")
            return (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                f'<worksheet xmlns="{namespace_main}">'
                f"<sheetData>{''.join(sheet_data_parts)}</sheetData>"
                "</worksheet>"
            )

        clean_sheet_name = (sheet_name or "Sheet1").replace('"', "'")
        metadata = metadata_rows or [["key", "value"]]

        content_types = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            '<Default Extension="xml" ContentType="application/xml"/>'
            '<Override PartName="/xl/workbook.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>'
            '<Override PartName="/xl/worksheets/sheet1.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>'
            '<Override PartName="/xl/worksheets/sheet2.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>'
            "</Types>"
        )
        root_rels = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            f'<Relationships xmlns="{namespace_pkg_rel}">'
            '<Relationship Id="rId1" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
            'Target="xl/workbook.xml"/>'
            "</Relationships>"
        )
        workbook = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            f'<workbook xmlns="{namespace_main}" xmlns:r="{namespace_rel}">'
            "<sheets>"
            f'<sheet name="{escape(clean_sheet_name)}" sheetId="1" r:id="rId1"/>'
            '<sheet name="Metadata" sheetId="2" r:id="rId2"/>'
            "</sheets>"
            "</workbook>"
        )
        workbook_rels = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            f'<Relationships xmlns="{namespace_pkg_rel}">'
            '<Relationship Id="rId1" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" '
            'Target="worksheets/sheet1.xml"/>'
            '<Relationship Id="rId2" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" '
            'Target="worksheets/sheet2.xml"/>'
            "</Relationships>"
        )
        sheet1 = build_sheet_xml(rows or [["No data"]])
        sheet2 = build_sheet_xml(metadata)

        buffer = BytesIO()
        with zipfile.ZipFile(buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("[Content_Types].xml", content_types)
            zf.writestr("_rels/.rels", root_rels)
            zf.writestr("xl/workbook.xml", workbook)
            zf.writestr("xl/_rels/workbook.xml.rels", workbook_rels)
            zf.writestr("xl/worksheets/sheet1.xml", sheet1)
            zf.writestr("xl/worksheets/sheet2.xml", sheet2)
        export_path.write_bytes(buffer.getvalue())

    def _export_docx_to_pdf(
        self,
        *,
        source_path: Path,
        export_path: Path,
        topic: str,
        source_filename: str,
    ) -> dict[str, str]:
        try:
            from word_engine.config import EngineConfig
            from word_engine.service import WordDocumentService

            service = WordDocumentService(
                config=EngineConfig(
                    allowed_roots=[Path(root) for root in self.allowed_roots] if self.allowed_roots else [Path.cwd()],
                    contract_version="v1",
                )
            )
            converted = service.convert_to_pdf(str(source_path), str(export_path))
            if (
                converted.get("status") == "ok"
                and export_path.exists()
                and export_path.stat().st_size > 0
            ):
                return {
                    "mode": "native_convert",
                    "backend": str(converted.get("method") or "word_engine_pdf"),
                }
            note = str(converted.get("message") or "word_engine conversion returned non-ok status")
        except Exception as exc:  # noqa: BLE001
            note = f"word_engine conversion unavailable: {exc}"

        lines = self._extract_docx_text_lines(source_path=source_path, max_lines=320)
        if not lines:
            lines = [f"Report export generated for {topic}.", f"Source file: {source_filename}"]
        self._write_minimal_pdf(export_path=export_path, title=topic, lines=lines)
        return {
            "mode": "fallback_convert",
            "backend": "minimal_pdf_writer",
            "note": note[:800],
        }

    @staticmethod
    def _extract_docx_text_lines(*, source_path: Path, max_lines: int) -> list[str]:
        from docx import Document

        doc = Document(str(source_path))
        lines: list[str] = []
        for paragraph in doc.paragraphs:
            text = (paragraph.text or "").strip()
            if not text:
                continue
            style_name = paragraph.style.name if paragraph.style else ""
            if style_name.lower().startswith("heading"):
                lines.append(f"[{style_name}] {text}")
            else:
                lines.append(text)
            if len(lines) >= max_lines:
                break
        return lines

    def _write_minimal_pdf(self, *, export_path: Path, title: str, lines: list[str]) -> None:
        wrapped_lines = [f"Report: {title}", ""] + self._wrap_pdf_lines(lines, max_chars=92)
        lines_per_page = 46
        pages = [wrapped_lines[index : index + lines_per_page] for index in range(0, len(wrapped_lines), lines_per_page)]
        if not pages:
            pages = [["Report export"]]

        page_count = len(pages)
        page_object_nums: list[int] = []
        content_object_nums: list[int] = []
        for index in range(page_count):
            page_object_nums.append(3 + index * 2)
            content_object_nums.append(4 + index * 2)
        font_object_num = 3 + page_count * 2

        objects: list[bytes] = []
        objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
        kids = " ".join(f"{number} 0 R" for number in page_object_nums)
        objects.append(f"<< /Type /Pages /Kids [{kids}] /Count {page_count} >>".encode("ascii"))

        for index, page_lines in enumerate(pages):
            page_obj = (
                f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                f"/Resources << /Font << /F1 {font_object_num} 0 R >> >> "
                f"/Contents {content_object_nums[index]} 0 R >>"
            ).encode("ascii")
            objects.append(page_obj)
            stream_bytes = self._build_pdf_stream(page_lines)
            stream_obj = (
                f"<< /Length {len(stream_bytes)} >>\nstream\n".encode("ascii")
                + stream_bytes
                + b"\nendstream"
            )
            objects.append(stream_obj)

        objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
        export_path.write_bytes(self._assemble_pdf(objects))

    @staticmethod
    def _wrap_pdf_lines(lines: list[str], *, max_chars: int) -> list[str]:
        wrapped: list[str] = []
        for raw in lines:
            text = raw.strip()
            if not text:
                wrapped.append("")
                continue
            while len(text) > max_chars:
                split_at = text.rfind(" ", 0, max_chars)
                if split_at <= 0:
                    split_at = max_chars
                wrapped.append(text[:split_at].rstrip())
                text = text[split_at:].strip()
            wrapped.append(text)
        return wrapped

    @staticmethod
    def _build_pdf_stream(lines: list[str]) -> bytes:
        commands = ["BT", "/F1 11 Tf", "50 760 Td", "14 TL"]
        for line in lines:
            escaped = (
                line.replace("\\", "\\\\")
                .replace("(", "\\(")
                .replace(")", "\\)")
            )
            commands.append(f"({escaped}) Tj")
            commands.append("T*")
        commands.append("ET")
        return "\n".join(commands).encode("latin-1", errors="replace")

    @staticmethod
    def _assemble_pdf(objects: list[bytes]) -> bytes:
        payload = bytearray()
        payload.extend(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
        offsets = [0]
        for index, obj in enumerate(objects, start=1):
            offsets.append(len(payload))
            payload.extend(f"{index} 0 obj\n".encode("ascii"))
            payload.extend(obj)
            payload.extend(b"\nendobj\n")

        xref_offset = len(payload)
        payload.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
        payload.extend(b"0000000000 65535 f \n")
        for offset in offsets[1:]:
            payload.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
        payload.extend(
            (
                f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
                f"startxref\n{xref_offset}\n%%EOF\n"
            ).encode("ascii")
        )
        return bytes(payload)

    @staticmethod
    def _mock_export_text(*, topic: str, source_filename: str, export_format: str) -> str:
        return (
            f"Mock {export_format.upper()} export for report topic: {topic}\n"
            f"Source report file: {source_filename}\n"
            "This placeholder export is generated during implementation phase testing."
        )

    def _write_report_working_doc(
        self,
        *,
        session_id: str,
        workflow: dict[str, Any],
        stage: str,
        instructions: dict[str, str] | None = None,
        content_map: dict[str, str] | None = None,
    ) -> None:
        working_artifact_id = workflow.get("working_artifact_id")
        if not working_artifact_id:
            return
        artifact = self.artifacts.get_artifact(session_id, working_artifact_id)
        if artifact is None:
            return
        file_path = Path(artifact["storage_uri"])
        from docx import Document

        topic = workflow.get("topic", "Report")
        section_entries = self._list_report_section_entries(workflow)
        instruction_map = instructions or workflow.get("section_instructions") or {}
        generated_map = content_map or {}

        doc = Document()
        doc.add_heading(f"Report: {topic}", level=0)
        for entry in section_entries:
            key = entry["key"]
            title = entry["title"]
            path = entry["path"]
            depth = int(entry["depth"])
            heading_level = max(1, min(depth, 6))
            doc.add_heading(title, level=heading_level)
            section_label = " > ".join(path)
            if stage == "content":
                text = generated_map.get(
                    key,
                    f"Content placeholder for {section_label} based on session context.",
                )
                doc.add_paragraph(text)
            elif stage == "instructions":
                instruction_text = instruction_map.get(
                    key,
                    f"Provide focused analysis for {section_label}.",
                )
                doc.add_paragraph(f"Instruction: {instruction_text}")
        doc.save(str(file_path))
        self.artifacts.update_artifact(
            session_id=session_id,
            artifact_id=working_artifact_id,
            metadata={
                "report_topic": topic,
                "report_state": stage,
                "updated_at": datetime.now(UTC).isoformat(),
            },
        )

    def _generate_report_content_map(self, workflow: dict[str, Any]) -> dict[str, str]:
        topic = workflow.get("topic", "Report")
        section_entries = self._list_report_section_entries(workflow)
        instructions: dict[str, str] = workflow.get("section_instructions") or {}
        section_assignments: dict[str, dict[str, Any]] = workflow.get("section_research_assignments") or {}
        dependency_map = self._build_report_section_dependency_map(
            workflow=workflow,
            section_entries=section_entries,
            section_assignments=section_assignments,
        )
        content: dict[str, str] = {}
        for entry in section_entries:
            key = entry["key"]
            title = entry["title"]
            section_label = " > ".join(entry["path"])
            instruction_text = instructions.get(key, "Use available sources and uploaded documents.")
            content[key] = (
                f"{title}: Generated analysis for '{topic}'. "
                f"Applied instruction: {instruction_text}"
            )
            if entry["depth"] == 1 and key in instructions:
                content[key] = (
                    f"{title}: This section was generated for '{topic}' "
                    f"following instruction: {instructions[key]}."
                )
            assignment = section_assignments.get(key) or {}
            assignment_text = assignment.get("reference_text")
            if assignment_text:
                content[key] = f"{content[key]}\n\nResearch reference: {assignment_text}"
            elif entry["depth"] >= 3:
                content[key] = f"{content[key]}\n\nSection path: {section_label}"
            dependencies = dependency_map.get(key) or []
            if dependencies:
                dep_text = "; ".join(dependencies[:4])
                content[key] = f"{content[key]}\n\nCross-section dependencies: {dep_text}"
        return content

    def _build_report_section_dependency_map(
        self,
        *,
        workflow: dict[str, Any],
        section_entries: list[dict[str, Any]],
        section_assignments: dict[str, dict[str, Any]],
    ) -> dict[str, list[str]]:
        by_key = {entry["key"]: entry for entry in section_entries}
        children_by_parent: dict[str, list[str]] = {}
        for entry in section_entries:
            parent_key = str(entry.get("parent_key") or "")
            children_by_parent.setdefault(parent_key, []).append(entry["key"])
        dependency_map: dict[str, list[str]] = {}
        for entry in section_entries:
            key = entry["key"]
            path = list(entry.get("path") or [])
            parent_key = str(entry.get("parent_key") or "")
            references: list[str] = []
            seen_refs: set[str] = set()

            parent_entry = by_key.get(parent_key)
            if parent_entry is not None:
                label = " > ".join(parent_entry.get("path") or [])
                if label and label not in seen_refs:
                    references.append(f"builds on parent section '{label}'")
                    seen_refs.add(label)

            sibling_keys = [item for item in children_by_parent.get(parent_key, []) if item != key]
            for sibling_key in sibling_keys[:2]:
                sibling_entry = by_key.get(sibling_key)
                if sibling_entry is None:
                    continue
                label = " > ".join(sibling_entry.get("path") or [])
                if label and label not in seen_refs:
                    references.append(f"align with sibling section '{label}'")
                    seen_refs.add(label)

            child_keys = children_by_parent.get(key, [])
            for child_key in child_keys[:2]:
                child_entry = by_key.get(child_key)
                if child_entry is None:
                    continue
                label = " > ".join(child_entry.get("path") or [])
                if label and label not in seen_refs:
                    references.append(f"summarize implications for child section '{label}'")
                    seen_refs.add(label)

            primary_path = path[0] if path else ""
            if primary_path and primary_path not in seen_refs and len(path) >= 2:
                references.append(f"tie findings back to primary section '{primary_path}'")
                seen_refs.add(primary_path)

            assignment = section_assignments.get(key) or {}
            source_id = str(assignment.get("source_id") or "").strip()
            if source_id and source_id not in seen_refs:
                references.append(f"maintain consistency with source evidence '{source_id}'")
                seen_refs.add(source_id)

            dependency_map[key] = references
        return dependency_map

    def _run_generate_report_document_tool(
        self,
        *,
        session_id: str,
        assistant_message_id: str,
        workflow: dict[str, Any],
        available_sources: list[dict[str, Any]],
        preferred_source_ids: list[str],
    ) -> dict[str, Any]:
        tool_name = "generate_report_document"
        call_id = str(uuid4())
        topic = str(workflow.get("topic") or "Report")
        section_keys = self._list_report_section_keys(workflow)
        section_count = len(section_keys)
        instruction_count = len(workflow.get("section_instructions") or {})

        self.events.create_event(
            session_id=session_id,
            message_id=assistant_message_id,
            event_type="tool_call_request",
            payload={
                "call_id": call_id,
                "tool_name": tool_name,
                "arguments": {
                    "topic": topic,
                    "section_count": section_count,
                    "instruction_count": instruction_count,
                    "available_source_count": len(available_sources),
                },
            },
        )

        try:
            research_items = self._list_report_research_context(session_id=session_id, limit=60)
            assignments = self._assign_research_context_to_sections(
                workflow=workflow,
                section_keys=section_keys,
                research_items=research_items,
            )
            source_pool = self._build_report_gap_fill_source_pool(
                available_sources=available_sources,
                preferred_source_ids=preferred_source_ids,
            )
            gap_fill_call_count = 0
            gap_fill_iteration_count = 0
            gap_fill_artifacts: list[str] = []
            source_map = {str(item.get("source_id") or ""): item for item in source_pool}
            min_section_quality_score = 2
            low_quality_targets = self._select_report_gap_fill_section_targets(
                section_keys=section_keys,
                assignments=assignments,
                min_quality_score=min_section_quality_score,
                max_targets=len(section_keys) or 1,
            )
            max_gap_fill_calls = min(10, max(3, len(low_quality_targets) * 2, len(section_keys) // 2))
            executed_signatures: set[str] = set()
            planner_invalid_iterations = 0
            section_source_attempts: dict[tuple[str, str], int] = {}
            low_quality_stable_iterations = 0
            previous_low_quality_count = len(low_quality_targets)
            gap_fill_stop_reason = ""
            gap_fill_mode = "llm_iterative"
            while gap_fill_call_count < max_gap_fill_calls and source_pool:
                remaining_budget = max_gap_fill_calls - gap_fill_call_count
                if remaining_budget <= 0:
                    gap_fill_stop_reason = "budget_exhausted"
                    break
                low_quality_targets = self._select_report_gap_fill_section_targets(
                    section_keys=section_keys,
                    assignments=assignments,
                    min_quality_score=min_section_quality_score,
                    max_targets=len(section_keys) or 1,
                )
                if not low_quality_targets:
                    gap_fill_stop_reason = "coverage_sufficient"
                    break
                planned = self._plan_report_gap_fill_calls(
                    workflow=workflow,
                    section_keys=section_keys,
                    assignments=assignments,
                    research_items=research_items,
                    source_pool=source_pool,
                    executed_signatures=executed_signatures,
                    low_quality_section_keys=low_quality_targets,
                    max_calls=min(3, remaining_budget),
                )
                planned_calls: list[dict[str, str]] = []
                if planned is not None:
                    action = str(planned.get("action") or "").strip().lower()
                    if action == "finish" and not planned.get("planned_calls"):
                        gap_fill_stop_reason = "planner_finish"
                        break
                    planned_calls = list(planned.get("planned_calls") or [])
                if not planned_calls:
                    planner_invalid_iterations += 1
                    if planner_invalid_iterations > 1 and gap_fill_call_count > 0:
                        gap_fill_stop_reason = "planner_invalid"
                        break
                    gap_fill_mode = "llm_with_heuristic_fallback"
                    gap_targets = low_quality_targets[: min(2, remaining_budget)]
                    for section_key in gap_targets:
                        section_query = self._build_report_section_gap_fill_query(
                            workflow=workflow,
                            section_key=section_key,
                        )
                        chosen_source_ids = self._select_sources_for_turn(
                            effective_sources=source_pool,
                            query_text=section_query,
                            max_sources=1,
                        )
                        if not chosen_source_ids:
                            continue
                        planned_calls.append(
                            {
                                "section_key": section_key,
                                "source_id": chosen_source_ids[0],
                                "research_statement": section_query,
                            }
                        )
                if not planned_calls:
                    gap_fill_stop_reason = "no_viable_calls"
                    break

                gap_fill_iteration_count += 1
                improved_this_iteration = False
                executed_this_iteration = 0
                for item in planned_calls:
                    if gap_fill_call_count >= max_gap_fill_calls:
                        break
                    section_key = str(item.get("section_key") or "")
                    source_id = str(item.get("source_id") or "")
                    default_query = self._build_report_section_gap_fill_query(
                        workflow=workflow,
                        section_key=section_key,
                    )
                    research_statement = self._normalize_research_statement(
                        str(item.get("research_statement") or "").strip(),
                        fallback=default_query,
                    )
                    if not section_key or not source_id or not research_statement:
                        continue
                    attempt_key = (section_key, source_id)
                    if section_source_attempts.get(attempt_key, 0) >= 2:
                        continue
                    signature = f"{section_key}::{source_id}::{research_statement.lower()}"
                    if signature in executed_signatures:
                        continue
                    source = source_map.get(source_id)
                    if source is None:
                        continue
                    executed_signatures.add(signature)
                    section_source_attempts[attempt_key] = section_source_attempts.get(attempt_key, 0) + 1
                    db_result = self._run_database_research_tool(
                        session_id=session_id,
                        assistant_message_id=assistant_message_id,
                        source=source,
                        research_statement=research_statement,
                    )
                    gap_fill_call_count += 1
                    executed_this_iteration += 1
                    gap_fill_artifacts.extend(db_result.get("artifacts_created") or [])
                    summary_text = str(db_result.get("summary_text") or "")
                    research_markdown = db_result.get("research_markdown") if isinstance(db_result, dict) else None
                    research_document = db_result.get("research_document") if isinstance(db_result, dict) else None
                    assigned_artifact_id = ""
                    if isinstance(research_markdown, dict):
                        assigned_artifact_id = str(research_markdown.get("artifact_id") or "")
                    if not assigned_artifact_id and isinstance(research_document, dict):
                        assigned_artifact_id = str(research_document.get("artifact_id") or "")
                    research_item = {
                        "artifact_id": assigned_artifact_id,
                        "filename": str(
                            (research_markdown or {}).get("filename")
                            or (research_document or {}).get("filename")
                            or ""
                        ),
                        "source_id": source_id,
                        "summary_text": summary_text,
                    }
                    research_items.insert(0, research_item)
                    if len(research_items) > 120:
                        research_items = research_items[:120]
                    new_quality_score = self._score_research_item_for_section(
                        workflow=workflow,
                        section_key=section_key,
                        research_item=research_item,
                    )
                    previous_assignment = assignments.get(section_key) or {}
                    previous_quality_score = int(previous_assignment.get("quality_score") or 0)
                    if new_quality_score >= previous_quality_score:
                        reference_text = (
                            f"artifact={assigned_artifact_id}, source={source_id}, quality={new_quality_score}"
                        )
                        summary_clip = self._clip_text(summary_text, max_chars=260)
                        if summary_clip:
                            reference_text = f"{reference_text}, summary={summary_clip}"
                        assignments[section_key] = {
                            "artifact_id": assigned_artifact_id,
                            "source_id": source_id,
                            "reference_text": reference_text[:420],
                            "quality_score": new_quality_score,
                        }
                    if new_quality_score > previous_quality_score:
                        improved_this_iteration = True
                if executed_this_iteration == 0:
                    gap_fill_stop_reason = "no_executed_calls"
                    break
                current_low_quality_count = len(
                    self._select_report_gap_fill_section_targets(
                        section_keys=section_keys,
                        assignments=assignments,
                        min_quality_score=min_section_quality_score,
                        max_targets=len(section_keys) or 1,
                    )
                )
                if current_low_quality_count < previous_low_quality_count:
                    low_quality_stable_iterations = 0
                elif not improved_this_iteration:
                    low_quality_stable_iterations += 1
                else:
                    low_quality_stable_iterations = 0
                previous_low_quality_count = current_low_quality_count
                if low_quality_stable_iterations >= 2:
                    gap_fill_stop_reason = "quality_stalled"
                    break

            low_quality_section_count = len(
                self._select_report_gap_fill_section_targets(
                    section_keys=section_keys,
                    assignments=assignments,
                    min_quality_score=min_section_quality_score,
                    max_targets=len(section_keys) or 1,
                )
            )
            workflow["section_research_assignments"] = assignments
            generated_content = self._generate_report_content_map(workflow)
            assignment_summary = [
                {
                    "section_key": key,
                    "artifact_id": value.get("artifact_id"),
                    "source_id": value.get("source_id"),
                    "quality_score": int(value.get("quality_score") or 0),
                }
                for key, value in assignments.items()
            ]
            result = {
                "status": "completed",
                "topic": topic,
                "section_count": section_count,
                "instruction_count": instruction_count,
                "research_artifact_count": len(research_items),
                "assigned_research_count": len(assignments),
                "gap_fill_call_count": gap_fill_call_count,
                "gap_fill_iteration_count": gap_fill_iteration_count,
                "gap_fill_mode": gap_fill_mode,
                "gap_fill_stop_reason": gap_fill_stop_reason or "completed",
                "low_quality_section_count": low_quality_section_count,
                "gap_fill_artifact_count": len(self._merge_artifact_ids(gap_fill_artifacts)),
                "section_assignments": assignment_summary[:80],
            }
            self.events.create_event(
                session_id=session_id,
                message_id=assistant_message_id,
                event_type="tool_call_response",
                payload={"call_id": call_id, "tool_name": tool_name, "result": result},
            )
            return {
                "status": "completed",
                "generated_content": generated_content,
                "research_artifact_count": len(research_items),
                "assigned_research_count": len(assignments),
                "gap_fill_call_count": gap_fill_call_count,
                "gap_fill_iteration_count": gap_fill_iteration_count,
                "gap_fill_mode": gap_fill_mode,
                "gap_fill_stop_reason": gap_fill_stop_reason or "completed",
                "low_quality_section_count": low_quality_section_count,
                "gap_fill_artifacts_created": self._merge_artifact_ids(gap_fill_artifacts),
            }
        except Exception as exc:  # noqa: BLE001
            result = {
                "status": "failed",
                "error": str(exc),
                "topic": topic,
            }
            self.events.create_event(
                session_id=session_id,
                message_id=assistant_message_id,
                event_type="tool_call_response",
                payload={"call_id": call_id, "tool_name": tool_name, "result": result},
            )
            return result

    def _list_report_section_keys(self, workflow: dict[str, Any]) -> list[str]:
        return [entry["key"] for entry in self._list_report_section_entries(workflow)]

    def _list_report_section_entries(self, workflow: dict[str, Any]) -> list[dict[str, Any]]:
        hierarchy = self._get_report_structure_hierarchy(workflow)
        entries: list[dict[str, Any]] = []

        def walk(parent_key: str, parent_path: list[str]) -> None:
            for title in hierarchy.get(parent_key, []):
                path = [*parent_path, title]
                key = self._section_key_from_path(path)
                entries.append(
                    {
                        "key": key,
                        "path": path,
                        "title": title,
                        "depth": len(path),
                        "parent_key": parent_key,
                    }
                )
                walk(key, path)

        walk("", [])
        return entries

    def _get_report_structure_hierarchy(self, workflow: dict[str, Any]) -> dict[str, list[str]]:
        raw = workflow.get("section_hierarchy")
        hierarchy: dict[str, list[str]] = {}
        if isinstance(raw, dict):
            for parent_key, children in raw.items():
                key = str(parent_key or "")
                if not isinstance(children, list):
                    continue
                hierarchy[key] = self._normalize_section_titles(children)

        if "" not in hierarchy or not hierarchy[""]:
            primary_sections: list[str] = workflow.get("primary_sections") or []
            hierarchy[""] = self._normalize_section_titles(primary_sections)
            subsections: dict[str, list[str]] = workflow.get("subsections") or {}
            for primary in hierarchy[""]:
                parent_key = self._section_key_from_path([primary])
                children = subsections.get(primary, [])
                hierarchy[parent_key] = self._normalize_section_titles(children)

        seen_keys = {""}
        queue = [""]
        while queue:
            parent_key = queue.pop(0)
            for child in hierarchy.get(parent_key, []):
                child_key = self._section_key_from_path([*self._section_path_from_key(parent_key), child])
                if child_key in seen_keys:
                    continue
                seen_keys.add(child_key)
                hierarchy.setdefault(child_key, [])
                queue.append(child_key)
        return hierarchy

    def _sync_legacy_structure_fields(self, workflow: dict[str, Any]) -> None:
        hierarchy = self._get_report_structure_hierarchy(workflow)
        workflow["section_hierarchy"] = hierarchy
        primary_sections = list(hierarchy.get("", []))
        workflow["primary_sections"] = primary_sections
        workflow["subsections"] = {
            primary: list(hierarchy.get(self._section_key_from_path([primary]), []))
            for primary in primary_sections
        }
        max_depth = 1
        for key in hierarchy:
            if not key:
                continue
            max_depth = max(max_depth, len(self._section_path_from_key(key)))
        workflow["structure_max_depth"] = max_depth

    def _expand_report_hierarchy_depth(
        self,
        *,
        hierarchy: dict[str, list[str]],
        topic: str,
        max_depth: int,
        max_parents: int,
    ) -> dict[str, list[str]]:
        if max_depth <= 1 or max_parents <= 0:
            return {}
        candidates: list[str] = []
        for key, children in hierarchy.items():
            if not key:
                continue
            depth = len(self._section_path_from_key(key))
            if depth >= max_depth:
                continue
            if children:
                continue
            candidates.append(key)
        candidates.sort(key=lambda item: (len(self._section_path_from_key(item)), item))
        added: dict[str, list[str]] = {}
        for parent_key in candidates[:max_parents]:
            parent_path = self._section_path_from_key(parent_key)
            suggestions = self._suggest_subsections_for_path(parent_path, topic)
            if not suggestions:
                continue
            hierarchy[parent_key] = suggestions
            for child in suggestions:
                child_key = self._section_key_from_path([*parent_path, child])
                hierarchy.setdefault(child_key, [])
            added[parent_key] = suggestions
        return added

    @staticmethod
    def _build_report_gap_fill_source_pool(
        *,
        available_sources: list[dict[str, Any]],
        preferred_source_ids: list[str],
    ) -> list[dict[str, Any]]:
        if not available_sources:
            return []
        by_id = {str(item.get("source_id") or ""): item for item in available_sources}
        ordered: list[dict[str, Any]] = []
        seen: set[str] = set()
        for source_id in preferred_source_ids:
            key = str(source_id or "")
            source = by_id.get(key)
            if source is None or key in seen:
                continue
            seen.add(key)
            ordered.append(source)
        for source in available_sources:
            key = str(source.get("source_id") or "")
            if not key or key in seen:
                continue
            seen.add(key)
            ordered.append(source)
        return ordered

    def _select_report_gap_fill_section_targets(
        self,
        *,
        section_keys: list[str],
        assignments: dict[str, dict[str, Any]],
        min_quality_score: int,
        max_targets: int,
    ) -> list[str]:
        if max_targets <= 0 or not section_keys:
            return []
        ranked: list[tuple[int, int, str]] = []
        for index, key in enumerate(section_keys):
            assignment = assignments.get(key) or {}
            quality_score = int(assignment.get("quality_score") or 0)
            if str(assignment.get("reference_text") or "").strip() == "":
                quality_score = 0
            if quality_score >= min_quality_score:
                continue
            ranked.append((quality_score, index, key))
        ranked.sort(key=lambda item: (item[0], item[1]))
        return [key for _, _, key in ranked[:max_targets]]

    def _build_report_section_gap_fill_query(self, *, workflow: dict[str, Any], section_key: str) -> str:
        topic = str(workflow.get("topic") or "Report")
        instructions = workflow.get("section_instructions") if isinstance(workflow, dict) else {}
        if not isinstance(instructions, dict):
            instructions = {}
        instruction_text = str(instructions.get(section_key) or "").strip()
        primary, subsection = self._split_section_key(section_key)
        section_label = primary if not subsection else f"{primary} - {subsection}"
        query = f"Research evidence for report topic '{topic}', section '{section_label}'."
        if instruction_text:
            query += f" Instruction focus: {instruction_text}"
        return query[:1400]

    def _score_research_item_for_section(
        self,
        *,
        workflow: dict[str, Any],
        section_key: str,
        research_item: dict[str, str],
    ) -> int:
        instructions = workflow.get("section_instructions") if isinstance(workflow, dict) else {}
        if not isinstance(instructions, dict):
            instructions = {}
        primary, subsection = self._split_section_key(section_key)
        section_terms = self._keyword_terms(
            " ".join(
                [
                    str(workflow.get("topic") or ""),
                    primary,
                    subsection,
                    str(instructions.get(section_key) or ""),
                ]
            )
        )
        if not section_terms:
            return 0
        haystack = " ".join(
            [
                str(research_item.get("source_id") or ""),
                str(research_item.get("filename") or ""),
                str(research_item.get("summary_text") or ""),
            ]
        ).lower()
        score = self._match_terms_count(haystack, section_terms)
        primary_token = primary.strip().lower()
        subsection_token = subsection.strip().lower()
        if primary_token and primary_token in haystack:
            score += 1
        if subsection_token and subsection_token in haystack:
            score += 1
        return max(0, min(score, 12))

    def _plan_report_gap_fill_calls(
        self,
        *,
        workflow: dict[str, Any],
        section_keys: list[str],
        assignments: dict[str, dict[str, Any]],
        research_items: list[dict[str, str]],
        source_pool: list[dict[str, Any]],
        executed_signatures: set[str],
        low_quality_section_keys: list[str],
        max_calls: int,
    ) -> dict[str, Any] | None:
        if max_calls <= 0 or not section_keys or not source_pool:
            return None

        instructions = workflow.get("section_instructions") if isinstance(workflow, dict) else {}
        if not isinstance(instructions, dict):
            instructions = {}
        topic = str(workflow.get("topic") or "Report")
        lines = [
            "REPORT_GAP_FILL_PLAN_JSON",
            "Plan additional database_research calls to fill report coverage gaps.",
            f"Topic: {topic}",
            f"Max calls this step: {max_calls}",
            "Sections and current coverage:",
        ]
        for key in section_keys[:60]:
            path = self._section_path_from_key(key)
            assignment = assignments.get(key) or {}
            lines.append(
                "- "
                + json.dumps(
                    {
                        "section_key": key,
                        "section_path": path,
                        "instruction": str(instructions.get(key) or "")[:260],
                        "quality_score": int(assignment.get("quality_score") or 0),
                        "current_reference": str(assignment.get("reference_text") or "")[:260],
                    }
                )
            )
        lines.append("Current research context:")
        if research_items:
            for item in research_items[:20]:
                lines.append(
                    "- "
                    + json.dumps(
                        {
                            "artifact_id": item.get("artifact_id"),
                            "source_id": item.get("source_id"),
                            "summary_text": str(item.get("summary_text") or "")[:260],
                        }
                    )
                )
        else:
            lines.append("- no_research_items")
        lines.append("Available sources:")
        for source in source_pool:
            lines.append(
                "- "
                + json.dumps(
                    {
                        "source_id": source.get("source_id"),
                        "name": source.get("name"),
                        "source_type": source.get("source_type"),
                    }
                )
            )
        lines.append("Already executed signatures (section::source::statement):")
        lines.append(", ".join(sorted(executed_signatures)) if executed_signatures else "none")
        lines.append("Prioritize these low-quality section keys first:")
        lines.append(", ".join(low_quality_section_keys[:20]) if low_quality_section_keys else "none")
        lines.append("Return JSON only with shape:")
        lines.append(
            json.dumps(
                {
                    "action": "call_more_or_finish",
                    "calls": [
                        {
                            "section_key": "section::path",
                            "source_id": "source_id",
                            "research_statement": "targeted research question",
                        }
                    ],
                    "reasoning": "short rationale",
                }
            )
        )
        lines.append(
            "Rules: use only listed section_key and source_id values; "
            "avoid duplicate signatures; prioritize low-quality sections; "
            "use concrete research_statement values tied to section instructions; "
            "choose finish when coverage is sufficient."
        )
        prompt = "\n".join(lines)
        response = self.agent.chat(
            [{"role": "user", "content": prompt}],
            system_context=(
                "REPORT_GAP_FILL_PLANNER_MODE: Return strict JSON only. "
                "Do not include markdown fences."
            ),
        )
        payload = self._extract_json_object_from_text(response)
        if not isinstance(payload, dict):
            return None
        action = str(payload.get("action") or "").strip().lower()
        calls = self._normalize_report_gap_fill_calls(
            calls=list(payload.get("calls") or payload.get("section_queries") or []),
            allowed_section_keys=section_keys,
            allowed_source_ids=[str(item.get("source_id") or "") for item in source_pool],
            executed_signatures=executed_signatures,
            workflow=workflow,
        )
        calls = calls[:max_calls]
        reasoning = str(payload.get("reasoning") or "").strip()
        if not reasoning:
            reasoning = "Gap-fill planner decision generated."
        if action in {"finish", "done", "none", "stop"} and not calls:
            return {"action": "finish", "planned_calls": [], "reasoning": reasoning[:600]}
        if not action:
            action = "call_more" if calls else "finish"
        if action not in {"call_more", "call", "database_research", "research", "finish"}:
            if not calls:
                return None
            action = "call_more"
        if action == "finish" and calls:
            action = "call_more"
        if action in {"call", "database_research", "research"}:
            action = "call_more"
        if action == "call_more" and not calls:
            return None
        return {"action": action, "planned_calls": calls, "reasoning": reasoning[:600]}

    def _normalize_report_gap_fill_calls(
        self,
        *,
        calls: list[dict[str, Any]],
        allowed_section_keys: list[str],
        allowed_source_ids: list[str],
        executed_signatures: set[str],
        workflow: dict[str, Any],
    ) -> list[dict[str, str]]:
        normalized: list[dict[str, str]] = []
        seen: set[str] = set()
        allowed_sections = set(allowed_section_keys)
        allowed_sources = set(allowed_source_ids)
        for item in calls:
            if not isinstance(item, dict):
                continue
            section_candidate = str(
                item.get("section_key")
                or item.get("section")
                or item.get("section_path")
                or ""
            ).strip()
            section_key = ""
            if section_candidate in allowed_sections:
                section_key = section_candidate
            else:
                section_key = str(
                    self._match_section_target(section_candidate, section_keys=allowed_section_keys) or ""
                )
            if not section_key:
                continue
            source_id = str(item.get("source_id") or "").strip()
            if not source_id or source_id not in allowed_sources:
                continue
            statement_raw = str(item.get("research_statement") or "").strip()
            default_statement = self._build_report_section_gap_fill_query(
                workflow=workflow,
                section_key=section_key,
            )
            statement = self._normalize_research_statement(
                statement_raw,
                fallback=default_statement,
            )
            signature = f"{section_key}::{source_id}::{statement.lower()}"
            if signature in executed_signatures or signature in seen:
                continue
            seen.add(signature)
            normalized.append(
                {
                    "section_key": section_key,
                    "source_id": source_id,
                    "research_statement": statement[:1400],
                }
            )
        return normalized

    @staticmethod
    def _split_section_key(section_key: str) -> tuple[str, str]:
        path = WordUIWorkspaceV2._section_path_from_key(section_key)
        if not path:
            return section_key.strip(), ""
        if len(path) == 1:
            return path[0], ""
        return path[0], " > ".join(path[1:])

    def _list_report_research_context(
        self,
        *,
        session_id: str,
        limit: int = 40,
    ) -> list[dict[str, str]]:
        rows = self.artifacts.list_artifacts(session_id, artifact_type="research_markdown", limit=500)
        out: list[dict[str, str]] = []
        for row in rows:
            metadata = row.get("metadata") if isinstance(row, dict) else {}
            if not isinstance(metadata, dict):
                metadata = {}
            source_id = str(metadata.get("source_id") or "")
            summary_text = str(metadata.get("summary_text") or "")
            out.append(
                {
                    "artifact_id": str(row.get("artifact_id") or ""),
                    "filename": str(row.get("filename") or ""),
                    "source_id": source_id,
                    "summary_text": summary_text,
                }
            )
            if len(out) >= limit:
                break
        return out

    def _assign_research_context_to_sections(
        self,
        *,
        workflow: dict[str, Any],
        section_keys: list[str],
        research_items: list[dict[str, str]],
    ) -> dict[str, dict[str, Any]]:
        if not section_keys or not research_items:
            return {}
        assignments: dict[str, dict[str, Any]] = {}
        used_artifact_ids: set[str] = set()
        for key in section_keys:
            best_item: dict[str, str] | None = None
            best_score = -1
            best_unused = -1
            for item in research_items:
                score = self._score_research_item_for_section(
                    workflow=workflow,
                    section_key=key,
                    research_item=item,
                )
                artifact_id = str(item.get("artifact_id") or "")
                unused_flag = 1 if artifact_id and artifact_id not in used_artifact_ids else 0
                if score > best_score or (score == best_score and unused_flag > best_unused):
                    best_item = item
                    best_score = score
                    best_unused = unused_flag
            if best_item is None:
                continue
            source_id = str(best_item.get("source_id") or "unknown_source")
            artifact_id = str(best_item.get("artifact_id") or "")
            if artifact_id:
                used_artifact_ids.add(artifact_id)
            snippet = self._clip_text(str(best_item.get("summary_text") or ""), max_chars=260)
            reference = f"artifact={artifact_id}, source={source_id}, quality={max(best_score, 0)}"
            if snippet:
                reference = f"{reference}, summary={snippet}"
            assignments[key] = {
                "artifact_id": artifact_id,
                "source_id": source_id,
                "reference_text": reference[:420],
                "quality_score": max(best_score, 0),
            }
        return assignments

    def _resolve_section_instructions(
        self,
        workflow: dict[str, Any],
        message_text: str,
    ) -> tuple[dict[str, str], dict[str, Any]]:
        topic = workflow.get("topic", "Report")
        section_entries = self._list_report_section_entries(workflow)
        text = (message_text or "").strip()
        lower = text.lower()
        instructions: dict[str, str] = {}
        all_keys = [entry["key"] for entry in section_entries]

        if "use defaults" in lower or lower in ("default", "defaults", "no", "no instructions"):
            default_global = f"Provide clear, evidence-based analysis for {topic}."
            for entry in section_entries:
                key = entry["key"]
                title = entry["title"]
                section_label = " > ".join(entry["path"])
                if entry["depth"] == 1:
                    instructions[key] = default_global
                else:
                    instructions[key] = (
                        f"Summarize {section_label} for {topic} with clear findings, evidence, and actions."
                    )
            return instructions, {
                "mode": "defaults",
                "provided_keys": [],
                "defaulted_keys": sorted(all_keys),
            }

        explicit_instruction_map, global_instruction = self._parse_instruction_overrides(
            text=text,
            section_keys=all_keys,
        )

        mode = "structured"
        if explicit_instruction_map:
            instructions.update(explicit_instruction_map)
            fallback_instruction = global_instruction or f"Provide clear, evidence-based analysis for {topic}."
            for key in all_keys:
                if key not in instructions:
                    instructions[key] = fallback_instruction
        else:
            mode = "global"
            fallback_instruction = text or f"Provide clear, evidence-based analysis for {topic}."
            for key in all_keys:
                instructions[key] = fallback_instruction

        provided_keys = sorted(explicit_instruction_map.keys())
        defaulted_keys = sorted(key for key in all_keys if key not in set(explicit_instruction_map.keys()))
        return instructions, {
            "mode": mode,
            "provided_keys": provided_keys,
            "defaulted_keys": defaulted_keys,
        }

    @staticmethod
    def _section_key(primary: str, subsection: str | None) -> str:
        path = [primary]
        if subsection:
            path.extend(WordUIWorkspaceV2._section_path_from_key(subsection))
        return WordUIWorkspaceV2._section_key_from_path(path)

    @staticmethod
    def _section_key_from_path(path: list[str]) -> str:
        cleaned = [str(item).strip() for item in path if str(item).strip()]
        return "::".join(cleaned)

    @staticmethod
    def _section_path_from_key(section_key: str) -> list[str]:
        if not section_key:
            return []
        return [part.strip() for part in str(section_key).split("::") if part.strip()]

    @staticmethod
    def _normalize_section_titles(items: list[str]) -> list[str]:
        out: list[str] = []
        seen: set[str] = set()
        for raw in items:
            value = str(raw or "").strip()
            if len(value) < 2:
                continue
            norm = WordUIWorkspaceV2._normalize_label(value)
            if not norm or norm in seen:
                continue
            seen.add(norm)
            out.append(value)
        return out[:20]

    def _report_output_dir(self, session_id: str) -> Path:
        root = self.allowed_roots[0] if self.allowed_roots else Path.cwd()
        return root / ".docx-agent-v2-reports" / session_id

    def _research_output_dir(self, session_id: str) -> Path:
        root = self.allowed_roots[0] if self.allowed_roots else Path.cwd()
        return root / ".docx-agent-v2-research" / session_id

    @staticmethod
    def _safe_slug(text: str) -> str:
        raw = "".join(ch.lower() if ch.isalnum() else "-" for ch in text.strip())
        compact = "-".join(part for part in raw.split("-") if part)
        return compact[:60] or "report"

    # Deprecated keyword methods — used only by _interpret_report_turn_fallback().
    # Primary classification is now handled by _interpret_report_turn() via LLM.

    @staticmethod
    def _is_report_request(text: str) -> bool:
        """Deprecated: keyword fallback for report intent detection."""
        lowered = text.lower()
        if "report" not in lowered:
            return False
        trigger_terms = ("create", "build", "generate", "make", "draft", "start", "write", "prepare")
        return any(term in lowered for term in trigger_terms)

    @staticmethod
    def _is_deeper_structure_request(text: str) -> bool:
        """Deprecated: keyword fallback for deeper-structure detection."""
        lowered = (text or "").lower()
        triggers = (
            "add deeper",
            "deeper",
            "nested",
            "expand structure",
            "continue depth",
            "add subsection depth",
        )
        return any(token in lowered for token in triggers)

    @staticmethod
    def _is_finalize_structure_request(text: str) -> bool:
        """Deprecated: keyword fallback for finalize-structure detection."""
        lowered = (text or "").lower().strip()
        exact = {"final structure", "finalize structure", "structure final", "done", "proceed"}
        if lowered in exact:
            return True
        contains = ("final structure", "structure is final", "finalize")
        return any(token in lowered for token in contains)

    @staticmethod
    def _extract_report_topic(text: str) -> str:
        """Deprecated: keyword fallback for topic extraction."""
        lowered = text.lower()
        markers = ["based on", "about", "for", "on"]
        for marker in markers:
            token = f"{marker} "
            if token in lowered:
                index = lowered.find(token)
                topic = text[index + len(token):].strip(" .,:;")
                if topic:
                    return topic[:180]
        return text.strip(" .,:;")[:180] or "Untitled Report"

    @staticmethod
    def _suggest_primary_sections(topic: str) -> list[str]:
        """Deprecated: hardcoded fallback section suggestions."""
        return [
            "Executive Summary",
            "Background and Scope",
            "Key Findings",
            "Risk and Opportunity Analysis",
            "Recommendations",
            "Conclusion",
        ]

    @staticmethod
    def _suggest_subsections(primary: str, topic: str) -> list[str]:
        """Deprecated: hardcoded fallback subsection suggestions."""
        mapping = {
            "Executive Summary": ["Purpose", "Top Insights", "Immediate Actions"],
            "Background and Scope": ["Context", "Objectives", "Method and Sources"],
            "Key Findings": ["Observed Trends", "Comparative Benchmarks", "Material Drivers"],
            "Risk and Opportunity Analysis": ["Risk Areas", "Opportunity Areas", "Impact Assessment"],
            "Recommendations": ["Priority Actions", "Implementation Plan", "Monitoring Metrics"],
            "Conclusion": ["Overall Assessment", "Decision Points", "Next Steps"],
        }
        if primary in mapping:
            return mapping[primary]
        return [f"{primary} Analysis", f"{primary} Evidence", f"{primary} Actions"]

    def _suggest_subsections_for_path(self, path: list[str], topic: str) -> list[str]:
        """Deprecated: hardcoded fallback for path-based subsection suggestions."""
        if not path:
            return []
        if len(path) == 1:
            return self._suggest_subsections(path[0], topic)
        label = path[-1]
        candidates = [
            f"{label} Context",
            f"{label} Evidence",
            f"{label} Actions",
        ]
        return self._normalize_section_titles(candidates)

    @staticmethod
    def _parse_list_from_text(text: str) -> list[str]:
        raw = text.replace("\n", ",").replace(";", ",")
        parts = [item.strip(" -") for item in raw.split(",")]
        cleaned = [item for item in parts if len(item) >= 2]
        if len(cleaned) == 1 and cleaned[0].lower() in ("yes", "use suggested", "use defaults", "ok", "okay"):
            return []
        return cleaned[:20]

    def _select_prior_agent_interactions(
        self,
        *,
        session_id: str,
        assistant_message_id: str,
        limit: int,
    ) -> list[dict[str, str]]:
        list_recent = getattr(self.events, "list_recent_events", None)
        if not callable(list_recent):
            return []
        rows = list_recent(
            session_id=session_id,
            limit=max(20, limit * 4),
            exclude_message_id=assistant_message_id,
            event_types=[
                "tool_call_response",
                "report_generation_completed",
                "report_section_instructions_captured",
                "artifact_created",
                "artifact_updated",
                "model_response",
            ],
        )
        interactions: list[dict[str, str]] = []
        for row in rows:
            event_type = str(row.get("event_type") or "")
            payload = row.get("payload") if isinstance(row, dict) else {}
            if not isinstance(payload, dict):
                payload = {}
            summary = self._summarize_prior_event_payload(event_type, payload)
            if not summary:
                continue
            interactions.append(
                {
                    "event_type": event_type,
                    "message_id": str(row.get("message_id") or ""),
                    "summary": summary,
                }
            )
            if len(interactions) >= limit:
                break
        return interactions

    @staticmethod
    def _summarize_prior_event_payload(event_type: str, payload: dict[str, Any]) -> str:
        if event_type == "tool_call_response":
            tool_name = str(payload.get("tool_name") or "tool")
            result = payload.get("result")
            status = "unknown"
            if isinstance(result, dict):
                status = str(result.get("status") or "unknown")
            return f"{tool_name} -> status={status}"
        if event_type == "artifact_created":
            artifact_id = str(payload.get("artifact_id") or "")
            artifact_type = str(payload.get("artifact_type") or "artifact")
            return f"{artifact_type} created ({artifact_id})" if artifact_id else f"{artifact_type} created"
        if event_type == "artifact_updated":
            artifact_id = str(payload.get("artifact_id") or "")
            reason = str(payload.get("reason") or "updated")
            return f"artifact {artifact_id} updated ({reason})" if artifact_id else f"artifact updated ({reason})"
        if event_type == "report_generation_completed":
            final_artifact_id = str(payload.get("final_artifact_id") or "")
            if final_artifact_id:
                return f"report generation completed, final artifact={final_artifact_id}"
            return "report generation completed"
        if event_type == "report_section_instructions_captured":
            mode = str(payload.get("mode") or "unknown")
            instruction_count = payload.get("instruction_count")
            return f"instructions captured mode={mode}, count={instruction_count}"
        if event_type == "model_response":
            text = str(payload.get("text") or "")
            if not text:
                return "model response generated"
            clipped = text[:120].replace("\n", " ").strip()
            return f"model response: {clipped}"
        return ""

    def _parse_structure_overrides(
        self,
        *,
        text: str,
        hierarchy: dict[str, list[str]],
    ) -> dict[str, list[str]]:
        overrides: dict[str, list[str]] = {}
        if not text.strip():
            return overrides
        for chunk in text.split("\n"):
            line = chunk.strip().strip(";")
            if not line or ":" not in line:
                continue
            parent_raw, children_raw = line.split(":", 1)
            parent_key = self._match_section_target(parent_raw.strip(), hierarchy=hierarchy)
            if parent_key is None:
                continue
            sub_items = self._parse_list_from_text(children_raw)
            if sub_items:
                overrides[parent_key] = self._normalize_section_titles(sub_items)
        return overrides

    def _parse_instruction_overrides(
        self,
        *,
        text: str,
        section_keys: list[str],
    ) -> tuple[dict[str, str], str | None]:
        instruction_map: dict[str, str] = {}
        global_instruction: str | None = None
        if not text.strip():
            return instruction_map, global_instruction

        lines = [line.strip() for line in text.replace(";", "\n").split("\n") if line.strip()]
        for line in lines:
            if "->" not in line:
                continue
            lhs_raw, rhs_raw = line.split("->", 1)
            lhs = lhs_raw.strip()
            instruction = rhs_raw.strip()
            if not lhs or not instruction:
                continue

            lhs_lower = lhs.lower()
            if lhs_lower in ("all", "global", "default", "overall"):
                global_instruction = instruction
                continue

            matched_key = self._match_section_target(lhs, section_keys=section_keys)
            if matched_key is None:
                continue
            instruction_map[matched_key] = instruction

        return instruction_map, global_instruction

    def _match_section_target(
        self,
        candidate: str,
        *,
        hierarchy: dict[str, list[str]] | None = None,
        section_keys: list[str] | None = None,
    ) -> str | None:
        keys = section_keys
        if keys is None and hierarchy is not None:
            keys = [key for key in hierarchy.keys() if key]
        if not keys:
            return None
        cand_norm = self._normalize_label(candidate)
        if not cand_norm:
            return None
        best_key: str | None = None
        best_score = -1
        for key in keys:
            path = self._section_path_from_key(key)
            if not path:
                continue
            aliases = [
                " > ".join(path),
                "::".join(path),
                path[0] if len(path) == 1 else f"{path[0]}: {' > '.join(path[1:])}",
                path[-1],
            ]
            alias_scores = [300, 280, 260, 120]
            for alias, base_score in zip(aliases, alias_scores):
                alias_norm = self._normalize_label(alias)
                if not alias_norm:
                    continue
                score = -1
                if cand_norm == alias_norm:
                    score = base_score + len(alias_norm)
                elif cand_norm in alias_norm or alias_norm in cand_norm:
                    score = (base_score // 2) + min(len(cand_norm), len(alias_norm))
                if score > best_score:
                    best_score = score
                    best_key = key
        return best_key if best_score >= 90 else None

    @staticmethod
    def _split_instruction_target(target: str) -> tuple[str, str | None]:
        stripped = target.strip()
        separators = (":", ">", "/")
        for sep in separators:
            if sep in stripped:
                left, right = stripped.split(sep, 1)
                return left.strip(), right.strip() or None
        return stripped, None

    @staticmethod
    def _match_label(candidate: str, options: list[str]) -> str | None:
        if not candidate or not options:
            return None
        cand_norm = WordUIWorkspaceV2._normalize_label(candidate)
        for option in options:
            if WordUIWorkspaceV2._normalize_label(option) == cand_norm:
                return option
        for option in options:
            opt_norm = WordUIWorkspaceV2._normalize_label(option)
            if cand_norm in opt_norm or opt_norm in cand_norm:
                return option
        return None

    @staticmethod
    def _normalize_label(value: str) -> str:
        compact = "".join(ch.lower() if ch.isalnum() else " " for ch in value)
        return " ".join(part for part in compact.split() if part)

    def _resolve_data_source_filters(self, filters: list[str] | None) -> dict[str, Any]:
        requested = [item.strip() for item in (filters or []) if item and item.strip()]
        if not requested:
            effective = self.data_sources.list_sources(enabled_only=True)
            return {
                "error": False,
                "mode": "unfiltered",
                "requested_source_ids": [],
                "effective_source_ids": [item["source_id"] for item in effective],
                "ignored_source_ids": [],
                "effective_sources": effective,
            }

        effective = self.data_sources.list_sources(enabled_only=True, source_ids=requested)
        effective_ids = [item["source_id"] for item in effective]
        ignored = [item for item in requested if item not in set(effective_ids)]
        if not effective:
            return {
                "error": True,
                "error_code": "INVALID_ARGUMENT",
                "error_message": "data_source_filters resolved to zero enabled sources",
            }
        return {
            "error": False,
            "mode": "filtered",
            "requested_source_ids": requested,
            "effective_source_ids": effective_ids,
            "ignored_source_ids": ignored,
            "effective_sources": effective,
        }

    def _build_llm_messages(self, session_id: str) -> list[dict[str, str]]:
        rows = self.messages.list_messages(session_id, limit=120)
        llm_messages: list[dict[str, str]] = []
        for row in rows:
            if row["processing_state"] != "completed":
                continue
            role = row["role"]
            if role not in ("user", "assistant"):
                continue
            text = row.get("content_text")
            if not text:
                continue
            llm_messages.append({"role": role, "content": text})
        recent = llm_messages[-40:]
        budget_chars = 12000
        selected: list[dict[str, str]] = []
        used_chars = 0
        for item in reversed(recent):
            text = item.get("content") or ""
            cost = len(text)
            if selected and used_chars + cost > budget_chars:
                break
            selected.append(item)
            used_chars += cost
            if len(selected) >= 20:
                break
        selected.reverse()
        return selected

    def _build_system_context(
        self,
        *,
        filter_mode: str,
        sources: list[dict[str, Any]],
        selected_source_ids: list[str],
        uploaded_context_units: list[dict[str, Any]],
        prior_agent_interactions: list[dict[str, Any]],
        internal_research: dict[str, Any],
        uploaded_research: dict[str, Any],
        report_plan_context_note: str | None = None,
    ) -> str:
        lines = [
            "You are the orchestration assistant for research/report workflows.",
            f"DATA_SOURCE_FILTER_MODE: {filter_mode}",
            "AVAILABLE_DATA_SOURCES:",
        ]
        for source in sources:
            lines.append(
                f"- {source['source_id']} | {source['name']} | {source['source_type']}"
            )
        lines.append("SELECTED_DATA_SOURCES_FOR_THIS_TURN:")
        if selected_source_ids:
            lines.extend(f"- {source_id}" for source_id in selected_source_ids)
        else:
            lines.append("- none selected")
        lines.append("UPLOADED_DOCUMENT_CONTEXT:")
        if uploaded_context_units:
            for unit in uploaded_context_units:
                unit_type = unit["unit_type"]
                artifact_id = unit["artifact_id"]
                snippet = unit["content"][:500].replace("\n", " ").strip()
                lines.append(f"- [{unit_type}] artifact={artifact_id}: {snippet}")
        else:
            lines.append("- no uploaded-document knowledge units selected")
        lines.append("PREVIOUS_AGENT_INTERACTIONS:")
        if prior_agent_interactions:
            for item in prior_agent_interactions[:8]:
                event_type = item.get("event_type") or "event"
                summary = item.get("summary") or "no summary"
                lines.append(f"- [{event_type}] {summary}")
        else:
            lines.append("- no prior agent interactions selected")
        lines.append("INTERNAL_RESEARCH_TOOL_OUTPUT:")
        lines.append(
            f"- status={internal_research.get('status', 'unknown')} "
            f"selected_sources={internal_research.get('selected_count', 0)} "
            f"findings={internal_research.get('finding_count', 0)}"
        )
        internal_artifacts = internal_research.get("artifacts_created") or []
        if internal_artifacts:
            lines.append(f"- output_artifact_ids: {', '.join(internal_artifacts[:10])}")
        summary_internal = (internal_research.get("summary_text") or "").strip()
        if summary_internal:
            lines.append(f"- summary: {summary_internal[:1200]}")
        lines.append("UPLOADED_DOCUMENT_RESEARCH_TOOL_OUTPUT:")
        lines.append(
            f"- status={uploaded_research.get('status', 'unknown')} "
            f"units={uploaded_research.get('unit_count', 0)} "
            f"artifacts={uploaded_research.get('artifact_count', 0)}"
        )
        uploaded_artifacts = uploaded_research.get("artifacts_created") or []
        if uploaded_artifacts:
            lines.append(f"- output_artifact_ids: {', '.join(uploaded_artifacts[:10])}")
        summary_uploaded = (uploaded_research.get("summary_text") or "").strip()
        if summary_uploaded:
            lines.append(f"- summary: {summary_uploaded[:1200]}")
        if report_plan_context_note:
            lines.append("REPORT_PLAN_CONTEXT_NOTE:")
            lines.append(f"- {report_plan_context_note}")
        return "\n".join(lines)

    def _ingest_artifact_content(self, session_id: str, artifact: dict[str, Any]) -> dict[str, Any]:
        artifact_id = artifact["artifact_id"]
        storage_uri = artifact["storage_uri"]
        path = Path(storage_uri)
        existing_meta = dict(artifact.get("metadata") or {})
        try:
            text = self._extract_text_for_ingestion(path)
            if not text.strip():
                final_meta = {
                    **existing_meta,
                    "ingestion_state": "completed",
                    "ingestion_note": "no extractable text detected",
                    "knowledge_unit_counts": {"summary": 0, "chunk": 0},
                }
                self.artifacts.update_artifact(
                    session_id=session_id,
                    artifact_id=artifact_id,
                    metadata=final_meta,
                )
                return {
                    "ingestion_state": "completed",
                    "summary_units": 0,
                    "chunk_units": 0,
                    "note": "no extractable text detected",
                }

            summary = self._summarize_ingestion_text(text, max_chars=1200)
            chunks = self._chunk_text_for_ingestion(text, chunk_chars=2000, overlap_chars=250, max_chunks=40)

            self.knowledge.create_knowledge_unit(
                session_id=session_id,
                artifact_id=artifact_id,
                unit_type="summary",
                sequence_no=0,
                content=summary,
                metadata={"source_uri": storage_uri},
            )
            for index, chunk in enumerate(chunks):
                self.knowledge.create_knowledge_unit(
                    session_id=session_id,
                    artifact_id=artifact_id,
                    unit_type="chunk",
                    sequence_no=index,
                    content=chunk,
                    metadata={"source_uri": storage_uri},
                )

            final_meta = {
                **existing_meta,
                "ingestion_state": "completed",
                "knowledge_unit_counts": {"summary": 1, "chunk": len(chunks)},
            }
            self.artifacts.update_artifact(
                session_id=session_id,
                artifact_id=artifact_id,
                metadata=final_meta,
            )
            return {
                "ingestion_state": "completed",
                "summary_units": 1,
                "chunk_units": len(chunks),
            }
        except Exception as exc:  # noqa: BLE001
            final_meta = {
                **existing_meta,
                "ingestion_state": "failed",
                "ingestion_error": str(exc),
            }
            self.artifacts.update_artifact(
                session_id=session_id,
                artifact_id=artifact_id,
                metadata=final_meta,
            )
            return {"ingestion_state": "failed", "error": str(exc)}

    def _select_uploaded_context(
        self,
        *,
        session_id: str,
        query_text: str,
        limit: int,
    ) -> list[dict[str, Any]]:
        summaries = self.knowledge.list_knowledge_units(session_id, unit_type="summary", limit=300)
        if not summaries:
            return []
        query_terms = self._keyword_terms(query_text)
        if not query_terms:
            return summaries[:limit]
        scored: list[tuple[int, dict[str, Any]]] = []
        for item in summaries:
            hay = item["content"].lower()
            score = sum(1 for term in query_terms if term in hay)
            scored.append((score, item))
        scored.sort(key=lambda it: it[0], reverse=True)
        selected = [item for score, item in scored if score > 0]
        if not selected:
            selected = [item for _, item in scored]
        return selected[:limit]

    def _select_sources_for_turn(
        self,
        *,
        effective_sources: list[dict[str, Any]],
        query_text: str,
        max_sources: int,
    ) -> list[str]:
        if not effective_sources:
            return []
        query_terms = self._keyword_terms(query_text)
        if not query_terms:
            return [item["source_id"] for item in effective_sources[:max_sources]]

        scored: list[tuple[int, str]] = []
        for source in effective_sources:
            source_id = source["source_id"]
            hay = f"{source_id} {source['name']} {source['source_type']}".lower()
            score = sum(1 for term in query_terms if term in hay)
            scored.append((score, source_id))
        scored.sort(key=lambda pair: pair[0], reverse=True)
        chosen = [source_id for score, source_id in scored if score > 0][:max_sources]
        if chosen:
            return chosen
        return [item["source_id"] for item in effective_sources[:max_sources]]

    @staticmethod
    def _extract_json_object_from_text(raw_text: str) -> dict[str, Any] | None:
        text = (raw_text or "").strip()
        if not text:
            return None
        candidates = [text]
        first = text.find("{")
        last = text.rfind("}")
        if first >= 0 and last > first:
            candidates.append(text[first : last + 1])
        for candidate in candidates:
            try:
                parsed = json.loads(candidate)
            except json.JSONDecodeError:
                continue
            if isinstance(parsed, dict):
                return parsed
        return None

    @staticmethod
    def _normalize_internal_research_calls(
        *,
        calls: list[dict[str, Any]],
        allowed_source_ids: list[str],
        default_statement: str,
    ) -> list[dict[str, str]]:
        allowed = set(allowed_source_ids)
        normalized: list[dict[str, str]] = []
        seen: set[str] = set()
        for item in calls:
            if not isinstance(item, dict):
                continue
            source_id = str(item.get("source_id") or "").strip()
            if not source_id or source_id not in allowed:
                continue
            statement_raw = str(item.get("research_statement") or "").strip()
            statement = WordUIWorkspaceV2._normalize_research_statement(
                statement_raw,
                fallback=default_statement,
            )
            if not statement:
                continue
            signature = f"{source_id}::{statement.lower()}"
            if signature in seen:
                continue
            seen.add(signature)
            normalized.append(
                {
                    "source_id": source_id,
                    "research_statement": statement[:1400],
                }
            )
        return normalized

    def _plan_followup_internal_research_calls(
        self,
        *,
        query_text: str,
        effective_sources: list[dict[str, Any]],
        current_source_results: list[dict[str, Any]],
        already_selected_source_ids: list[str],
        executed_signatures: set[str] | None,
        max_sources: int,
    ) -> dict[str, Any] | None:
        if max_sources <= 0:
            return None
        if not effective_sources:
            return None

        lines = [
            "INTERNAL_RESEARCH_ITERATION_JSON",
            "Decide the next internal database_research tool calls for this query.",
            f"User query: {query_text}",
            f"Max source calls this step: {max_sources}",
            "Already used source IDs:",
            ", ".join(already_selected_source_ids) if already_selected_source_ids else "none",
            "Already executed call signatures (source::statement):",
            ", ".join(sorted(executed_signatures)) if executed_signatures else "none",
            "Current results summary:",
        ]
        if current_source_results:
            for item in current_source_results[-8:]:
                lines.append(
                    "- "
                    + json.dumps(
                        {
                            "source_id": item.get("source_id"),
                            "status": item.get("status"),
                            "retriever_tool": item.get("retriever_tool"),
                            "research_statement": item.get("research_statement"),
                            "summary_text": str(item.get("summary_text") or "")[:420],
                        }
                    )
                )
        else:
            lines.append("- no results yet")
        lines.append("Available sources:")
        for source in effective_sources:
            lines.append(
                "- "
                + json.dumps(
                    {
                        "source_id": source.get("source_id"),
                        "name": source.get("name"),
                        "source_type": source.get("source_type"),
                    }
                )
            )
        lines.append("Return JSON object only with shape:")
        lines.append(
            json.dumps(
                {
                    "action": "call_more_or_finish",
                    "source_queries": [
                        {
                            "source_id": "candidate_source_id",
                            "research_statement": "specific follow-up question",
                        }
                    ],
                    "reasoning": "short reason for decision",
                }
            )
        )
        lines.append(
            "Rules: Use only listed source IDs; avoid duplicate source::statement signatures; "
            "if no further calls are needed set action='finish' and source_queries=[]."
        )
        prompt = "\n".join(lines)
        response = self.agent.chat(
            [{"role": "user", "content": prompt}],
            system_context=(
                "INTERNAL_RESEARCH_ITERATION_MODE: Return strict JSON only. "
                "Do not include markdown fences."
            ),
        )
        payload = self._extract_json_object_from_text(response)
        if not isinstance(payload, dict):
            return None

        raw_action = str(payload.get("action") or "").strip().lower()
        allowed_ids = [str(source.get("source_id") or "") for source in effective_sources]
        calls = self._normalize_internal_research_calls(
            calls=list(payload.get("source_queries") or payload.get("calls") or []),
            allowed_source_ids=allowed_ids,
            default_statement=query_text,
        )
        if executed_signatures:
            calls = [
                item
                for item in calls
                if f"{item['source_id']}::{item['research_statement'].strip().lower()}" not in executed_signatures
            ]
        calls = calls[:max_sources]
        reasoning = str(payload.get("reasoning") or "").strip()
        if not reasoning:
            reasoning = "Model selected next internal research action."

        if raw_action in {"finish", "done", "none", "stop"} and not calls:
            return {
                "action": "finish",
                "planned_calls": [],
                "reasoning": reasoning[:600],
            }
        if not raw_action:
            raw_action = "call_more" if calls else "finish"
        if raw_action not in {"call_more", "call", "database_research", "research", "finish"}:
            if not calls:
                return None
            raw_action = "call_more"
        if raw_action in {"finish"} and calls:
            raw_action = "call_more"
        if raw_action in {"call", "database_research", "research"}:
            raw_action = "call_more"
        if raw_action == "call_more" and not calls:
            return None
        return {
            "action": raw_action,
            "planned_calls": calls,
            "reasoning": reasoning[:600],
        }

    def _run_internal_data_source_research_tool(
        self,
        *,
        session_id: str,
        assistant_message_id: str,
        query_text: str,
        effective_sources: list[dict[str, Any]],
        selected_source_ids: list[str],
        planned_calls: list[dict[str, str]] | None = None,
        selection_mode: str = "heuristic_fallback",
        selection_reasoning_summary: str = "",
    ) -> dict[str, Any]:
        tool_name = "research_internal_data_sources"
        call_id = str(uuid4())
        normalized_calls = self._normalize_internal_research_calls(
            calls=planned_calls or [],
            allowed_source_ids=[str(item.get("source_id") or "") for item in effective_sources],
            default_statement=(query_text or "").strip(),
        )
        self.events.create_event(
            session_id=session_id,
            message_id=assistant_message_id,
            event_type="tool_call_request",
            payload={
                "call_id": call_id,
                "tool_name": tool_name,
                "arguments": {
                    "query_text": query_text,
                    "selected_source_ids": selected_source_ids,
                    "selection_mode": selection_mode,
                    "planned_calls": normalized_calls[:12],
                },
            },
        )

        query = (query_text or "").strip()
        if not query:
            result = {
                "status": "skipped",
                "reason": "empty_query",
                "selected_count": len(selected_source_ids),
                "finding_count": 0,
                "selected_source_ids": selected_source_ids,
                "selection_mode": selection_mode,
                "selection_reasoning_summary": selection_reasoning_summary,
                "summary_text": "No query text provided for internal source research.",
                "source_results": [],
                "artifacts_created": [],
                "planned_call_count": len(normalized_calls),
                "executed_call_count": 0,
            }
            self.events.create_event(
                session_id=session_id,
                message_id=assistant_message_id,
                event_type="tool_call_response",
                payload={"call_id": call_id, "tool_name": tool_name, "result": result},
            )
            return result

        source_map = {item["source_id"]: item for item in effective_sources}
        if not source_map:
            result = {
                "status": "skipped",
                "reason": "no_selected_sources",
                "selected_count": 0,
                "finding_count": 0,
                "selected_source_ids": [],
                "selection_mode": "heuristic_fallback",
                "selection_reasoning_summary": "No enabled data sources were available for this turn.",
                "planned_call_count": 0,
                "executed_call_count": 0,
                "summary_text": "No enabled data sources were available for this turn.",
                "source_results": [],
                "artifacts_created": [],
            }
            self.events.create_event(
                session_id=session_id,
                message_id=assistant_message_id,
                event_type="tool_call_response",
                payload={"call_id": call_id, "tool_name": tool_name, "result": result},
            )
            return result

        initial_selected = [source_id for source_id in selected_source_ids if source_id in source_map]
        if not initial_selected and normalized_calls:
            initial_selected = [item["source_id"] for item in normalized_calls if item["source_id"] in source_map]

        final_selection_mode = selection_mode
        final_selection_reasoning = selection_reasoning_summary or ""
        source_results: list[dict[str, Any]] = []
        findings: list[dict[str, Any]] = []
        created_artifacts: list[str] = []
        selected_source_order: list[str] = []
        seen_signatures: set[str] = set()
        calls_executed = 0
        max_iterations = 4
        max_total_calls = 8
        planned_call_count = len(normalized_calls)
        call_queue = list(normalized_calls)
        if not call_queue and initial_selected:
            call_queue = [
                {"source_id": source_id, "research_statement": query}
                for source_id in initial_selected
            ]
            planned_call_count += len(call_queue)
        model_decision_count = 0
        fallback_attempted = False
        source_call_counts: dict[str, int] = {}
        internal_stop_reason = ""
        high_signal_count = 0
        stalled_signal_iterations = 0

        for _ in range(max_iterations):
            remaining_budget = max_total_calls - calls_executed
            if remaining_budget <= 0:
                internal_stop_reason = "budget_exhausted"
                break

            if not call_queue:
                followup = self._plan_followup_internal_research_calls(
                    query_text=query,
                    effective_sources=effective_sources,
                    current_source_results=source_results,
                    already_selected_source_ids=selected_source_order,
                    executed_signatures=seen_signatures,
                    max_sources=min(2, remaining_budget),
                )
                if followup is not None:
                    model_decision_count += 1
                    followup_reasoning = str(followup.get("reasoning") or "").strip()
                    if followup_reasoning:
                        if final_selection_reasoning:
                            final_selection_reasoning = (
                                f"{final_selection_reasoning} Iteration: {followup_reasoning}"
                            )[:900]
                        else:
                            final_selection_reasoning = followup_reasoning[:900]
                    action = str(followup.get("action") or "").strip().lower()
                    if action == "finish" and not followup.get("planned_calls"):
                        internal_stop_reason = "planner_finish"
                        break
                    planned_next = self._normalize_internal_research_calls(
                        calls=list(followup.get("planned_calls") or []),
                        allowed_source_ids=list(source_map.keys()),
                        default_statement=query,
                    )
                    if planned_next:
                        call_queue = planned_next[:remaining_budget]
                        planned_call_count += len(call_queue)
                        final_selection_mode = "model_iterative"

            if not call_queue:
                if fallback_attempted:
                    internal_stop_reason = "no_additional_calls"
                    break
                fallback_attempted = True
                if calls_executed > 0:
                    internal_stop_reason = "planner_no_followup"
                    break
                fallback_ids = initial_selected or self._select_sources_for_turn(
                    effective_sources=effective_sources,
                    query_text=query,
                    max_sources=min(2, remaining_budget),
                )
                call_queue = self._normalize_internal_research_calls(
                    calls=[
                        {"source_id": source_id, "research_statement": query}
                        for source_id in fallback_ids
                    ],
                    allowed_source_ids=list(source_map.keys()),
                    default_statement=query,
                )
                if call_queue:
                    planned_call_count += len(call_queue)
                    final_selection_mode = "heuristic_fallback"
                    if not final_selection_reasoning:
                        final_selection_reasoning = (
                            "Model planning unavailable; heuristic source selection applied."
                        )
                else:
                    internal_stop_reason = "heuristic_no_sources"
                    break

            pending = self._normalize_internal_research_calls(
                calls=call_queue,
                allowed_source_ids=list(source_map.keys()),
                default_statement=query,
            )
            call_queue = []
            executable: list[dict[str, str]] = []
            for item in pending:
                signature = f"{item['source_id']}::{item['research_statement'].strip().lower()}"
                if signature in seen_signatures:
                    continue
                if source_call_counts.get(item["source_id"], 0) >= 2:
                    continue
                executable.append(item)
            if not executable:
                internal_stop_reason = "dedup_or_source_limit"
                continue
            if remaining_budget <= 0:
                internal_stop_reason = "budget_exhausted"
                break
            executable = executable[:remaining_budget]

            prior_high_signal_count = high_signal_count
            for planned in executable:
                source_id = planned["source_id"]
                source = source_map.get(source_id)
                if source is None:
                    continue
                research_statement = planned["research_statement"]
                signature = f"{source_id}::{research_statement.strip().lower()}"
                seen_signatures.add(signature)
                source_call_counts[source_id] = source_call_counts.get(source_id, 0) + 1
                if source_id not in selected_source_order:
                    selected_source_order.append(source_id)

                db_result = self._run_database_research_tool(
                    session_id=session_id,
                    assistant_message_id=assistant_message_id,
                    source=source,
                    research_statement=research_statement,
                )
                calls_executed += 1
                execution = dict(db_result.get("retriever_output") or {})
                location = source.get("location") if isinstance(source, dict) else None
                schema_json = source.get("schema_json") if isinstance(source, dict) else None
                relevance_score = int(execution.get("relevance_score") or 0)
                if relevance_score >= 2:
                    high_signal_count += 1
                findings.append(
                    {
                        "source_id": source_id,
                        "name": source.get("name"),
                        "source_type": source.get("source_type"),
                        "relevance_score": relevance_score,
                        "location_summary": self._summarize_mapping(location, limit=4),
                        "schema_summary": self._summarize_schema(schema_json, limit=8),
                        "execution": execution,
                    }
                )
                source_results.append(
                    {
                        "source_id": source_id,
                        "status": db_result.get("status", "unknown"),
                        "retriever_tool": db_result.get("retriever_tool"),
                        "summary_text": db_result.get("summary_text"),
                        "research_output_format": db_result.get("research_output_format"),
                        "research_statement": research_statement,
                        "research_markdown": db_result.get("research_markdown"),
                        "research_document": db_result.get("research_document"),
                    }
                )
                created_artifacts.extend(db_result.get("artifacts_created") or [])
            if high_signal_count <= prior_high_signal_count:
                stalled_signal_iterations += 1
            else:
                stalled_signal_iterations = 0
            if stalled_signal_iterations >= 2 and calls_executed > 0:
                internal_stop_reason = "signal_stalled"
                break

        if not final_selection_reasoning:
            if final_selection_mode == "model_iterative":
                final_selection_reasoning = "Model iterative source planning executed."
            elif final_selection_mode == "heuristic_fallback":
                final_selection_reasoning = "Heuristic source selection executed."
            else:
                final_selection_reasoning = "Internal research selection executed."
        if model_decision_count > 0 and calls_executed > 0:
            final_selection_mode = "model_iterative"

        findings.sort(key=lambda item: int(item.get("relevance_score") or 0), reverse=True)
        findings = findings[:6]
        summary_text = (
            "Internal data-source findings: "
            + "; ".join(
                f"{item['source_id']} ({item['execution'].get('handler', 'handler')}:{item['execution'].get('status', 'unknown')}, score={item['relevance_score']})"
                for item in findings
                if item.get("source_id")
            )
        )
        result = {
            "status": "completed",
            "selected_count": len(selected_source_order),
            "selected_source_ids": selected_source_order,
            "selection_mode": final_selection_mode,
            "selection_reasoning_summary": final_selection_reasoning,
            "finding_count": len(findings),
            "planned_call_count": planned_call_count,
            "executed_call_count": calls_executed,
            "stop_reason": internal_stop_reason or "completed",
            "summary_text": summary_text,
            "findings": findings,
            "source_results": source_results,
            "artifacts_created": self._merge_artifact_ids(created_artifacts),
        }
        self.events.create_event(
            session_id=session_id,
            message_id=assistant_message_id,
            event_type="tool_call_response",
            payload={"call_id": call_id, "tool_name": tool_name, "result": result},
        )
        return result

    def _run_uploaded_documents_research_tool(
        self,
        *,
        session_id: str,
        assistant_message_id: str,
        query_text: str,
        uploaded_context_units: list[dict[str, Any]],
    ) -> dict[str, Any]:
        tool_name = "research_uploaded_documents"
        call_id = str(uuid4())
        self.events.create_event(
            session_id=session_id,
            message_id=assistant_message_id,
            event_type="tool_call_request",
            payload={
                "call_id": call_id,
                "tool_name": tool_name,
                "arguments": {
                    "query_text": query_text,
                    "context_unit_count": len(uploaded_context_units),
                },
            },
        )

        query = (query_text or "").strip()
        if not query:
            result = {
                "status": "skipped",
                "reason": "empty_query",
                "unit_count": len(uploaded_context_units),
                "artifact_count": 0,
                "summary_text": "No query text provided for uploaded-document research.",
                "artifacts_created": [],
            }
            self.events.create_event(
                session_id=session_id,
                message_id=assistant_message_id,
                event_type="tool_call_response",
                payload={"call_id": call_id, "tool_name": tool_name, "result": result},
            )
            return result

        if not uploaded_context_units:
            result = {
                "status": "skipped",
                "reason": "no_uploaded_context_units",
                "unit_count": 0,
                "artifact_count": 0,
                "summary_text": "No uploaded-document context was available for this turn.",
                "artifacts_created": [],
            }
            self.events.create_event(
                session_id=session_id,
                message_id=assistant_message_id,
                event_type="tool_call_response",
                payload={"call_id": call_id, "tool_name": tool_name, "result": result},
            )
            return result

        query_terms = self._keyword_terms(query)
        scored: list[tuple[int, dict[str, Any]]] = []
        for item in uploaded_context_units:
            content = str(item.get("content") or "")
            hay = content.lower()
            score = sum(1 for term in query_terms if term in hay)
            scored.append((score, item))
        scored.sort(key=lambda pair: pair[0], reverse=True)
        selected = [item for _, item in scored][:4]

        findings = [
            {
                "artifact_id": item.get("artifact_id"),
                "unit_type": item.get("unit_type"),
                "relevance_score": score,
                "excerpt": self._clip_text(str(item.get("content") or ""), max_chars=420),
            }
            for score, item in scored[:4]
        ]
        artifact_ids = sorted(
            {
                str(item.get("artifact_id"))
                for item in selected
                if item.get("artifact_id") is not None
            }
        )
        summary_text = (
            "Uploaded-document findings: "
            + "; ".join(f"{item['artifact_id']} ({item['unit_type']})" for item in findings if item.get("artifact_id"))
        )
        persisted = self._persist_tool_research_outputs(
            session_id=session_id,
            assistant_message_id=assistant_message_id,
            source_id="uploaded_documents",
            source_type="uploaded_documents",
            retriever_tool="uploaded_document_retriever",
            research_statement=query,
            summary_text=summary_text,
            evidence_items=findings,
            preferred_output_format="docx",
            extra_metadata={
                "uploaded_artifact_ids": artifact_ids,
                "unit_count": len(selected),
            },
        )
        result = {
            "status": "completed",
            "unit_count": len(selected),
            "artifact_count": len(artifact_ids),
            "artifact_ids": artifact_ids,
            "summary_text": summary_text,
            "findings": findings,
            "research_markdown": persisted.get("research_markdown"),
            "research_document": persisted.get("research_document"),
            "artifacts_created": persisted.get("artifacts_created", []),
        }
        self.events.create_event(
            session_id=session_id,
            message_id=assistant_message_id,
            event_type="tool_call_response",
            payload={"call_id": call_id, "tool_name": tool_name, "result": result},
        )
        return result

    def _run_database_research_tool(
        self,
        *,
        session_id: str,
        assistant_message_id: str,
        source: dict[str, Any],
        research_statement: str,
    ) -> dict[str, Any]:
        tool_name = "database_research"
        call_id = str(uuid4())
        source_id = str(source.get("source_id") or "")
        source_type = str(source.get("source_type") or "")
        self.events.create_event(
            session_id=session_id,
            message_id=assistant_message_id,
            event_type="tool_call_request",
            payload={
                "call_id": call_id,
                "tool_name": tool_name,
                "arguments": {
                    "source_id": source_id,
                    "source_type": source_type,
                    "research_statement": research_statement,
                },
            },
        )

        query_terms = self._keyword_terms(research_statement)
        retriever_output = self._run_source_retriever(
            source=source,
            research_statement=research_statement,
            query_terms=query_terms,
        )

        retriever_tool = str(retriever_output.get("handler") or "generic_source_probe")
        retriever_status = str(retriever_output.get("status") or "unknown")
        evidence_items = self._extract_evidence_items_from_retriever_output(retriever_output)
        summary_text = self._compose_database_research_summary(
            source_id=source_id,
            source_type=source_type,
            retriever_tool=retriever_tool,
            retriever_status=retriever_status,
            evidence_items=evidence_items,
        )
        research_output_format = self._resolve_research_output_format(
            source=source,
            retriever_output=retriever_output,
        )
        persisted = self._persist_tool_research_outputs(
            session_id=session_id,
            assistant_message_id=assistant_message_id,
            source_id=source_id,
            source_type=source_type,
            retriever_tool=retriever_tool,
            research_statement=research_statement,
            summary_text=summary_text,
            evidence_items=evidence_items,
            preferred_output_format=research_output_format,
            extra_metadata={
                "retriever_status": retriever_status,
                "retriever_mode": retriever_output.get("mode"),
                "research_output_format": research_output_format,
            },
        )
        result = {
            "status": "completed",
            "source_id": source_id,
            "source_type": source_type,
            "retriever_tool": retriever_tool,
            "retriever_output": retriever_output,
            "summary_text": summary_text,
            "research_output_format": research_output_format,
            "research_markdown": persisted.get("research_markdown"),
            "research_document": persisted.get("research_document"),
            "artifacts_created": persisted.get("artifacts_created", []),
        }
        self.events.create_event(
            session_id=session_id,
            message_id=assistant_message_id,
            event_type="tool_call_response",
            payload={"call_id": call_id, "tool_name": tool_name, "result": result},
        )
        return result

    def _run_source_retriever(
        self,
        *,
        source: dict[str, Any],
        research_statement: str,
        query_terms: list[str],
    ) -> dict[str, Any]:
        registry = getattr(self, "retriever_registry", None)
        if not isinstance(registry, RetrieverRegistry):
            registry = self._build_retriever_registry()
            self.retriever_registry = registry
        try:
            return registry.run(
                source=source,
                research_statement=research_statement,
                query_terms=query_terms,
            )
        except Exception as exc:  # noqa: BLE001
            return {
                "status": "failed",
                "handler": "retriever_registry_error",
                "mode": "registry",
                "error": str(exc),
                "relevance_score": 0,
            }

    def _build_retriever_registry(
        self,
        *,
        custom_source_retrievers: dict[str, Callable[..., dict[str, Any]]] | None = None,
    ) -> RetrieverRegistry:
        registry = RetrieverRegistry(
            source_overrides={
                "risk_db.suppliers": "mock_supplier_risk",
                "sales_db.orders": "mock_sales_orders",
                "finance_db.quarterly_metrics": "mock_financial_data",
                "compliance_db.audit_findings": "mock_compliance_findings",
                "hr_db.team_performance": "mock_employee_metrics",
            },
            type_overrides={
                "postgres_table": "postgres_relation_probe",
                "warehouse_view": "postgres_relation_probe",
                "search_index": "search_index_metadata",
            },
        )
        registry.register(MockSupplierRiskRetriever())
        registry.register(MockSalesOrdersRetriever())
        registry.register(MockFinancialDataRetriever())
        registry.register(MockComplianceRetriever())
        registry.register(MockEmployeeRetriever())
        registry.register(
            PostgresRelationProbeRetriever(
                probe_fn=self._execute_postgres_relation_probe,
                max_rows=5,
            )
        )
        registry.register(SearchIndexMetadataRetriever())
        registry.register(GenericMetadataRetriever())

        # Register supplementary financials retriever if DB is available
        try:
            ds_config = DataSourcesConfig.from_env()
            ds_db = DataSourcesDB(self.store)
            registry.register(SuppFinancialsRetriever(config=ds_config, db=ds_db))
            registry.source_overrides["supp_financials"] = "supp_financials"
        except Exception:
            # DB or API key not configured — skip supp financials retriever
            pass

        for source_id, fn in (custom_source_retrievers or {}).items():
            retriever_id = f"custom:{source_id}"
            registry.register(
                FunctionSourceRetriever(
                    retriever_id=retriever_id,
                    fn=fn,
                )
            )
            registry.source_overrides[source_id] = retriever_id
        return registry

    def _persist_tool_research_outputs(
        self,
        *,
        session_id: str,
        assistant_message_id: str,
        source_id: str,
        source_type: str,
        retriever_tool: str,
        research_statement: str,
        summary_text: str,
        evidence_items: list[dict[str, Any]],
        preferred_output_format: str | None = None,
        extra_metadata: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        output_dir = self._research_output_dir(session_id)
        output_dir.mkdir(parents=True, exist_ok=True)
        stamp = datetime.now(UTC).strftime("%Y%m%d-%H%M%S")
        base_name = f"{self._safe_slug(source_id)[:34]}-{stamp}-{str(assistant_message_id)[:8]}"
        artifact_group_id = str(uuid4())

        markdown_path = output_dir / f"{base_name}.md"
        markdown_text = self._build_tool_research_markdown(
            source_id=source_id,
            source_type=source_type,
            retriever_tool=retriever_tool,
            research_statement=research_statement,
            summary_text=summary_text,
            evidence_items=evidence_items,
        )
        markdown_path.write_text(markdown_text, encoding="utf-8")

        output_format = self._normalize_output_doc_format(preferred_output_format)
        metadata = {
            "source_id": source_id,
            "source_type": source_type,
            "retriever_tool": retriever_tool,
            "research_statement": research_statement[:800],
            "summary_text": summary_text[:1500],
            "research_output_format": output_format,
            **(extra_metadata or {}),
        }
        markdown_artifact = self.artifacts.create_artifact(
            session_id=session_id,
            artifact_group_id=artifact_group_id,
            artifact_type="research_markdown",
            lifecycle_state="final",
            format="md",
            filename=markdown_path.name,
            storage_uri=str(markdown_path),
            mime_type="text/markdown",
            size_bytes=markdown_path.stat().st_size,
            created_from_message_id=assistant_message_id,
            metadata=metadata,
        )
        self.events.create_event(
            session_id=session_id,
            message_id=assistant_message_id,
            event_type="artifact_created",
            payload={
                "artifact_id": markdown_artifact["artifact_id"],
                "artifact_group_id": artifact_group_id,
                "artifact_type": "research_markdown",
                "source_id": source_id,
            },
        )

        created_ids = [markdown_artifact["artifact_id"]]
        research_markdown = {
            "artifact_id": markdown_artifact["artifact_id"],
            "filename": markdown_artifact["filename"],
            "storage_uri": markdown_artifact["storage_uri"],
            "format": "md",
        }
        research_document: dict[str, Any] | None = None
        try:
            from docx import Document

            docx_path = output_dir / f"{base_name}.docx"
            doc = Document()
            doc.add_heading("Database Research Output", level=0)
            doc.add_heading("Source", level=1)
            doc.add_paragraph(f"{source_id} ({source_type})")
            doc.add_heading("Retriever", level=1)
            doc.add_paragraph(retriever_tool)
            doc.add_heading("Research Statement", level=1)
            doc.add_paragraph(research_statement)
            doc.add_heading("Summary", level=1)
            doc.add_paragraph(summary_text)
            doc.add_heading("Evidence", level=1)
            if evidence_items:
                for item in evidence_items[:20]:
                    label = item.get("label") or item.get("type") or "evidence"
                    text = item.get("text") or ""
                    doc.add_paragraph(f"{label}: {text}", style="List Bullet")
            else:
                doc.add_paragraph("No evidence items captured.")
            doc.save(str(docx_path))

            output_path = docx_path
            output_detail = {"mode": "native_generate", "backend": "python-docx"}
            if output_format != "docx":
                output_path = output_dir / f"{base_name}.{output_format}"
                output_detail = self._materialize_report_export(
                    source_path=docx_path,
                    export_path=output_path,
                    export_format=output_format,
                    topic=f"Research Output: {source_id}",
                    source_filename=docx_path.name,
                )
                docx_path.unlink(missing_ok=True)

            doc_artifact = self.artifacts.create_artifact(
                session_id=session_id,
                artifact_group_id=artifact_group_id,
                artifact_type="research_output_doc",
                lifecycle_state="final",
                format=output_format,
                filename=output_path.name,
                storage_uri=str(output_path),
                mime_type=self._mime_for_export_format(output_format),
                size_bytes=output_path.stat().st_size,
                created_from_message_id=assistant_message_id,
                source_artifact_id=markdown_artifact["artifact_id"],
                metadata={
                    **metadata,
                    "research_output_mode": output_detail["mode"],
                    "research_output_backend": output_detail["backend"],
                    **(
                        {"research_output_note": output_detail["note"]}
                        if output_detail.get("note")
                        else {}
                    ),
                },
            )
            self.events.create_event(
                session_id=session_id,
                message_id=assistant_message_id,
                event_type="artifact_created",
                payload={
                    "artifact_id": doc_artifact["artifact_id"],
                    "artifact_group_id": artifact_group_id,
                    "artifact_type": "research_output_doc",
                    "source_id": source_id,
                    "source_artifact_id": markdown_artifact["artifact_id"],
                    "format": output_format,
                },
            )
            created_ids.append(doc_artifact["artifact_id"])
            research_document = {
                "artifact_id": doc_artifact["artifact_id"],
                "filename": doc_artifact["filename"],
                "storage_uri": doc_artifact["storage_uri"],
                "format": output_format,
            }
        except Exception as exc:  # noqa: BLE001
            self.events.create_event(
                session_id=session_id,
                message_id=assistant_message_id,
                event_type="error",
                payload={
                    "message": "failed to create research output document artifact",
                    "source_id": source_id,
                    "error": str(exc),
                },
            )

        return {
            "research_markdown": research_markdown,
            "research_document": research_document,
            "artifacts_created": created_ids,
        }

    @staticmethod
    def _build_tool_research_markdown(
        *,
        source_id: str,
        source_type: str,
        retriever_tool: str,
        research_statement: str,
        summary_text: str,
        evidence_items: list[dict[str, Any]],
    ) -> str:
        lines = [
            "# Database Research Output",
            "",
            f"Generated at: {datetime.now(UTC).isoformat()}",
            "",
            f"- Source ID: `{source_id}`",
            f"- Source Type: `{source_type}`",
            f"- Retriever Tool: `{retriever_tool}`",
            "",
            "## Research Statement",
            research_statement,
            "",
            "## Summary",
            summary_text,
            "",
            "## Evidence",
        ]
        if evidence_items:
            for item in evidence_items[:20]:
                label = item.get("label") or item.get("type") or "evidence"
                text = item.get("text") or ""
                lines.append(f"- **{label}**: {text}")
        else:
            lines.append("- No evidence items captured.")
        return "\n".join(lines)

    @staticmethod
    def _extract_evidence_items_from_retriever_output(retriever_output: dict[str, Any]) -> list[dict[str, Any]]:
        evidence: list[dict[str, Any]] = []
        sample_rows = retriever_output.get("sample_rows")
        if isinstance(sample_rows, list):
            for index, row in enumerate(sample_rows[:8], start=1):
                if isinstance(row, dict):
                    rendered = ", ".join(f"{key}={value}" for key, value in row.items())
                else:
                    rendered = str(row)
                evidence.append(
                    {
                        "type": "row_sample",
                        "label": f"row_{index}",
                        "text": rendered[:500],
                    }
                )
        matched_terms = retriever_output.get("matched_terms")
        if isinstance(matched_terms, list):
            for term in matched_terms[:8]:
                evidence.append(
                    {
                        "type": "matched_term",
                        "label": "term",
                        "text": str(term),
                    }
                )
        if not evidence:
            relation = retriever_output.get("relation")
            if relation:
                evidence.append(
                    {
                        "type": "relation",
                        "label": "relation",
                        "text": str(relation),
                    }
                )
        return evidence

    @staticmethod
    def _compose_database_research_summary(
        *,
        source_id: str,
        source_type: str,
        retriever_tool: str,
        retriever_status: str,
        evidence_items: list[dict[str, Any]],
    ) -> str:
        evidence_count = len(evidence_items)
        return (
            f"Database research completed for {source_id} ({source_type}) using {retriever_tool}. "
            f"Retriever status: {retriever_status}. Evidence items captured: {evidence_count}."
        )

    @staticmethod
    def _merge_artifact_ids(*groups: list[str] | tuple[str, ...]) -> list[str]:
        merged: list[str] = []
        seen: set[str] = set()
        for group in groups:
            for item in group:
                value = str(item)
                if not value or value in seen:
                    continue
                seen.add(value)
                merged.append(value)
        return merged

    @staticmethod
    def _clip_text(text: str, *, max_chars: int) -> str:
        if len(text) <= max_chars:
            return text
        return text[: max_chars - 3].rstrip() + "..."

    @staticmethod
    def _summarize_mapping(value: Any, *, limit: int) -> list[str]:
        if not isinstance(value, dict):
            return []
        pairs: list[str] = []
        for key, item in value.items():
            if len(pairs) >= limit:
                break
            if isinstance(item, (str, int, float, bool)):
                pairs.append(f"{key}={item}")
            else:
                pairs.append(str(key))
        return pairs

    @staticmethod
    def _summarize_schema(value: Any, *, limit: int) -> list[str]:
        if not isinstance(value, dict):
            return []
        fields = value.get("fields")
        if isinstance(fields, list):
            out: list[str] = []
            for field in fields:
                if len(out) >= limit:
                    break
                if isinstance(field, dict):
                    name = field.get("name")
                    if name:
                        out.append(str(name))
                elif isinstance(field, str):
                    out.append(field)
            if out:
                return out
        return [str(key) for key in list(value.keys())[:limit]]

    def _execute_postgres_relation_probe(
        self,
        *,
        source: dict[str, Any],
        query_terms: list[str],
        max_rows: int,
    ) -> dict[str, Any]:
        source_id = str(source.get("source_id") or "")
        location = source.get("location") if isinstance(source, dict) else {}
        schema_json = source.get("schema_json") if isinstance(source, dict) else {}
        schema_name, relation_name = self._resolve_relation_name(location, source_id)

        metadata_hay = self._source_haystack(source)
        metadata_score = self._match_terms_count(metadata_hay, query_terms)
        if not relation_name:
            return {
                "status": "completed",
                "handler": "postgres_relation_probe",
                "mode": "metadata_only",
                "reason": "missing_relation_name",
                "relation": None,
                "rows_returned": 0,
                "sample_rows": [],
                "relevance_score": metadata_score,
            }

        relation = f"{schema_name}.{relation_name}"
        if getattr(self, "store", None) is None:
            return {
                "status": "completed",
                "handler": "postgres_relation_probe",
                "mode": "metadata_only",
                "reason": "store_unavailable",
                "relation": relation,
                "rows_returned": 0,
                "sample_rows": [],
                "location_summary": self._summarize_mapping(location, limit=4),
                "schema_summary": self._summarize_schema(schema_json, limit=8),
                "relevance_score": metadata_score,
            }

        try:
            from psycopg import sql

            with self.store.connection() as conn:
                with conn.cursor() as cur:
                    cur.execute("select to_regclass(%s) as relation_oid", (relation,))
                    regclass = cur.fetchone()
                    if not regclass or regclass.get("relation_oid") is None:
                        return {
                            "status": "failed",
                            "handler": "postgres_relation_probe",
                            "mode": "live_probe",
                            "relation": relation,
                            "error": "relation_not_found",
                            "rows_returned": 0,
                            "sample_rows": [],
                            "relevance_score": metadata_score,
                        }

                    cur.execute(
                        """
                        select column_name, data_type
                        from information_schema.columns
                        where table_schema = %s and table_name = %s
                        order by ordinal_position
                        """,
                        (schema_name, relation_name),
                    )
                    column_rows = cur.fetchall() or []
                    column_names = [str(row.get("column_name")) for row in column_rows if row.get("column_name")]
                    selected_columns = column_names[:8]
                    if not selected_columns:
                        return {
                            "status": "completed",
                            "handler": "postgres_relation_probe",
                            "mode": "live_probe",
                            "relation": relation,
                            "reason": "no_columns_detected",
                            "rows_returned": 0,
                            "sample_rows": [],
                            "relevance_score": metadata_score,
                        }

                    text_columns = []
                    for row in column_rows:
                        col_name = row.get("column_name")
                        data_type = str(row.get("data_type") or "").lower()
                        if not col_name:
                            continue
                        if (
                            "char" in data_type
                            or "text" in data_type
                            or "json" in data_type
                            or "uuid" in data_type
                            or "date" in data_type
                            or "time" in data_type
                        ):
                            text_columns.append(str(col_name))

                    filter_columns = text_columns[:3]
                    patterns = [f"%{term}%" for term in query_terms[:5]]
                    where_clause = sql.SQL("")
                    params: list[Any] = []
                    if patterns and filter_columns:
                        predicates = []
                        for col in filter_columns:
                            predicates.append(sql.SQL("{}::text ilike any(%s)").format(sql.Identifier(col)))
                            params.append(patterns)
                        where_clause = sql.SQL(" where ") + sql.SQL(" or ").join(predicates)

                    query_sql = sql.SQL(
                        "select {cols} from {schema}.{table}{where_clause} limit {limit}"
                    ).format(
                        cols=sql.SQL(", ").join(sql.Identifier(col) for col in selected_columns),
                        schema=sql.Identifier(schema_name),
                        table=sql.Identifier(relation_name),
                        where_clause=where_clause,
                        limit=sql.Literal(max_rows),
                    )
                    cur.execute(query_sql, params)
                    rows = cur.fetchall() or []
                    sample_rows = self._sanitize_sample_rows(
                        rows,
                        max_rows=max_rows,
                        max_value_chars=240,
                    )
                    sample_hay = " ".join(
                        " ".join(str(value) for value in row.values())
                        for row in sample_rows
                    ).lower()
                    sample_score = self._match_terms_count(sample_hay, query_terms)
                    query_preview = query_sql.as_string(conn)
                    return {
                        "status": "completed",
                        "handler": "postgres_relation_probe",
                        "mode": "live_probe",
                        "relation": relation,
                        "query_preview": query_preview,
                        "selected_columns": selected_columns,
                        "filter_columns": filter_columns,
                        "rows_returned": len(sample_rows),
                        "sample_rows": sample_rows,
                        "relevance_score": metadata_score + sample_score,
                    }
        except Exception as exc:  # noqa: BLE001
            return {
                "status": "failed",
                "handler": "postgres_relation_probe",
                "mode": "live_probe",
                "relation": relation,
                "error": str(exc),
                "rows_returned": 0,
                "sample_rows": [],
                "relevance_score": metadata_score,
            }

    def _execute_search_index_probe(
        self,
        *,
        source: dict[str, Any],
        query_terms: list[str],
    ) -> dict[str, Any]:
        source_id = str(source.get("source_id") or "")
        location = source.get("location") if isinstance(source, dict) else {}
        if not isinstance(location, dict):
            location = {}
        index_name = (
            str(location.get("index") or "").strip()
            or str(location.get("collection") or "").strip()
            or source_id
        )
        hay = self._source_haystack(source)
        matched_terms = [term for term in query_terms if term in hay]
        return {
            "status": "completed",
            "handler": "search_index_probe",
            "mode": "metadata_only",
            "index_name": index_name,
            "matched_terms": matched_terms[:8],
            "relevance_score": len(matched_terms),
        }

    def _execute_generic_source_probe(
        self,
        *,
        source: dict[str, Any],
        query_terms: list[str],
    ) -> dict[str, Any]:
        hay = self._source_haystack(source)
        score = self._match_terms_count(hay, query_terms)
        return {
            "status": "completed",
            "handler": "generic_source_probe",
            "mode": "metadata_only",
            "relevance_score": score,
        }

    @staticmethod
    def _resolve_relation_name(location: Any, source_id: str) -> tuple[str, str]:
        if not isinstance(location, dict):
            location = {}

        schema_name = str(location.get("schema") or location.get("schema_name") or "").strip()
        relation_name = str(
            location.get("table")
            or location.get("table_name")
            or location.get("view")
            or location.get("view_name")
            or ""
        ).strip()
        relation_full = str(
            location.get("relation")
            or location.get("qualified_name")
            or location.get("table_fqn")
            or ""
        ).strip()

        if relation_full and not relation_name:
            if "." in relation_full:
                left, right = relation_full.split(".", 1)
                schema_name = schema_name or left
                relation_name = right
            else:
                relation_name = relation_full

        if not relation_name and "." in source_id:
            left, right = source_id.split(".", 1)
            if left and right:
                schema_name = schema_name or left
                relation_name = relation_name or right

        if not schema_name:
            schema_name = "public"
        return schema_name, relation_name

    @staticmethod
    def _source_haystack(source: dict[str, Any]) -> str:
        location = source.get("location")
        schema_json = source.get("schema_json")
        return " ".join(
            [
                str(source.get("source_id") or ""),
                str(source.get("name") or ""),
                str(source.get("source_type") or ""),
                str(location or ""),
                str(schema_json or ""),
            ]
        ).lower()

    @staticmethod
    def _match_terms_count(haystack: str, query_terms: list[str]) -> int:
        if not haystack or not query_terms:
            return 0
        return sum(1 for term in query_terms if term in haystack)

    @staticmethod
    def _sanitize_sample_rows(
        rows: list[dict[str, Any]],
        *,
        max_rows: int,
        max_value_chars: int,
    ) -> list[dict[str, Any]]:
        sanitized: list[dict[str, Any]] = []
        for row in rows[:max_rows]:
            if not isinstance(row, dict):
                continue
            normalized: dict[str, Any] = {}
            for key, value in row.items():
                if value is None:
                    normalized[str(key)] = None
                    continue
                if isinstance(value, (int, float, bool)):
                    normalized[str(key)] = value
                    continue
                rendered = str(value)
                if len(rendered) > max_value_chars:
                    rendered = rendered[: max_value_chars - 3].rstrip() + "..."
                normalized[str(key)] = rendered
            sanitized.append(normalized)
        return sanitized

    @staticmethod
    def _extract_text_for_ingestion(path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix in (".txt", ".md", ".csv", ".json", ".html", ".xml"):
            return path.read_text(encoding="utf-8", errors="replace")
        if suffix == ".docx":
            try:
                from docx import Document

                doc = Document(str(path))
                parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                return "\n".join(parts)
            except Exception:
                return ""
        return ""

    @staticmethod
    def _summarize_ingestion_text(text: str, max_chars: int = 1200) -> str:
        normalized = " ".join(text.split())
        if len(normalized) <= max_chars:
            return normalized
        return normalized[: max_chars - 3].rstrip() + "..."

    @staticmethod
    def _chunk_text_for_ingestion(
        text: str,
        *,
        chunk_chars: int,
        overlap_chars: int,
        max_chunks: int,
    ) -> list[str]:
        normalized = " ".join(text.split())
        if not normalized:
            return []
        if chunk_chars <= 0:
            return [normalized]
        if overlap_chars >= chunk_chars:
            overlap_chars = max(0, chunk_chars // 4)
        chunks: list[str] = []
        start = 0
        while start < len(normalized) and len(chunks) < max_chunks:
            end = min(len(normalized), start + chunk_chars)
            chunk = normalized[start:end].strip()
            if chunk:
                chunks.append(chunk)
            if end >= len(normalized):
                break
            start = max(start + 1, end - overlap_chars)
        return chunks

    @staticmethod
    def _keyword_terms(text: str) -> list[str]:
        parts = [item.strip().lower() for item in text.split() if item.strip()]
        deduped: list[str] = []
        for token in parts:
            if len(token) < 3:
                continue
            if token not in deduped:
                deduped.append(token)
        return deduped[:20]

    @staticmethod
    def _is_generic_research_statement(text: str) -> bool:
        normalized = " ".join(str(text or "").lower().split())
        if not normalized:
            return True
        generic_markers = (
            "research this",
            "look into this",
            "more details",
            "general analysis",
            "analyze this",
            "investigate this",
            "find info",
            "more context",
            "summarize this",
        )
        if normalized in {"research", "analyze", "investigate", "details", "summary"}:
            return True
        return any(marker == normalized for marker in generic_markers)

    @staticmethod
    def _normalize_research_statement(statement: str, *, fallback: str) -> str:
        cleaned = " ".join(str(statement or "").split())
        if len(cleaned) < 16 or WordUIWorkspaceV2._is_generic_research_statement(cleaned):
            cleaned = " ".join(str(fallback or "").split())
        if len(cleaned) < 16:
            return ""
        return cleaned[:1400]

    def _resolve_allowed_roots(self, allowed_roots: list[str | Path] | None) -> list[Path]:
        if not allowed_roots:
            return []
        return [Path(item).expanduser().resolve() for item in allowed_roots]

    def _is_path_allowed(self, path: Path) -> bool:
        if not self.allowed_roots:
            return True
        for root in self.allowed_roots:
            try:
                path.relative_to(root)
                return True
            except ValueError:
                continue
        return False

    def _normalize_path(self, file_path: str) -> str | dict[str, Any]:
        if not file_path or not isinstance(file_path, str):
            return self._error("INVALID_ARGUMENT", "file_path must be a non-empty string")
        return str(Path(file_path).expanduser().resolve())

    def _is_valid_user_id(self, user_id: str) -> bool:
        return bool(user_id and len(user_id) == 9 and user_id.isdigit())

    def _maybe_auto_rename_session(self, session_id: str) -> None:
        """Auto-generate a session title after the 2nd user message.

        Skipped when the session already has a title_source set (manual or auto).
        Failures are logged but never propagate.
        """
        try:
            session = self.sessions.get_session(session_id)
            if session is None:
                return
            meta = session.get("metadata") or {}
            if meta.get("title_source"):
                return
            user_count = self.messages.count_user_messages(session_id)
            if user_count < 2:
                return
            rows = self.messages.list_messages(session_id, limit=3)
            user_texts = [r["content_text"] for r in rows if r["role"] == "user" and r.get("content_text")]
            if not user_texts:
                return
            title = self.agent.generate_session_title(user_texts)
            if not title:
                return
            existing_meta = dict(meta)
            existing_meta["title_source"] = "auto"
            self.sessions.update_session(session_id, title=title, metadata=existing_meta)
        except Exception:
            logger.exception("auto-rename failed for session %s", session_id)

    def _render_pdf_preview_html(self, path: Path) -> str:
        import base64

        size_bytes = path.stat().st_size
        max_inline_bytes = 4 * 1024 * 1024
        pretty_size = self._human_readable_size(size_bytes)
        if size_bytes > max_inline_bytes:
            return (
                "<html><body style='font-family:Arial,sans-serif;padding:16px;background:#fafafa;color:#1f2937;'>"
                "<h3 style='margin-top:0;'>PDF Preview (Large File Mode)</h3>"
                f"<p>File size: <strong>{self._escape_html(pretty_size)}</strong> "
                f"({size_bytes} bytes). Inline rendering is disabled above {max_inline_bytes} bytes.</p>"
                "<p>You can open the source file directly from the artifact list.</p>"
                f"<p style='word-break:break-all;'>Path: <code>{self._escape_html(str(path))}</code></p>"
                "</body></html>"
            )
        payload = base64.b64encode(path.read_bytes()).decode("ascii")
        src = f"data:application/pdf;base64,{payload}"
        return (
            "<html><body style='margin:0;padding:0;background:#f7f7f7;'>"
            "<div style='padding:8px 12px;font-family:Arial,sans-serif;border-bottom:1px solid #ddd;background:#fff;'>"
            f"<strong>PDF Preview:</strong> {self._escape_html(path.name)}"
            f"<span style='margin-left:8px;color:#555;'>({self._escape_html(pretty_size)})</span>"
            "</div>"
            f"<embed src='{src}' type='application/pdf' style='width:100%;height:calc(100vh - 44px);' />"
            "</body></html>"
        )

    def _render_xlsx_preview_html(self, path: Path) -> str:
        size_bytes = path.stat().st_size
        if size_bytes > 8 * 1024 * 1024:
            max_sheets, max_rows, max_cols = 1, 14, 8
            mode_label = "large-file mode"
        elif size_bytes > 2 * 1024 * 1024:
            max_sheets, max_rows, max_cols = 2, 22, 10
            mode_label = "balanced mode"
        else:
            max_sheets, max_rows, max_cols = 3, 35, 12
            mode_label = "full mode"
        sheets = self._extract_xlsx_preview_data(
            path=path,
            max_sheets=max_sheets,
            max_rows=max_rows,
            max_cols=max_cols,
        )
        if not sheets:
            return (
                "<html><body style='font-family:Arial,sans-serif;padding:16px;'>"
                "<h3>XLSX Preview</h3><p>No readable worksheet data found.</p>"
                "</body></html>"
            )
        parts = [
            "<html><body style='font-family:Arial,sans-serif;padding:16px;background:#fafafa;'>",
            f"<h3 style='margin-top:0;'>Spreadsheet Preview: {self._escape_html(path.name)}</h3>",
            (
                "<p style='color:#444;'>"
                f"Preview mode: {self._escape_html(mode_label)}. "
                f"Showing up to {max_sheets} sheet(s), {max_rows} rows, {max_cols} columns per sheet."
                "</p>"
            ),
        ]
        for sheet in sheets:
            parts.append(f"<h4>{self._escape_html(sheet['name'])}</h4>")
            parts.append("<div style='overflow:auto;border:1px solid #ddd;background:#fff;margin-bottom:16px;'>")
            parts.append("<table style='border-collapse:collapse;min-width:640px;'>")
            for row_idx, row in enumerate(sheet["rows"]):
                parts.append("<tr>")
                for value in row:
                    tag = "th" if row_idx == 0 else "td"
                    style = (
                        "border:1px solid #e3e3e3;padding:6px 8px;text-align:left;"
                        "font-size:12px;vertical-align:top;white-space:pre-wrap;"
                    )
                    if tag == "th":
                        style += "background:#f3f4f6;font-weight:600;"
                    parts.append(f"<{tag} style='{style}'>{self._escape_html(str(value or ''))}</{tag}>")
                parts.append("</tr>")
            parts.append("</table></div>")
        parts.append("</body></html>")
        return "".join(parts)

    @staticmethod
    def _human_readable_size(size_bytes: int) -> str:
        units = ["B", "KB", "MB", "GB"]
        value = float(max(size_bytes, 0))
        unit_index = 0
        while value >= 1024 and unit_index < len(units) - 1:
            value /= 1024
            unit_index += 1
        if unit_index == 0:
            return f"{int(value)} {units[unit_index]}"
        return f"{value:.1f} {units[unit_index]}"

    def _extract_xlsx_preview_data(
        self,
        *,
        path: Path,
        max_sheets: int,
        max_rows: int,
        max_cols: int,
    ) -> list[dict[str, Any]]:
        try:
            from openpyxl import load_workbook

            wb = load_workbook(filename=str(path), read_only=True, data_only=True)
            out: list[dict[str, Any]] = []
            for sheet_name in wb.sheetnames[:max_sheets]:
                ws = wb[sheet_name]
                rows: list[list[str]] = []
                for row in ws.iter_rows(min_row=1, max_row=max_rows, max_col=max_cols, values_only=True):
                    rows.append([str(value) if value is not None else "" for value in row])
                if rows:
                    out.append({"name": sheet_name, "rows": rows})
            wb.close()
            if out:
                return out
        except Exception:
            pass

        try:
            return self._extract_xlsx_preview_data_ooxml(
                path=path,
                max_sheets=max_sheets,
                max_rows=max_rows,
                max_cols=max_cols,
            )
        except Exception:
            return []

    @staticmethod
    def _extract_xlsx_preview_data_ooxml(
        *,
        path: Path,
        max_sheets: int,
        max_rows: int,
        max_cols: int,
    ) -> list[dict[str, Any]]:
        import re
        import xml.etree.ElementTree as ET
        import zipfile

        ns_main = {"main": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
        ns_rel = {"rel": "http://schemas.openxmlformats.org/officeDocument/2006/relationships"}
        ns_pkg_rel = {"pkg": "http://schemas.openxmlformats.org/package/2006/relationships"}

        with zipfile.ZipFile(path, mode="r") as zf:
            shared_strings: list[str] = []
            if "xl/sharedStrings.xml" in zf.namelist():
                root = ET.fromstring(zf.read("xl/sharedStrings.xml"))
                for item in root.findall(".//main:si", ns_main):
                    texts = [node.text or "" for node in item.findall(".//main:t", ns_main)]
                    shared_strings.append("".join(texts))

            workbook_root = ET.fromstring(zf.read("xl/workbook.xml"))
            rels_root = ET.fromstring(zf.read("xl/_rels/workbook.xml.rels"))
            rel_map: dict[str, str] = {}
            for rel in rels_root.findall("pkg:Relationship", ns_pkg_rel):
                rid = rel.attrib.get("Id") or ""
                target = rel.attrib.get("Target") or ""
                if rid and target:
                    rel_map[rid] = target

            sheets: list[dict[str, Any]] = []
            for sheet in workbook_root.findall(".//main:sheet", ns_main)[:max_sheets]:
                sheet_name = sheet.attrib.get("name") or "Sheet"
                rid = sheet.attrib.get(f"{{{ns_rel['rel']}}}id") or ""
                target = rel_map.get(rid, "")
                if not target:
                    continue
                sheet_path = f"xl/{target}" if not target.startswith("xl/") else target
                if sheet_path not in zf.namelist():
                    continue
                sheet_root = ET.fromstring(zf.read(sheet_path))
                rows_data: list[list[str]] = []
                for row in sheet_root.findall(".//main:sheetData/main:row", ns_main):
                    row_values = [""] * max_cols
                    for cell in row.findall("main:c", ns_main):
                        ref = cell.attrib.get("r") or ""
                        col_match = re.match(r"([A-Z]+)", ref)
                        if not col_match:
                            continue
                        col_letters = col_match.group(1)
                        col_index = 0
                        for ch in col_letters:
                            col_index = (col_index * 26) + (ord(ch) - 64)
                        col_index -= 1
                        if col_index < 0 or col_index >= max_cols:
                            continue
                        cell_type = cell.attrib.get("t") or ""
                        value = ""
                        if cell_type == "inlineStr":
                            texts = [node.text or "" for node in cell.findall(".//main:t", ns_main)]
                            value = "".join(texts)
                        else:
                            v = cell.find("main:v", ns_main)
                            raw = v.text if v is not None and v.text is not None else ""
                            if cell_type == "s" and raw.isdigit():
                                idx = int(raw)
                                value = shared_strings[idx] if 0 <= idx < len(shared_strings) else raw
                            else:
                                value = raw
                        row_values[col_index] = value
                    rows_data.append(row_values)
                    if len(rows_data) >= max_rows:
                        break
                if rows_data:
                    sheets.append({"name": sheet_name, "rows": rows_data})
            return sheets

    @staticmethod
    def _escape_html(value: str) -> str:
        return (
            str(value)
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
            .replace("'", "&#39;")
        )

    def _text_to_html(self, text: str) -> str:
        escaped = self._escape_html(text)
        return f"<html><body><pre>{escaped}</pre></body></html>"

    def _ok(self, **payload: Any) -> dict[str, Any]:
        return {"status": "ok", "contract_version": self.contract_version, **payload}

    def _error(self, error_code: str, message: str, **payload: Any) -> dict[str, Any]:
        return {
            "status": "error",
            "contract_version": self.contract_version,
            "error_code": error_code,
            "message": message,
            **payload,
        }
