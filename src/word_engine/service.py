"""DOCX service implementing the v1 contract."""

from __future__ import annotations

import logging
import os
from pathlib import Path
import re
import shutil
import subprocess
import tempfile
import time
from typing import Any
import importlib.util
import xml.etree.ElementTree as ET
import zipfile

from docx import Document

from .config import EngineConfig
from .errors import ErrorCode
from .locking import FileLockManager
from .models import ParagraphStyleTemplate, RunStyleTemplate, SelectorRange
from .responses import error, ok

logger = logging.getLogger(__name__)

HEADING_STYLE_PATTERN = re.compile(r"^heading\s+(\d+)$", re.IGNORECASE)
WORDPROCESSINGML_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WORD_NS_MAP = {"w": WORDPROCESSINGML_NS}


class ServiceError(Exception):
    """Domain error that maps to stable API error responses."""

    def __init__(self, error_code: ErrorCode, message: str):
        super().__init__(message)
        self.error_code = error_code
        self.message = message


class WordDocumentService:
    """Main application service used by MCP tool handlers and agent workflows."""

    def __init__(
        self,
        config: EngineConfig | None = None,
        lock_manager: FileLockManager | None = None,
    ) -> None:
        self.config = config or EngineConfig()
        self.lock_manager = lock_manager or FileLockManager()

    # Public API -----------------------------------------------------------------

    def create_document(
        self,
        file_path: str,
        template_path: str | None = None,
        title: str | None = None,
        author: str | None = None,
    ) -> dict[str, Any]:
        event = "create_document"
        started = time.perf_counter()
        try:
            output_path = self._resolve_docx_path(file_path, must_exist=False)
            with self.lock_manager.acquire(output_path):
                if template_path:
                    template = self._resolve_docx_path(template_path, must_exist=True)
                    doc = Document(str(template))
                else:
                    doc = Document()

                if title:
                    doc.core_properties.title = title
                if author:
                    doc.core_properties.author = author

                self._save_document_atomic(output_path, doc)
                response = ok(self.config.contract_version, file_path=str(output_path))
                self._log_success(event, output_path, started)
                return response
        except ServiceError as exc:
            self._log_failure(event, file_path, exc.error_code, started)
            return error(self.config.contract_version, exc.error_code, exc.message)
        except Exception as exc:  # pragma: no cover - unexpected protection path
            self._log_failure(event, file_path, ErrorCode.INTERNAL_ERROR, started)
            return error(
                self.config.contract_version,
                ErrorCode.INTERNAL_ERROR,
                f"Unexpected error creating document: {exc}",
            )

    def copy_document(self, source_path: str, destination_path: str | None = None) -> dict[str, Any]:
        event = "copy_document"
        started = time.perf_counter()
        try:
            source = self._resolve_docx_path(source_path, must_exist=True)
            destination = self._resolve_docx_path(
                destination_path or self._default_copy_path(source),
                must_exist=False,
            )

            with self.lock_manager.acquire_many([source, destination]):
                shutil.copy2(source, destination)

            response = ok(self.config.contract_version, file_path=str(destination))
            self._log_success(event, destination, started)
            return response
        except ServiceError as exc:
            self._log_failure(event, source_path, exc.error_code, started)
            return error(self.config.contract_version, exc.error_code, exc.message)
        except Exception as exc:  # pragma: no cover
            self._log_failure(event, source_path, ErrorCode.INTERNAL_ERROR, started)
            return error(
                self.config.contract_version,
                ErrorCode.INTERNAL_ERROR,
                f"Unexpected error copying document: {exc}",
            )

    def get_document_info(self, file_path: str) -> dict[str, Any]:
        event = "get_document_info"
        started = time.perf_counter()
        try:
            path = self._resolve_docx_path(file_path, must_exist=True)
            doc = Document(str(path))
            full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            words = re.findall(r"\b\w+\b", full_text)

            metadata = {
                "title": doc.core_properties.title,
                "author": doc.core_properties.author,
                "subject": doc.core_properties.subject,
                "created": str(doc.core_properties.created) if doc.core_properties.created else None,
                "modified": str(doc.core_properties.modified) if doc.core_properties.modified else None,
            }
            counts = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "words": len(words),
                "characters": len(full_text),
            }
            self._log_success(event, path, started)
            return ok(self.config.contract_version, metadata=metadata, counts=counts)
        except ServiceError as exc:
            self._log_failure(event, file_path, exc.error_code, started)
            return error(self.config.contract_version, exc.error_code, exc.message)
        except Exception as exc:  # pragma: no cover
            self._log_failure(event, file_path, ErrorCode.INTERNAL_ERROR, started)
            return error(
                self.config.contract_version,
                ErrorCode.INTERNAL_ERROR,
                f"Unexpected error reading document info: {exc}",
            )

    def get_document_outline(self, file_path: str) -> dict[str, Any]:
        event = "get_document_outline"
        started = time.perf_counter()
        try:
            path = self._resolve_docx_path(file_path, must_exist=True)
            doc = Document(str(path))
            headings: list[dict[str, Any]] = []
            for index, paragraph in enumerate(doc.paragraphs):
                style_name = paragraph.style.name if paragraph.style else ""
                level = self._extract_heading_level(style_name)
                if level is None:
                    continue
                headings.append(
                    {
                        "level": level,
                        "text": paragraph.text,
                        "paragraph_index": index,
                        "style_name": style_name,
                    }
                )
            self._log_success(event, path, started)
            return ok(self.config.contract_version, headings=headings)
        except ServiceError as exc:
            self._log_failure(event, file_path, exc.error_code, started)
            return error(self.config.contract_version, exc.error_code, exc.message)
        except Exception as exc:  # pragma: no cover
            self._log_failure(event, file_path, ErrorCode.INTERNAL_ERROR, started)
            return error(
                self.config.contract_version,
                ErrorCode.INTERNAL_ERROR,
                f"Unexpected error reading outline: {exc}",
            )

    def get_paragraph_text(self, file_path: str, paragraph_index: int) -> dict[str, Any]:
        event = "get_paragraph_text"
        started = time.perf_counter()
        try:
            path = self._resolve_docx_path(file_path, must_exist=True)
            doc = Document(str(path))
            paragraph = self._paragraph_by_index(doc, paragraph_index)
            self._log_success(event, path, started)
            return ok(
                self.config.contract_version,
                paragraph={
                    "index": paragraph_index,
                    "text": paragraph.text,
                    "style_name": paragraph.style.name if paragraph.style else None,
                },
            )
        except ServiceError as exc:
            self._log_failure(event, file_path, exc.error_code, started)
            return error(self.config.contract_version, exc.error_code, exc.message)
        except Exception as exc:  # pragma: no cover
            self._log_failure(event, file_path, ErrorCode.INTERNAL_ERROR, started)
            return error(
                self.config.contract_version,
                ErrorCode.INTERNAL_ERROR,
                f"Unexpected error reading paragraph: {exc}",
            )

    def find_text(
        self,
        file_path: str,
        query: str,
        match_case: bool = False,
        whole_word: bool = False,
    ) -> dict[str, Any]:
        event = "find_text"
        started = time.perf_counter()
        try:
            if not query:
                raise ServiceError(ErrorCode.INVALID_ARGUMENT, "query must be a non-empty string")
            path = self._resolve_docx_path(file_path, must_exist=True)
            doc = Document(str(path))
            matches: list[dict[str, Any]] = []

            if whole_word:
                regex = re.compile(
                    rf"\b{re.escape(query)}\b",
                    0 if match_case else re.IGNORECASE,
                )
                for paragraph_index, paragraph in enumerate(doc.paragraphs):
                    for found in regex.finditer(paragraph.text):
                        matches.append(
                            {
                                "paragraph_index": paragraph_index,
                                "start": found.start(),
                                "end": found.end(),
                                "match": found.group(0),
                            }
                        )
            else:
                needle = query if match_case else query.lower()
                for paragraph_index, paragraph in enumerate(doc.paragraphs):
                    haystack = paragraph.text if match_case else paragraph.text.lower()
                    offset = 0
                    while True:
                        idx = haystack.find(needle, offset)
                        if idx < 0:
                            break
                        matches.append(
                            {
                                "paragraph_index": paragraph_index,
                                "start": idx,
                                "end": idx + len(query),
                                "match": paragraph.text[idx : idx + len(query)],
                            }
                        )
                        offset = idx + len(query)

            self._log_success(event, path, started)
            return ok(
                self.config.contract_version,
                matches=matches,
                total_matches=len(matches),
            )
        except ServiceError as exc:
            self._log_failure(event, file_path, exc.error_code, started)
            return error(self.config.contract_version, exc.error_code, exc.message)
        except Exception as exc:  # pragma: no cover
            self._log_failure(event, file_path, ErrorCode.INTERNAL_ERROR, started)
            return error(
                self.config.contract_version,
                ErrorCode.INTERNAL_ERROR,
                f"Unexpected error searching text: {exc}",
            )

    def insert_paragraphs(
        self,
        file_path: str,
        after_paragraph_index: int,
        paragraphs: list[dict[str, Any]],
    ) -> dict[str, Any]:
        event = "insert_paragraphs"
        started = time.perf_counter()
        try:
            normalized = self._normalize_paragraph_inputs(paragraphs)
            path = self._resolve_docx_path(file_path, must_exist=True)

            def mutator(doc: Document) -> dict[str, Any]:
                if after_paragraph_index < -1 or after_paragraph_index >= len(doc.paragraphs):
                    raise ServiceError(
                        ErrorCode.PARAGRAPH_INDEX_OUT_OF_RANGE,
                        f"after_paragraph_index {after_paragraph_index} is out of range for document with {len(doc.paragraphs)} paragraphs",
                    )
                anchor = doc.paragraphs[after_paragraph_index] if after_paragraph_index >= 0 else None
                for item in normalized:
                    new_paragraph = doc.add_paragraph(item["text"])
                    if item["style_hint"]:
                        self._try_apply_style(doc, new_paragraph, item["style_hint"])
                    elif anchor is not None and anchor.style is not None:
                        new_paragraph.style = anchor.style
                    self._insert_paragraph_after_anchor(doc, new_paragraph, anchor)
                    anchor = new_paragraph
                return {"inserted_count": len(normalized)}

            result = self._mutate_document(path, mutator)
            self._log_success(event, path, started)
            return ok(self.config.contract_version, **result)
        except ServiceError as exc:
            self._log_failure(event, file_path, exc.error_code, started)
            return error(self.config.contract_version, exc.error_code, exc.message)
        except Exception as exc:  # pragma: no cover
            self._log_failure(event, file_path, ErrorCode.INTERNAL_ERROR, started)
            return error(
                self.config.contract_version,
                ErrorCode.INTERNAL_ERROR,
                f"Unexpected error inserting paragraphs: {exc}",
            )

    def delete_paragraph_range(
        self,
        file_path: str,
        start_index: int,
        end_index: int,
    ) -> dict[str, Any]:
        event = "delete_paragraph_range"
        started = time.perf_counter()
        try:
            path = self._resolve_docx_path(file_path, must_exist=True)

            def mutator(doc: Document) -> dict[str, Any]:
                first, last = self._normalize_range(start_index, end_index, len(doc.paragraphs))
                deleted = 0
                for idx in range(last, first - 1, -1):
                    paragraph = doc.paragraphs[idx]
                    paragraph._element.getparent().remove(paragraph._element)
                    deleted += 1
                return {"deleted_count": deleted}

            result = self._mutate_document(path, mutator)
            self._log_success(event, path, started)
            return ok(self.config.contract_version, **result)
        except ServiceError as exc:
            self._log_failure(event, file_path, exc.error_code, started)
            return error(self.config.contract_version, exc.error_code, exc.message)
        except Exception as exc:  # pragma: no cover
            self._log_failure(event, file_path, ErrorCode.INTERNAL_ERROR, started)
            return error(
                self.config.contract_version,
                ErrorCode.INTERNAL_ERROR,
                f"Unexpected error deleting paragraph range: {exc}",
            )

    def search_and_replace(
        self,
        file_path: str,
        find_text: str,
        replace_text: str,
        max_replacements: int | None = None,
    ) -> dict[str, Any]:
        event = "search_and_replace"
        started = time.perf_counter()
        try:
            if not find_text:
                raise ServiceError(ErrorCode.INVALID_ARGUMENT, "find_text must be a non-empty string")
            if max_replacements is not None and max_replacements <= 0:
                raise ServiceError(ErrorCode.INVALID_ARGUMENT, "max_replacements must be > 0 when provided")
            path = self._resolve_docx_path(file_path, must_exist=True)

            def mutator(doc: Document) -> dict[str, Any]:
                replaced = 0
                for paragraph in doc.paragraphs:
                    if find_text not in paragraph.text:
                        continue
                    if max_replacements is None:
                        local_count = paragraph.text.count(find_text)
                        if local_count == 0:
                            continue
                        updated = paragraph.text.replace(find_text, replace_text)
                    else:
                        if replaced >= max_replacements:
                            break
                        remaining = max_replacements - replaced
                        local_count = paragraph.text.count(find_text)
                        if local_count == 0:
                            continue
                        local_to_apply = min(local_count, remaining)
                        updated = paragraph.text.replace(find_text, replace_text, local_to_apply)
                    self._set_paragraph_text_preserve_style(paragraph, updated)
                    replaced += local_count if max_replacements is None else min(local_count, remaining)
                return {"replacements": replaced}

            result = self._mutate_document(path, mutator)
            self._log_success(event, path, started)
            return ok(self.config.contract_version, **result)
        except ServiceError as exc:
            self._log_failure(event, file_path, exc.error_code, started)
            return error(self.config.contract_version, exc.error_code, exc.message)
        except Exception as exc:  # pragma: no cover
            self._log_failure(event, file_path, ErrorCode.INTERNAL_ERROR, started)
            return error(
                self.config.contract_version,
                ErrorCode.INTERNAL_ERROR,
                f"Unexpected error in search_and_replace: {exc}",
            )

    def replace_section_content(
        self,
        file_path: str,
        selector: dict[str, Any],
        new_paragraphs: list[str | dict[str, Any]],
        preserve_style: bool = True,
        dry_run: bool = False,
    ) -> dict[str, Any]:
        event = "replace_section_content"
        started = time.perf_counter()
        try:
            path = self._resolve_docx_path(file_path, must_exist=True)
            normalized = self._normalize_new_section_paragraphs(new_paragraphs)

            def compute_preview(doc: Document) -> dict[str, Any]:
                resolved = self._resolve_selector_range(doc, selector)
                existing = [doc.paragraphs[i].text for i in range(resolved.start_index, resolved.end_exclusive)]
                preview = {
                    "existing_paragraphs": existing,
                    "new_paragraphs": [item["text"] for item in normalized],
                }
                replaced_range = {
                    "start_index": resolved.start_index,
                    "end_index": resolved.end_exclusive - 1,
                    "selector_mode": resolved.selector_mode,
                    "selector_details": resolved.selector_details,
                }
                return {"preview": preview, "replaced_range": replaced_range, "resolved": resolved}

            if dry_run:
                doc = Document(str(path))
                payload = compute_preview(doc)
                self._log_success(event, path, started)
                return ok(
                    self.config.contract_version,
                    replaced_range=payload["replaced_range"],
                    preview=payload["preview"],
                )

            def mutator(doc: Document) -> dict[str, Any]:
                payload = compute_preview(doc)
                resolved: SelectorRange = payload["resolved"]
                templates = (
                    self._capture_style_templates(doc, resolved.start_index, resolved.end_exclusive)
                    if preserve_style
                    else []
                )
                if preserve_style and not templates:
                    templates = self._fallback_style_templates(doc, resolved.start_index)

                for idx in range(resolved.end_exclusive - 1, resolved.start_index - 1, -1):
                    paragraph = doc.paragraphs[idx]
                    paragraph._element.getparent().remove(paragraph._element)

                anchor = doc.paragraphs[resolved.start_index - 1] if resolved.start_index > 0 and doc.paragraphs else None
                for idx, item in enumerate(normalized):
                    template = self._select_template_for_index(templates, idx)
                    new_paragraph = doc.add_paragraph()
                    self._apply_template(doc, new_paragraph, template, item["style_hint"] if preserve_style else item["style_hint"])
                    run = new_paragraph.add_run(item["text"])
                    if preserve_style and template is not None:
                        self._apply_run_style(run, template.run_style)
                    self._insert_paragraph_after_anchor(doc, new_paragraph, anchor)
                    anchor = new_paragraph

                return {
                    "replaced_range": payload["replaced_range"],
                    "preview": payload["preview"],
                }

            result = self._mutate_document(path, mutator)
            self._log_success(event, path, started)
            return ok(self.config.contract_version, **result)
        except ServiceError as exc:
            self._log_failure(event, file_path, exc.error_code, started)
            return error(self.config.contract_version, exc.error_code, exc.message)
        except Exception as exc:  # pragma: no cover
            self._log_failure(event, file_path, ErrorCode.INTERNAL_ERROR, started)
            return error(
                self.config.contract_version,
                ErrorCode.INTERNAL_ERROR,
                f"Unexpected error replacing section content: {exc}",
            )

    def save_as(self, file_path: str, output_path: str) -> dict[str, Any]:
        event = "save_as"
        started = time.perf_counter()
        try:
            source = self._resolve_docx_path(file_path, must_exist=True)
            destination = self._resolve_docx_path(output_path, must_exist=False)
            with self.lock_manager.acquire_many([source, destination]):
                doc = Document(str(source))
                self._save_document_atomic(destination, doc)
            self._log_success(event, destination, started)
            return ok(self.config.contract_version, output_path=str(destination))
        except ServiceError as exc:
            self._log_failure(event, file_path, exc.error_code, started)
            return error(self.config.contract_version, exc.error_code, exc.message)
        except Exception as exc:  # pragma: no cover
            self._log_failure(event, file_path, ErrorCode.INTERNAL_ERROR, started)
            return error(
                self.config.contract_version,
                ErrorCode.INTERNAL_ERROR,
                f"Unexpected error saving copy: {exc}",
            )

    def convert_to_pdf(self, file_path: str, output_path: str | None = None) -> dict[str, Any]:
        event = "convert_to_pdf"
        started = time.perf_counter()
        try:
            source = self._resolve_docx_path(file_path, must_exist=True)
            destination = self._resolve_pdf_path(
                output_path or str(source.with_suffix(".pdf")),
                must_exist=False,
            )
            with self.lock_manager.acquire_many([source, destination]):
                method = self._perform_pdf_conversion(source, destination)
            self._log_success(event, destination, started)
            return ok(
                self.config.contract_version,
                output_path=str(destination),
                method=method,
                experimental=True,
            )
        except ServiceError as exc:
            self._log_failure(event, file_path, exc.error_code, started)
            return error(self.config.contract_version, exc.error_code, exc.message)
        except Exception as exc:  # pragma: no cover
            self._log_failure(event, file_path, ErrorCode.INTERNAL_ERROR, started)
            return error(
                self.config.contract_version,
                ErrorCode.INTERNAL_ERROR,
                f"Unexpected error converting to PDF: {exc}",
            )

    def get_document_comments(self, file_path: str) -> dict[str, Any]:
        event = "get_document_comments"
        started = time.perf_counter()
        try:
            source = self._resolve_docx_path(file_path, must_exist=True)
            comments_xml = self._read_docx_xml_part(source, "word/comments.xml")
            if comments_xml is None:
                self._log_success(event, source, started)
                return ok(
                    self.config.contract_version,
                    comments=[],
                    total_comments=0,
                    experimental=True,
                )

            try:
                root = ET.fromstring(comments_xml)
            except ET.ParseError as exc:
                raise ServiceError(ErrorCode.DOCX_ERROR, f"invalid comments.xml: {exc}") from exc

            comments: list[dict[str, Any]] = []
            id_key = f"{{{WORDPROCESSINGML_NS}}}id"
            date_key = f"{{{WORDPROCESSINGML_NS}}}date"
            author_key = f"{{{WORDPROCESSINGML_NS}}}author"
            initials_key = f"{{{WORDPROCESSINGML_NS}}}initials"

            for comment in root.findall("w:comment", WORD_NS_MAP):
                comment_text = self._extract_wordml_text(comment)
                raw_id = comment.attrib.get(id_key, "")
                try:
                    sort_id = int(raw_id)
                except ValueError:
                    sort_id = 10**9
                comments.append(
                    {
                        "id": raw_id,
                        "author": comment.attrib.get(author_key),
                        "initials": comment.attrib.get(initials_key),
                        "date": comment.attrib.get(date_key),
                        "text": comment_text,
                        "_sort_id": sort_id,
                    }
                )

            comments.sort(key=lambda item: (item["_sort_id"], item["id"]))
            for item in comments:
                item.pop("_sort_id", None)

            self._log_success(event, source, started)
            return ok(
                self.config.contract_version,
                comments=comments,
                total_comments=len(comments),
                experimental=True,
            )
        except ServiceError as exc:
            self._log_failure(event, file_path, exc.error_code, started)
            return error(self.config.contract_version, exc.error_code, exc.message)
        except Exception as exc:  # pragma: no cover
            self._log_failure(event, file_path, ErrorCode.INTERNAL_ERROR, started)
            return error(
                self.config.contract_version,
                ErrorCode.INTERNAL_ERROR,
                f"Unexpected error reading comments: {exc}",
            )

    def get_document_footnotes(self, file_path: str) -> dict[str, Any]:
        event = "get_document_footnotes"
        started = time.perf_counter()
        try:
            source = self._resolve_docx_path(file_path, must_exist=True)
            footnotes_xml = self._read_docx_xml_part(source, "word/footnotes.xml")
            if footnotes_xml is None:
                self._log_success(event, source, started)
                return ok(
                    self.config.contract_version,
                    footnotes=[],
                    total_footnotes=0,
                    experimental=True,
                )

            try:
                root = ET.fromstring(footnotes_xml)
            except ET.ParseError as exc:
                raise ServiceError(ErrorCode.DOCX_ERROR, f"invalid footnotes.xml: {exc}") from exc

            id_key = f"{{{WORDPROCESSINGML_NS}}}id"
            type_key = f"{{{WORDPROCESSINGML_NS}}}type"
            footnotes: list[dict[str, Any]] = []

            for footnote in root.findall("w:footnote", WORD_NS_MAP):
                raw_id = footnote.attrib.get(id_key, "")
                footnote_type = footnote.attrib.get(type_key)
                try:
                    footnote_id = int(raw_id)
                except ValueError:
                    # Non-integer IDs are treated as non-user footnotes.
                    continue
                # Skip system entries like separator/continuationSeparator.
                if footnote_id <= 0 or footnote_type is not None:
                    continue

                footnotes.append(
                    {
                        "id": raw_id,
                        "text": self._extract_wordml_text(footnote),
                    }
                )

            footnotes.sort(key=lambda item: int(item["id"]))
            self._log_success(event, source, started)
            return ok(
                self.config.contract_version,
                footnotes=footnotes,
                total_footnotes=len(footnotes),
                experimental=True,
            )
        except ServiceError as exc:
            self._log_failure(event, file_path, exc.error_code, started)
            return error(self.config.contract_version, exc.error_code, exc.message)
        except Exception as exc:  # pragma: no cover
            self._log_failure(event, file_path, ErrorCode.INTERNAL_ERROR, started)
            return error(
                self.config.contract_version,
                ErrorCode.INTERNAL_ERROR,
                f"Unexpected error reading footnotes: {exc}",
            )

    def list_available_documents(self, directory: str = ".") -> dict[str, Any]:
        event = "list_available_documents"
        started = time.perf_counter()
        try:
            path = self._resolve_directory_path(directory)
            files = [
                {"name": entry.name, "path": str(entry.resolve()), "size_bytes": entry.stat().st_size}
                for entry in sorted(path.iterdir(), key=lambda p: p.name.lower())
                if entry.is_file() and entry.suffix.lower() == ".docx"
            ]
            self._log_success(event, path, started)
            return ok(self.config.contract_version, files=files)
        except ServiceError as exc:
            self._log_failure(event, directory, exc.error_code, started)
            return error(self.config.contract_version, exc.error_code, exc.message)
        except Exception as exc:  # pragma: no cover
            self._log_failure(event, directory, ErrorCode.INTERNAL_ERROR, started)
            return error(
                self.config.contract_version,
                ErrorCode.INTERNAL_ERROR,
                f"Unexpected error listing documents: {exc}",
            )

    # Internal helpers -----------------------------------------------------------

    def _resolve_docx_path(self, value: str, must_exist: bool) -> Path:
        if not value or not isinstance(value, str):
            raise ServiceError(ErrorCode.INVALID_PATH, "file path must be a non-empty string")
        path = Path(value).expanduser()
        if not path.is_absolute():
            path = (Path.cwd() / path).resolve()
        else:
            path = path.resolve()
        if path.suffix.lower() != ".docx":
            raise ServiceError(ErrorCode.INVALID_PATH, "only .docx files are supported")
        self._ensure_path_allowed(path)
        if must_exist and not path.exists():
            raise ServiceError(ErrorCode.FILE_NOT_FOUND, f"document does not exist: {path}")
        if path.exists():
            self._check_file_size(path)
        return path

    def _read_docx_xml_part(self, source: Path, part_name: str) -> str | None:
        try:
            with zipfile.ZipFile(source, "r") as archive:
                names = set(archive.namelist())
                if part_name not in names:
                    return None
                return archive.read(part_name).decode("utf-8")
        except zipfile.BadZipFile as exc:
            raise ServiceError(ErrorCode.DOCX_ERROR, f"invalid docx archive: {exc}") from exc
        except UnicodeDecodeError as exc:
            raise ServiceError(ErrorCode.DOCX_ERROR, f"invalid XML encoding in {part_name}: {exc}") from exc

    def _extract_wordml_text(self, root: ET.Element) -> str:
        parts: list[str] = []
        for node in root.findall(".//w:t", WORD_NS_MAP):
            if node.text:
                parts.append(node.text)
        return "".join(parts).strip()

    def _resolve_pdf_path(self, value: str, must_exist: bool) -> Path:
        if not value or not isinstance(value, str):
            raise ServiceError(ErrorCode.INVALID_PATH, "output path must be a non-empty string")
        path = Path(value).expanduser()
        if not path.is_absolute():
            path = (Path.cwd() / path).resolve()
        else:
            path = path.resolve()
        if path.suffix.lower() != ".pdf":
            raise ServiceError(ErrorCode.INVALID_PATH, "output path must end with .pdf")
        self._ensure_path_allowed(path)
        if must_exist and not path.exists():
            raise ServiceError(ErrorCode.FILE_NOT_FOUND, f"file does not exist: {path}")
        if path.exists():
            self._check_file_size(path)
        return path

    def _resolve_directory_path(self, value: str) -> Path:
        if not value:
            raise ServiceError(ErrorCode.INVALID_PATH, "directory must be a non-empty string")
        path = Path(value).expanduser()
        if not path.is_absolute():
            path = (Path.cwd() / path).resolve()
        else:
            path = path.resolve()
        self._ensure_path_allowed(path)
        if not path.exists() or not path.is_dir():
            raise ServiceError(ErrorCode.FILE_NOT_FOUND, f"directory does not exist: {path}")
        return path

    def _ensure_path_allowed(self, path: Path) -> None:
        for allowed_root in self.config.normalized_allowed_roots():
            if path == allowed_root or path.is_relative_to(allowed_root):
                return
        raise ServiceError(
            ErrorCode.PATH_NOT_ALLOWED,
            f"path is outside allowed roots: {path}",
        )

    def _check_file_size(self, path: Path) -> None:
        size = path.stat().st_size
        if size > self.config.max_file_size_bytes:
            raise ServiceError(
                ErrorCode.FILE_TOO_LARGE,
                f"file exceeds max size ({self.config.max_file_size_bytes} bytes): {path}",
            )

    def _default_copy_path(self, source: Path) -> Path:
        stem = source.stem
        candidate = source.with_name(f"{stem}-copy{source.suffix}")
        counter = 2
        while candidate.exists():
            candidate = source.with_name(f"{stem}-copy-{counter}{source.suffix}")
            counter += 1
        return candidate

    def _paragraph_by_index(self, doc: Document, index: int):
        if index < 0 or index >= len(doc.paragraphs):
            raise ServiceError(
                ErrorCode.PARAGRAPH_INDEX_OUT_OF_RANGE,
                f"paragraph_index {index} is out of range for {len(doc.paragraphs)} paragraphs",
            )
        return doc.paragraphs[index]

    def _extract_heading_level(self, style_name: str | None) -> int | None:
        if not style_name:
            return None
        match = HEADING_STYLE_PATTERN.match(style_name.strip())
        if not match:
            return None
        return int(match.group(1))

    def _normalize_paragraph_inputs(self, paragraphs: list[dict[str, Any]]) -> list[dict[str, Any]]:
        if not isinstance(paragraphs, list) or not paragraphs:
            raise ServiceError(ErrorCode.INVALID_ARGUMENT, "paragraphs must be a non-empty list")
        normalized = []
        for item in paragraphs:
            if not isinstance(item, dict):
                raise ServiceError(ErrorCode.INVALID_ARGUMENT, "each paragraph input must be an object")
            text = item.get("text")
            if text is None or not isinstance(text, str):
                raise ServiceError(ErrorCode.INVALID_ARGUMENT, "paragraph text must be a string")
            normalized.append({"text": text, "style_hint": item.get("style_hint")})
        return normalized

    def _normalize_new_section_paragraphs(
        self,
        new_paragraphs: list[str | dict[str, Any]],
    ) -> list[dict[str, Any]]:
        if not isinstance(new_paragraphs, list):
            raise ServiceError(ErrorCode.INVALID_ARGUMENT, "new_paragraphs must be a list")
        normalized: list[dict[str, Any]] = []
        for item in new_paragraphs:
            if isinstance(item, str):
                normalized.append({"text": item, "style_hint": None})
            elif isinstance(item, dict):
                text = item.get("text")
                if text is None or not isinstance(text, str):
                    raise ServiceError(ErrorCode.INVALID_ARGUMENT, "new paragraph object text must be a string")
                normalized.append({"text": text, "style_hint": item.get("style_hint")})
            else:
                raise ServiceError(ErrorCode.INVALID_ARGUMENT, "new_paragraphs entries must be string or object")
        return normalized

    def _normalize_range(self, start_index: int, end_index: int, total: int) -> tuple[int, int]:
        if total == 0:
            raise ServiceError(ErrorCode.PARAGRAPH_INDEX_OUT_OF_RANGE, "document has no paragraphs")
        first = min(start_index, end_index)
        last = max(start_index, end_index)
        if first < 0 or last >= total:
            raise ServiceError(
                ErrorCode.PARAGRAPH_INDEX_OUT_OF_RANGE,
                f"range [{start_index}, {end_index}] is out of bounds for {total} paragraphs",
            )
        return first, last

    def _mutate_document(self, path: Path, mutator):
        with self.lock_manager.acquire(path):
            doc = Document(str(path))
            payload = mutator(doc)
            self._save_document_atomic(path, doc)
            return payload

    def _perform_pdf_conversion(self, source: Path, destination: Path) -> str:
        errors: list[str] = []

        if importlib.util.find_spec("docx2pdf") is not None:
            try:
                from docx2pdf import convert  # type: ignore

                convert(str(source), str(destination))
                if destination.exists() and destination.stat().st_size > 0:
                    return "docx2pdf"
                errors.append("docx2pdf ran but did not produce a valid output file")
            except Exception as exc:  # pragma: no cover - backend/platform specific
                errors.append(f"docx2pdf failed: {exc}")
        else:
            errors.append("docx2pdf not installed")

        outdir = destination.parent
        expected_libreoffice_output = outdir / f"{source.stem}.pdf"
        for command in ("soffice", "libreoffice"):
            try:
                result = subprocess.run(
                    [
                        command,
                        "--headless",
                        "--convert-to",
                        "pdf",
                        "--outdir",
                        str(outdir),
                        str(source),
                    ],
                    check=False,
                    capture_output=True,
                    text=True,
                    timeout=120,
                )
                if result.returncode != 0:
                    stderr = result.stderr.strip() or "unknown error"
                    errors.append(f"{command} failed: {stderr}")
                    continue
                if expected_libreoffice_output.exists():
                    if expected_libreoffice_output != destination:
                        os.replace(str(expected_libreoffice_output), str(destination))
                    if destination.exists() and destination.stat().st_size > 0:
                        return command
                errors.append(f"{command} did not produce expected output")
            except FileNotFoundError:
                errors.append(f"{command} not installed")
            except Exception as exc:  # pragma: no cover - backend/platform specific
                errors.append(f"{command} conversion error: {exc}")

        raise ServiceError(
            ErrorCode.DOCX_ERROR,
            "PDF conversion failed: " + "; ".join(errors),
        )

    def _save_document_atomic(self, destination: Path, doc: Document) -> None:
        temp_file_path: Path | None = None
        try:
            with tempfile.NamedTemporaryFile(
                suffix=".docx",
                prefix=f".{destination.stem}-",
                dir=str(destination.parent),
                delete=False,
            ) as handle:
                temp_file_path = Path(handle.name)
            doc.save(str(temp_file_path))
            os.replace(str(temp_file_path), str(destination))
        except Exception as exc:
            if temp_file_path is not None and temp_file_path.exists():
                temp_file_path.unlink(missing_ok=True)
            raise ServiceError(ErrorCode.DOCX_ERROR, f"failed to save document atomically: {exc}") from exc

    def _set_paragraph_text_preserve_style(self, paragraph, text: str) -> None:
        template = self._capture_template(paragraph)
        for run in list(paragraph.runs):
            paragraph._element.remove(run._element)
        run = paragraph.add_run(text)
        if template is not None:
            self._apply_run_style(run, template.run_style)

    def _resolve_selector_range(self, doc: Document, selector: dict[str, Any]) -> SelectorRange:
        if not isinstance(selector, dict):
            raise ServiceError(ErrorCode.INVALID_ARGUMENT, "selector must be an object")
        mode = selector.get("mode")
        if mode == "heading_exact":
            value = selector.get("value")
            occurrence = selector.get("occurrence", 1)
            if not value or not isinstance(value, str):
                raise ServiceError(ErrorCode.INVALID_ARGUMENT, "heading_exact selector requires string value")
            if not isinstance(occurrence, int) or occurrence <= 0:
                raise ServiceError(ErrorCode.INVALID_ARGUMENT, "occurrence must be a positive integer")
            matches = [idx for idx, para in enumerate(doc.paragraphs) if para.text.strip() == value.strip()]
            if len(matches) < occurrence:
                raise ServiceError(
                    ErrorCode.SELECTOR_NOT_FOUND,
                    f"heading not found for occurrence {occurrence}: {value}",
                )
            heading_index = matches[occurrence - 1]
            start = heading_index + 1
            end = len(doc.paragraphs)
            for idx in range(start, len(doc.paragraphs)):
                if self._extract_heading_level(doc.paragraphs[idx].style.name if doc.paragraphs[idx].style else None) is not None:
                    end = idx
                    break
            return SelectorRange(
                start_index=start,
                end_exclusive=end,
                selector_mode=mode,
                selector_details={"value": value, "occurrence": occurrence, "heading_index": heading_index},
            )
        if mode == "anchors":
            start_text = selector.get("start_text")
            end_text = selector.get("end_text")
            if not start_text or not isinstance(start_text, str):
                raise ServiceError(ErrorCode.INVALID_ARGUMENT, "anchors selector requires start_text")
            if not end_text or not isinstance(end_text, str):
                raise ServiceError(ErrorCode.INVALID_ARGUMENT, "anchors selector requires end_text")

            start_index = None
            for idx, para in enumerate(doc.paragraphs):
                if para.text.strip() == start_text.strip():
                    start_index = idx
                    break
            if start_index is None:
                raise ServiceError(ErrorCode.SELECTOR_NOT_FOUND, f"start anchor not found: {start_text}")

            end_index = None
            for idx in range(start_index + 1, len(doc.paragraphs)):
                if doc.paragraphs[idx].text.strip() == end_text.strip():
                    end_index = idx
                    break
            if end_index is None:
                raise ServiceError(ErrorCode.SELECTOR_NOT_FOUND, f"end anchor not found: {end_text}")

            return SelectorRange(
                start_index=start_index + 1,
                end_exclusive=end_index,
                selector_mode=mode,
                selector_details={"start_text": start_text, "end_text": end_text},
            )
        raise ServiceError(ErrorCode.INVALID_ARGUMENT, f"unsupported selector mode: {mode}")

    def _capture_template(self, paragraph) -> ParagraphStyleTemplate | None:
        if paragraph is None:
            return None
        run = paragraph.runs[0] if paragraph.runs else None
        paragraph_format = paragraph.paragraph_format
        return ParagraphStyleTemplate(
            style_name=paragraph.style.name if paragraph.style else None,
            alignment=paragraph.alignment,
            left_indent=paragraph_format.left_indent,
            right_indent=paragraph_format.right_indent,
            first_line_indent=paragraph_format.first_line_indent,
            space_before=paragraph_format.space_before,
            space_after=paragraph_format.space_after,
            line_spacing=paragraph_format.line_spacing,
            line_spacing_rule=paragraph_format.line_spacing_rule,
            run_style=RunStyleTemplate(
                bold=run.bold if run else None,
                italic=run.italic if run else None,
                underline=run.underline if run else None,
                font_name=run.font.name if run else None,
                font_size=run.font.size if run else None,
                font_color=run.font.color.rgb if run and run.font.color else None,
            ),
        )

    def _capture_style_templates(
        self,
        doc: Document,
        start_index: int,
        end_exclusive: int,
    ) -> list[ParagraphStyleTemplate]:
        templates: list[ParagraphStyleTemplate] = []
        for idx in range(start_index, end_exclusive):
            template = self._capture_template(doc.paragraphs[idx])
            if template is not None:
                templates.append(template)
        return templates

    def _fallback_style_templates(self, doc: Document, start_index: int) -> list[ParagraphStyleTemplate]:
        if start_index > 0:
            template = self._capture_template(doc.paragraphs[start_index - 1])
            if template is not None:
                return [template]
        return [ParagraphStyleTemplate(style_name="Normal")]

    def _select_template_for_index(
        self,
        templates: list[ParagraphStyleTemplate],
        index: int,
    ) -> ParagraphStyleTemplate | None:
        if not templates:
            return None
        if index < len(templates):
            return templates[index]
        return templates[-1]

    def _apply_template(
        self,
        doc: Document,
        paragraph,
        template: ParagraphStyleTemplate | None,
        style_hint: str | None,
    ) -> None:
        if style_hint:
            self._try_apply_style(doc, paragraph, style_hint)
        elif template is not None and template.style_name:
            self._try_apply_style(doc, paragraph, template.style_name)
        elif paragraph.style is None:
            self._try_apply_style(doc, paragraph, "Normal")

        if template is None:
            return
        paragraph.alignment = template.alignment
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = template.left_indent
        paragraph_format.right_indent = template.right_indent
        paragraph_format.first_line_indent = template.first_line_indent
        paragraph_format.space_before = template.space_before
        paragraph_format.space_after = template.space_after
        paragraph_format.line_spacing = template.line_spacing
        paragraph_format.line_spacing_rule = template.line_spacing_rule

    def _apply_run_style(self, run, template: RunStyleTemplate) -> None:
        run.bold = template.bold
        run.italic = template.italic
        run.underline = template.underline
        if template.font_name:
            run.font.name = template.font_name
        if template.font_size:
            run.font.size = template.font_size
        if template.font_color:
            run.font.color.rgb = template.font_color

    def _try_apply_style(self, doc: Document, paragraph, style_name: str) -> bool:
        try:
            paragraph.style = style_name
            return True
        except Exception:
            available = {style.name for style in doc.styles}
            if "Normal" in available:
                paragraph.style = "Normal"
            return False

    def _insert_paragraph_after_anchor(self, doc: Document, paragraph, anchor) -> None:
        if anchor is None:
            body = doc._body._body
            body.insert(0, paragraph._element)
            return
        anchor._element.addnext(paragraph._element)

    def _log_success(self, event: str, file_path: Path | str, started: float) -> None:
        duration_ms = int((time.perf_counter() - started) * 1000)
        logger.info(
            "word_engine_success",
            extra={
                "event": event,
                "file_path": str(file_path),
                "status": "ok",
                "duration_ms": duration_ms,
            },
        )

    def _log_failure(
        self,
        event: str,
        file_path: Path | str,
        error_code: ErrorCode,
        started: float,
    ) -> None:
        duration_ms = int((time.perf_counter() - started) * 1000)
        logger.error(
            "word_engine_failure",
            extra={
                "event": event,
                "file_path": str(file_path),
                "status": "error",
                "error_code": error_code.value,
                "duration_ms": duration_ms,
            },
        )
